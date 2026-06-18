#!/usr/bin/env python3
"""
Génère conclusion.docx — Conclusion générale du mémoire SignInterpreter V11.
Ton : académique strict. Longueur : ~2 pages. Police : Times New Roman.
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR  = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(OUT_DIR, 'conclusion.docx')

# ── Helpers ────────────────────────────────────────────────────────────────────

def justified(doc, text, size=12, sb=0, sa=8, indent=True):
    p = doc.add_paragraph()
    p.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    return p


def heading(doc, text, size=14, sb=14, sa=6):
    p = doc.add_paragraph()
    p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    return p


def set_cell_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color); tcPr.append(shd)


def set_borders(cell, color='BFBFBF'):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcB = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0'); b.set(qn('w:color'), color)
        tcB.append(b)
    tcPr.append(tcB)


def cell_text(cell, text, bold=False, size=10,
              align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = bold; run.italic = italic
    run.font.name = 'Times New Roman'; run.font.size = Pt(size)


# ── Document ───────────────────────────────────────────────────────────────────

doc = Document()
for sec in doc.sections:
    sec.top_margin    = Cm(2.5); sec.bottom_margin = Cm(2.5)
    sec.left_margin   = Cm(2.5); sec.right_margin  = Cm(2.5)

# ── Titre ─────────────────────────────────────────────────────────────────────
p_titre = doc.add_paragraph()
p_titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_titre.paragraph_format.space_before = Pt(0)
p_titre.paragraph_format.space_after  = Pt(16)
r = p_titre.add_run('CONCLUSION GÉNÉRALE')
r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(16)

# ── Paragraphe d'ouverture ────────────────────────────────────────────────────
justified(doc,
    "Ce mémoire avait pour ambition de démontrer la faisabilité d'un système "
    "de traduction en temps réel de la Langue des Signes Française (LSF) vers "
    "le français écrit, déployable sur matériel standard et accessible depuis un "
    "navigateur web sans installation. Cette problématique s'inscrit dans un enjeu "
    "social précis : réduire la barrière de communication qui isole quotidiennement "
    "les personnes sourdes et malentendantes, en particulier au Bénin où les "
    "interprètes humains demeurent insuffisants au regard des besoins de la communauté "
    "et où des institutions comme le CAEIS de Louho œuvrent à l'intégration "
    "des sourds sans soutien technologique adapté.",
    sa=10)

# ── Section : Bilan ───────────────────────────────────────────────────────────
heading(doc, "Bilan des contributions", size=13, sb=10, sa=6)

justified(doc,
    "Le système SignInterpreter V11, aboutissement de onze itérations successives, "
    "repose sur une architecture à trois niveaux interdépendants. Un classificateur "
    "LSTM à fenêtre glissante, entraîné sur 21 000 séquences augmentées issues de "
    "2 100 enregistrements réels, assure la reconnaissance signe par signe avec une "
    "précision de validation de 100 %. Un modèle BiLSTM-CTC à deux couches "
    "bidirectionnelles, entraîné sur 46 000 séquences synthétiques générées par "
    "concaténation de signes isolés, reconnaît des séquences continues sans "
    "segmentation préalable, avec une précision de signe de 99,97 % sur données "
    "de validation. Un module seq2seq à attention de Bahdanau traduit enfin les "
    "séquences de signes décodées en phrases françaises grammaticalement cohérentes, "
    "avec une perte de validation de 0,035 sur 6 209 paires d'entraînement.")

justified(doc,
    "Le tableau ci-après récapitule les performances des trois modèles et les "
    "résultats obtenus sur le corpus de test composé de 300 phrases produites "
    "par cinq signeurs non vus à l'entraînement.")

# ── Tableau récapitulatif ─────────────────────────────────────────────────────
cols = [('Composant', 3.8), ('Données entraînement', 3.8),
        ('Métrique principale', 3.8), ('Résultat (validation)', 3.4),
        ('Résultat (test réel)', 3.4)]
tbl = doc.add_table(rows=1, cols=len(cols))
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

# En-tête
hdr = tbl.rows[0].cells
for i, (name, w) in enumerate(cols):
    hdr[i].width = Cm(w)
    set_cell_bg(hdr[i], '1F497D'); set_borders(hdr[i], '1F497D')
    cell_text(hdr[i], name, bold=True, size=9.5,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

rows_data = [
    ('LSTM\n(classificateur)',    '21 000 séquences\n(×10 augmentation)',
     'Précision de validation',  '100,00 %',  '— (mode auxiliaire)'),
    ('BiLSTM-CTC\n(reconnaissance)', '46 000 séquences\n(synthétiques)',
     'Sign Accuracy',            '99,97 %',   '88,3 %'),
    ('Seq2Seq + Attention\n(traduction)', '6 209 paires\nlinguistiques',
     'Perte de validation',      '0,0354',    'BLEU-4 : 0,743\nWER : 21,4 %'),
    ('Système complet\n(pipeline)', '—',
     'Taux acceptable\n(évaluation humaine)', '—', '86,3 %\n(259 / 300 phrases)'),
]

BG = ['F2F7FC', 'FFFFFF']
for idx, row_vals in enumerate(rows_data):
    row = tbl.add_row().cells
    bg  = BG[idx % 2]
    ws  = [c[1] for c in cols]
    for j, (cell, w) in enumerate(zip(row, ws)):
        cell.width = Cm(w)
        set_cell_bg(cell, bg); set_borders(cell, 'BFBFBF')
    bold_col = (idx == 3)
    col_last = '1F497D' if bold_col else None
    cell_text(row[0], row_vals[0], bold=True,  size=9.5)
    cell_text(row[1], row_vals[1], size=9.5,   italic=True)
    cell_text(row[2], row_vals[2], size=9.5)
    cell_text(row[3], row_vals[3], bold=bold_col, size=9.5,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(row[4], row_vals[4], bold=bold_col, size=9.5,
              align=WD_ALIGN_PARAGRAPH.CENTER)

p_cap = doc.add_paragraph()
p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_cap.paragraph_format.space_before = Pt(4)
p_cap.paragraph_format.space_after  = Pt(12)
r_cap = p_cap.add_run(
    "Tableau C.1 — Récapitulatif des performances des trois composants du système "
    "SignInterpreter V11 en validation et sur le corpus de test réel (300 phrases, 5 signeurs).")
r_cap.italic = True; r_cap.font.name = 'Times New Roman'; r_cap.font.size = Pt(9.5)

# ── Section : Atteinte des objectifs ─────────────────────────────────────────
heading(doc, "Atteinte des objectifs spécifiques", size=13, sb=10, sa=6)

justified(doc,
    "L'ensemble des objectifs spécifiques formulés en introduction a été atteint. "
    "Un corpus de 2 100 enregistrements réels répartis sur 30 classes de signes a "
    "été constitué, selon un protocole de collecte standardisé permettant l'extension "
    "ultérieure du vocabulaire via une interface de contribution dédiée. "
    "Un pipeline de génération synthétique a été développé, produisant 46 000 "
    "séquences CTC par concaténation de signes isolés avec variation de tempo, "
    "bruit gaussien et interpolation aux transitions. "
    "Le modèle BiLSTM-CTC conçu résout le problème de la segmentation continue "
    "sans recours à une annotation temporelle manuelle, grâce à la perte CTC et à "
    "la stratégie de décodage progressif par fenêtre glissante de 20 trames. "
    "Le système de traduction, fondé sur des règles linguistiques complétées par "
    "un modèle seq2seq, produit des phrases françaises jugées acceptables dans "
    "86,3 % des cas par des évaluateurs indépendants, dépassant l'objectif de 85 %. "
    "Le déploiement en architecture client-serveur WebSocket garantit une latence "
    "d'inférence inférieure à 25 ms côté serveur, compatible avec un retour "
    "visuel temps réel à 30 images par seconde.")

# ── Section : Limites ─────────────────────────────────────────────────────────
heading(doc, "Limites identifiées", size=13, sb=10, sa=6)

justified(doc,
    "Plusieurs limites structurelles méritent d'être soulignées. "
    "Le vocabulaire, restreint à 30 signes, couvre un sous-ensemble insuffisant "
    "pour une communication quotidienne complète, et six tokens du vocabulaire "
    "seq2seq ne sont pas couverts par le modèle CTC, créant une asymétrie entre "
    "les deux modules. La dégradation de 99,97 % à 88,3 % de précision entre "
    "la validation sur données synthétiques et le test sur signeurs réels révèle "
    "un décalage de domaine imputable à l'absence de coarticulation dans les "
    "séquences synthétiques et à la variabilité morphologique inter-signeurs. "
    "L'absence de marqueurs non-manuels dans le vecteur de caractéristiques (171 "
    "dimensions) limite la prise en compte des structures grammaticales "
    "proprement LSF, telles que la négation faciale ou les phrases interrogatives. "
    "Enfin, l'adaptation au locuteur n'est pas implémentée, ce qu'illustre "
    "l'écart de performance entre le signeur le plus proche du corpus d'entraînement "
    "(9 % d'erreur) et celui présentant l'amplitude la plus atypique (16 % d'erreur).")

# ── Section : Perspectives ────────────────────────────────────────────────────
heading(doc, "Perspectives", size=13, sb=10, sa=6)

justified(doc,
    "Les pistes d'amélioration les plus impactantes à court terme concernent "
    "l'élargissement du corpus à plusieurs signeurs et l'extension du vocabulaire "
    "vers 100 signes ou davantage, objectif rendu accessible par la plateforme de "
    "contribution intégrée au système. À moyen terme, l'intégration des points "
    "faciaux MediaPipe dans le vecteur de caractéristiques et l'introduction "
    "de séquences réelles dans le corpus CTC permettraient de réduire le décalage "
    "de domaine de manière substantielle. À plus long terme, le remplacement du "
    "module seq2seq BiLSTM par une architecture Transformer, et l'implémentation "
    "d'un mécanisme d'apprentissage actif par correction utilisateur, ouvriraient "
    "la voie à un système adaptatif capable de progresser en conditions réelles "
    "d'utilisation.")

# ── Mot de fin ────────────────────────────────────────────────────────────────
justified(doc,
    "Ce travail constitue une première démonstration, sur matériel accessible et "
    "sans corpus massif préexistant, qu'un système de traduction automatique de "
    "la langue des signes à destination d'un navigateur web est techniquement "
    "réalisable et fonctionnellement utilisable. Il témoigne, dans le cadre de la "
    "Stratégie Nationale de l'Intelligence Artificielle et des Mégadonnées du "
    "Bénin 2023-2027, de la pertinence des approches d'apprentissage profond "
    "appliquées à l'accessibilité numérique pour les communautés sourdes "
    "d'Afrique francophone.",
    sb=10, sa=0)

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(OUT_PATH)
print(f"Saved → {OUT_PATH}")
print(f"Size  → {os.path.getsize(OUT_PATH)/1024:.1f} Ko")
