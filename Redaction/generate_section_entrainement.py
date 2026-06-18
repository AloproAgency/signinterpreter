#!/usr/bin/env python3
"""Génère la section Protocole d'entraînement — CTC BiLSTM + Seq2Seq.
Toutes les valeurs sont vérifiées contre le code source."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/Users/alopro/Desktop/AI/RECHERCHE/SignInterpreter/V11/Redaction/section_entrainement.docx"

doc = Document()
sec = doc.sections[0]
sec.page_width  = Cm(21); sec.page_height = Cm(29.7)
sec.left_margin = sec.right_margin = Cm(2.5)
sec.top_margin  = sec.bottom_margin = Cm(2.5)

# ── helpers ───────────────────────────────────────────────────────────────────

def shd(cell, hex_c):
    tc = cell._tc; p = tc.get_or_add_tcPr()
    e = OxmlElement('w:shd')
    e.set(qn('w:val'), 'clear'); e.set(qn('w:color'), 'auto')
    e.set(qn('w:fill'), hex_c); p.append(e)

def h1(text, sb=20, sa=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    r = p.add_run(text); r.bold = True
    r.font.size = Pt(14); r.font.name = "Times New Roman"
    r.font.color.rgb = RGBColor.from_string("1F3864")

def h2(text, sb=14, sa=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    r = p.add_run(text); r.bold = True
    r.font.size = Pt(12); r.font.name = "Times New Roman"
    r.font.color.rgb = RGBColor.from_string("2E4057")

def j(text, indent=True, sb=0, sa=6):
    p = doc.add_paragraph()
    p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    r = p.add_run(text)
    r.font.size = Pt(12); r.font.name = "Times New Roman"

def j_mix(parts, indent=True, sb=0, sa=6):
    """parts = [(text, bold, italic)]"""
    p = doc.add_paragraph()
    p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    for text, bold, italic in parts:
        r = p.add_run(text)
        r.bold = bold; r.italic = italic
        r.font.size = Pt(12); r.font.name = "Times New Roman"

def dash(label, text, sb=0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1.5)
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(4)
    r1 = p.add_run("– " + label + " ")
    r1.bold = True; r1.font.size = Pt(12); r1.font.name = "Times New Roman"
    r2 = p.add_run(text)
    r2.font.size = Pt(12); r2.font.name = "Times New Roman"

def code_block(lines, sa=8):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(sa)
    r = p.add_run("\n".join(lines))
    r.font.name = "Courier New"; r.font.size = Pt(9.5)

def caption(text):
    p = doc.add_paragraph()
    p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(14)
    r = p.add_run(text); r.italic = True
    r.font.size = Pt(10); r.font.name = "Times New Roman"

def table(headers, rows, cap=None, widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shd(c, "1F3864")
        r = c.paragraphs[0].runs[0]
        r.bold = True; r.font.size = Pt(10)
        r.font.name = "Times New Roman"
        r.font.color.rgb = RGBColor(255, 255, 255)
    for ri, row in enumerate(rows):
        fill = "EBF0F7" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = str(val)
            shd(cell, fill)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10); run.font.name = "Times New Roman"
    if widths:
        for row in t.rows:
            for ci, w in enumerate(widths):
                row.cells[ci].width = Cm(w)
    if cap:
        caption(cap)
    return t

def hr():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '4')
    bot.set(qn('w:space'), '1'); bot.set(qn('w:color'), '2E4057')
    pb.append(bot); pPr.append(pb)

def encadre(title, lines):
    """Boîte encadrée pour les formules ou points clés."""
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    c = t.rows[0].cells[0]
    shd(c, "EBF0F7")
    # Title run
    p0 = c.paragraphs[0]
    p0.paragraph_format.space_before = Pt(4)
    p0.paragraph_format.space_after  = Pt(2)
    r0 = p0.add_run(title)
    r0.bold = True; r0.font.size = Pt(10); r0.font.name = "Times New Roman"
    r0.font.color.rgb = RGBColor.from_string("1F3864")
    # Content lines
    for line in lines:
        px = c.add_paragraph()
        px.paragraph_format.space_before = Pt(0)
        px.paragraph_format.space_after  = Pt(2)
        rx = px.add_run(line)
        rx.font.name = "Courier New"; rx.font.size = Pt(9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ══════════════════════════════════════════════════════════════════════════════
# TITRE
# ══════════════════════════════════════════════════════════════════════════════
tp = doc.add_paragraph()
tp.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.CENTER
tp.paragraph_format.space_before = Pt(0)
tp.paragraph_format.space_after  = Pt(6)
tr = tp.add_run("2.8  Protocole d'entraînement")
tr.bold = True; tr.font.size = Pt(16); tr.font.name = "Times New Roman"
tr.font.color.rgb = RGBColor.from_string("1F3864")
hr()

j("Cette section décrit les protocoles d'entraînement des deux modèles séquentiels principaux du système : le modèle CTC BiLSTM pour la reconnaissance continue des signes, et le modèle Seq2Seq à attention pour la traduction. Pour chaque modèle, nous précisons les contraintes techniques spécifiques rencontrées, les choix d'hyperparamètres et leurs justifications, ainsi que les résultats obtenus.", sb=4)

# ══════════════════════════════════════════════════════════════════════════════
# 2.8.1 — CTC BiLSTM
# ══════════════════════════════════════════════════════════════════════════════
h1("2.8.1  Entraînement du modèle CTC BiLSTM")

j("L'entraînement d'un modèle CTC diffère d'un entraînement supervisé classique sur deux aspects fondamentaux : la nature de la fonction de perte et la contrainte sur les longueurs de séquence.")

h2("Contrainte de longueur CTC")
j("La perte CTC (Connectionist Temporal Classification) calcule la vraisemblance d'une séquence cible en marginalisant sur tous les alignements temporels possibles entre la séquence d'entrée et la séquence de signes cible. Cette marginalisation impose une contrainte stricte sur les longueurs : la séquence d'entrée de T images doit satisfaire :")

encadre("Contrainte CTC :", [
    "T  ≥  2 × S − 1",
    "",
    "où  T = nombre d'images de la séquence d'entrée",
    "     S = nombre de signes dans la séquence cible",
    "",
    "Exemple : S = 5 signes  →  T ≥ 9 images minimum",
    "(intercaler un blank entre chaque signe : s₁ _ s₂ _ s₃ _ s₄ _ s₅)",
])

j("Dans le corpus synthétique, certaines séquences générées avec un facteur de tempo faible (0.50×) peuvent ne pas satisfaire cette contrainte pour des séquences cibles longues. Le paramètre zero_infinity=True de CTCLoss gère ces cas en remplaçant les pertes infinies par zéro plutôt que d'interrompre l'entraînement :")

code_block([
    "criterion = nn.CTCLoss(",
    "    blank        = N_CLASSES,    # index 30 (= nombre de classes)",
    "    reduction    = 'mean',",
    "    zero_infinity = True          # ignore les pertes infinies (contrainte violée)",
    ")",
])

h2("Stratégie de calcul sur Apple Silicon (MPS)")
j("PyTorch ne supporte pas l'opération CTCLoss sur le backend MPS d'Apple Silicon. La stratégie adoptée consiste à effectuer le calcul de la perte sur CPU, tandis que les paramètres du modèle et les passes forward restent sur le GPU MPS :")

code_block([
    "_CTC_ON_CPU = (DEVICE.type == 'mps')",
    "",
    "# Dans la boucle d'entraînement :",
    "log_p = model(padded, in_lens)              # forward sur MPS (rapide)",
    "loss  = criterion(",
    "    log_p.cpu() if _CTC_ON_CPU else log_p,  # perte sur CPU si MPS",
    "    tgt, in_lens, tgt_lens",
    ")",
    "loss.backward()                             # gradient remonte vers MPS",
])

j("Cette stratégie hybride permet de bénéficier de l'accélération MPS pour les opérations matricielles du BiLSTM tout en contournant la limitation de CTCLoss, sans surcoût significatif grâce au faible volume des transferts (logits CPU → gradient MPS uniquement).")

h2("Hyperparamètres d'entraînement")

table(
    ["Hyperparamètre", "Valeur", "Justification"],
    [
        ["Architecture",
         "BiLSTM — 2 couches\nhidden = 256 par direction",
         "Modélisation des dépendances temporelles dans les deux sens ;\n"
         "256 unités offrent une capacité suffisante sans saturer la mémoire M1."],
        ["Dimension d'entrée",
         "171 dims",
         "Vecteur de caractéristiques complet\n"
         "(39 pose + 66 main gauche + 66 main droite).\n"
         "Non enrichi avec les 6 dims de vitesse."],
        ["Token blank CTC",
         "Index 30\n(= n_classes)",
         "Convention standard PyTorch : le blank est placé\n"
         "à l'indice n_classes, après les 30 classes réelles."],
        ["Batch size",
         "32",
         "Contrainte mémoire : M1 16 Go unifiée CPU/GPU.\n"
         "Batchs plus grands risquent de saturer la mémoire partagée."],
        ["Optimiseur",
         "Adam\nlr = 1×10⁻³\nweight_decay = 1×10⁻⁵",
         "Adam : convergence rapide sur des données synthétiques variées.\n"
         "weight_decay : régularisation L2 légère pour limiter le sur-apprentissage."],
        ["Scheduler LR",
         "CosineAnnealingLR\nT_max = 80\nη_min = 1×10⁻⁵",
         "Décroissance cosinus : exploration large en début d'entraînement,\n"
         "affinage progressif vers η_min sans paliers brusques."],
        ["Gradient clipping",
         "‖g‖ ≤ 5.0",
         "Prévention de l'explosion du gradient, phénomène fréquent\n"
         "dans les LSTM profonds sur des séquences longues."],
        ["Arrêt anticipé",
         "patience = 15 époques\nsur sign_accuracy (val)",
         "Métrique de référence : sign accuracy = 1 − SER\n"
         "(Sign Error Rate, distance de Levenshtein normalisée).\n"
         "Pas de gain sur 15 époques consécutives → arrêt."],
        ["Époques maximum",
         "80",
         "Le CosineAnnealingLR est paramétré sur T_max = 80,\n"
         "cohérence avec la durée maximale d'entraînement."],
        ["Division train / val",
         "90 % / 10 %\n41 400 / 4 600 séquences",
         "Calculé sur 46 000 séquences synthétiques totales\n"
         "(6 000 mono-signe + 40 000 multi-signes).\n"
         "n_val = max(300, int(0.10 × 46 000)) = 4 600."],
    ],
    cap="Tableau 2.6 — Hyperparamètres d'entraînement du modèle CTC BiLSTM\n"
       "(vérifiés dans scripts/train_ctc.py)",
    widths=[4.0, 4.0, 8.5]
)

h2("Métrique d'évaluation : Sign Accuracy")
j("La métrique de validation n'est pas la perte CTC mais la Sign Accuracy, définie comme le complément du Sign Error Rate (SER). Le SER est calculé par la distance de Levenshtein entre la séquence de signes prédite (après décodage glouton) et la séquence cible, normalisée par la longueur de la séquence de référence :")

encadre("Sign Accuracy :", [
    "SER         =  Σ Levenshtein(prédit, référence)  /  Σ len(référence)",
    "Sign Acc    =  1 − SER",
    "",
    "Résultat obtenu :  Sign Accuracy = 99.97 %  (sur 4 600 séquences de validation)",
    "                   (source : data/ctc_model/config.json — best_sign_acc)",
])

j("Le choix de la Sign Accuracy comme critère d'arrêt anticipé (plutôt que la perte CTC) garantit que le modèle sauvegardé est celui qui décode le mieux les séquences de signes, ce qui est l'objectif final du système.")

# ══════════════════════════════════════════════════════════════════════════════
# 2.8.2 — Seq2Seq
# ══════════════════════════════════════════════════════════════════════════════
h1("2.8.2  Entraînement du modèle Seq2Seq")

j("L'entraînement du modèle Seq2Seq soulève un problème spécifique aux architectures auto-régressives : le biais d'exposition (exposure bias). En entraînement standard avec teacher forcing, le décodeur reçoit à chaque pas de temps le vrai mot précédent comme entrée, ce qui ne correspond pas aux conditions réelles d'inférence où il doit utiliser ses propres prédictions. Un écart entre les distributions d'entraînement et d'inférence s'accumule alors à chaque pas de génération.")

h2("Teacher Forcing Annealing")
j("Pour atténuer ce biais, une technique de teacher forcing annealing est utilisée : la probabilité d'utiliser le vrai mot précédent (teacher forcing) décroît progressivement au fil des époques, forçant le modèle à s'appuyer de plus en plus sur ses propres prédictions pendant l'entraînement lui-même.")

encadre("Formule du teacher forcing annealing :", [
    "tf(époque)  =  max(0.3,  1.0 − époque × 0.008)",
    "",
    "époque  1  :  tf = max(0.3, 0.992)  ≈  0.99   (quasi-teacher-forcing)",
    "époque 50  :  tf = max(0.3, 0.600)  =  0.60",
    "époque 88  :  tf = max(0.3, 0.296)  =  0.30   (plancher atteint)",
    "époque 100 :  tf = max(0.3, 0.200)  =  0.30   (plancher maintenu)",
    "",
    "À chaque pas de temps t du décodeur :",
    "  si rand() < tf  →  entrée = vrai mot précédent  (teacher forcing)",
    "  sinon           →  entrée = argmax(logits_{t-1})  (scheduled sampling)",
])

j("Le plancher de 0.30 empêche le modèle de basculer vers une génération entièrement auto-régressive pendant l'entraînement, ce qui entraînerait une instabilité due à l'accumulation d'erreurs sur les longues séquences.")

h2("Configuration d'entraînement")

table(
    ["Hyperparamètre", "Valeur", "Justification"],
    [
        ["Architecture encodeur",
         "BiLSTM 2 couches\nhidden = 128 par direction",
         "Taille modérée adaptée au vocabulaire source\nlimité (40 tokens)."],
        ["Architecture décodeur",
         "LSTM 2 couches\nhidden = 128",
         "Symétrie avec l'encodeur.\nAttention de Bahdanau intégrée."],
        ["Dimension d'embedding",
         "64 dims (src et tgt)",
         "Compromis entre expressivité et taille du modèle\npour de petits vocabulaires."],
        ["Vocabulaire source",
         "40 tokens",
         "30 signes + 4 tokens spéciaux (<PAD>, <SOS>, <EOS>, <UNK>)\n"
         "+ 6 tokens additionnels présents dans les données\n"
         "(attendons, encore, ensemble, maintenant, toi, voulons)."],
        ["Vocabulaire cible",
         "195 tokens",
         "Mots français distincts présents dans les 6 209 paires\nd'entraînement."],
        ["Teacher forcing",
         "Annealing\nmax(0.3, 1.0 − ep × 0.008)",
         "Atténue le biais d'exposition progressivement.\nPlancher à 0.30 pour stabilité."],
        ["Fonction de perte",
         "CrossEntropyLoss\n(ignore_index = PAD)",
         "Les positions de padding ne contribuent pas à la perte,\n"
         "ce qui évite de biaiser l'apprentissage vers les séquences courtes."],
        ["Optimiseur",
         "Adam  (lr = 1×10⁻³)",
         "Standard pour les modèles seq2seq."],
        ["Scheduler LR",
         "ReduceLROnPlateau\nfacteur = 0.5, patience = 5",
         "Réduit le LR de moitié si la perte de validation ne s'améliore\npas sur 5 époques consécutives."],
        ["Gradient clipping",
         "‖g‖ ≤ 1.0",
         "Seuil plus conservateur que pour le CTC,\nles séquences cibles étant plus courtes."],
        ["Arrêt anticipé",
         "patience = 15 époques\nsur val_loss",
         "Sauvegarde du meilleur modèle selon la perte de validation."],
        ["Époques maximum",
         "100",
         "Le plancher tf = 0.30 est atteint à l'époque 88,\n"
         "100 époques permettent un affinage après stabilisation."],
        ["Division train / val",
         "90 % / 10 %\n5 589 / 620 paires",
         "n_val = max(200, int(0.10 × 6 209)) = 620.\n"
         "Mélange aléatoire (seed = 42) avant la division."],
    ],
    cap="Tableau 2.7 — Hyperparamètres d'entraînement du modèle Seq2Seq à attention de Bahdanau\n"
       "(vérifiés dans scripts/train_seq2seq.py)",
    widths=[4.0, 4.0, 8.5]
)

h2("Décodage en inférence")
j("En inférence, le décodeur opère de façon entièrement auto-régressive : le token <SOS> initialise la génération, et chaque mot produit devient l'entrée du pas de temps suivant. La génération s'arrête dès l'émission du token <EOS> ou après 20 mots au maximum (MAX_LEN = 20 dans Seq2SeqEngine.translate()). Le décodage est glouton : à chaque pas de temps, le mot retenu est celui de plus haute probabilité (argmax des logits).")

# ══════════════════════════════════════════════════════════════════════════════
# 2.8.3 — Résultats
# ══════════════════════════════════════════════════════════════════════════════
h1("2.8.3  Résultats d'entraînement")

table(
    ["Modèle", "Corpus d'entraînement", "Corpus de validation", "Meilleure métrique"],
    [
        ["CTC BiLSTM",
         "41 400 séquences\n(90 % × 46 000 synthétiques)",
         "4 600 séquences\n(10 % × 46 000)",
         "Sign Accuracy : 99.97 %\n(source : data/ctc_model/config.json)"],
        ["Seq2Seq attention",
         "5 589 paires\n(90 % × 6 209)",
         "620 paires\n(10 % × 6 209)",
         "Perte de validation : 0.0354\ncross-entropie (ignore PAD)\n(source : data/seq2seq/config.json)"],
    ],
    cap="Tableau 2.8 — Résultats d'entraînement des modèles CTC BiLSTM et Seq2Seq\n"
       "Valeurs extraites directement des fichiers config.json générés à l'entraînement.",
    widths=[3.5, 4.5, 4.0, 4.5]
)

j("La Sign Accuracy de 99.97 % du modèle CTC sur les séquences de validation synthétiques indique que le modèle a appris à décoder correctement les séquences continues sans alignement préalable. Il convient de noter que cette métrique est calculée sur des données synthétiques issues du même générateur que les données d'entraînement ; les performances sur des séquences réelles enregistrées en conditions variées peuvent différer.")
j("La perte de validation Seq2Seq de 0.0354 correspond à une perplexité de e^{0.0354} ≈ 1.036, ce qui indique que le modèle assigne une probabilité très élevée aux mots corrects sur les paires de validation. Les traductions produites sur les séquences du vocabulaire couvert ont été vérifiées manuellement et sont grammaticalement correctes.")

hr()

# Références
h2("Références", sb=16, sa=6)
refs_list = [
    "Graves, A., Fernández, S., Gomez, F., & Schmidhuber, J. (2006). Connectionist Temporal Classification: Labelling unsegmented sequence data with recurrent neural networks. ICML 2006, pp. 369–376.",
    "Bahdanau, D., Cho, K., & Bengio, Y. (2015). Neural Machine Translation by Jointly Learning to Align and Translate. ICLR 2015. arXiv:1409.0473.",
    "Bengio, S., Vinyals, O., Jaitly, N., & Shazeer, N. (2015). Scheduled Sampling for Sequence Prediction with Recurrent Neural Networks. NeurIPS 2015. arXiv:1506.03099. (Fondement du scheduled sampling / teacher forcing annealing.)",
    "PyTorch Contributors. (2024). torch.nn.CTCLoss. PyTorch Documentation. pytorch.org/docs/stable/generated/torch.nn.CTCLoss.html.",
    "Loshchilov, I., & Hutter, F. (2017). SGDR: Stochastic Gradient Descent with Warm Restarts. ICLR 2017. arXiv:1608.03983. (Fondement du CosineAnnealingLR.)",
]
for i, ref in enumerate(refs_list, 1):
    rp = doc.add_paragraph()
    rp.paragraph_format.left_indent       = Cm(1.0)
    rp.paragraph_format.first_line_indent = Cm(-1.0)
    rp.paragraph_format.space_after       = Pt(5)
    rr = rp.add_run(f"[{i}]  {ref}")
    rr.font.size = Pt(10); rr.font.name = "Times New Roman"

doc.save(OUT)
size_kb = __import__('os').path.getsize(OUT) / 1024
print(f"Généré : {OUT}")
print(f"Taille : {size_kb:.1f} Ko")
