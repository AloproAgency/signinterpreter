#!/usr/bin/env python3
"""Génère le mémoire de master en format .docx"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx.opc.constants
import copy

OUT = "/Users/alopro/Desktop/AI/RECHERCHE/SignInterpreter/V11/Redaction/memoire_master_sign_interpreter.docx"

doc = Document()

# ── Page setup ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ── Helper functions ─────────────────────────────────────────────────────────
def add_paragraph(text="", bold=False, italic=False, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  space_before=0, space_after=6, font_name="Times New Roman", color=None):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.first_line_indent = Pt(0)
    if text:
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.size = Pt(size)
        run.font.name = font_name
        if color:
            run.font.color.rgb = RGBColor(*color)
    return p

def add_heading(text, level=1, size=None, align=WD_ALIGN_PARAGRAPH.LEFT,
                space_before=12, space_after=6):
    sizes = {1: 16, 2: 14, 3: 12, 4: 11}
    if size is None:
        size = sizes.get(level, 12)
    p = doc.add_paragraph()
    p.paragraph_format.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    return p

def add_justified(text, size=12, indent=True, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    return p

def add_centered(text, size=12, bold=False, space_before=6, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    return p

def add_page_break():
    doc.add_page_break()

def add_section_title(text, size=18):
    add_page_break()
    add_paragraph()
    add_paragraph()
    add_centered(text, size=size, bold=True, space_before=60, space_after=20)
    add_paragraph()

def add_chapter_title(text, num=""):
    add_page_break()
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(72)
    p.paragraph_format.space_after  = Pt(24)
    # Border box effect via XML
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ['top','left','bottom','right']:
        bdr = OxmlElement(f'w:{side}')
        bdr.set(qn('w:val'), 'single')
        bdr.set(qn('w:sz'), '12')
        bdr.set(qn('w:space'), '6')
        bdr.set(qn('w:color'), 'E07000')
        pBdr.append(bdr)
    pPr.append(pBdr)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "Times New Roman"
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].runs[0].font.name = "Times New Roman"
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows):
        tr = table.rows[ri+1]
        for ci, val in enumerate(row):
            tr.cells[ci].text = str(val)
            tr.cells[ci].paragraphs[0].runs[0].font.size = Pt(10)
            tr.cells[ci].paragraphs[0].runs[0].font.name = "Times New Roman"
    doc.add_paragraph()
    return table

def add_footer_text(text):
    """Add small footer note below content"""
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.italic = True
    run.font.name = "Times New Roman"

def add_horizontal_rule():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

# ═══════════════════════════════════════════════════════════════════════════════
# PAGE DE GARDE
# ═══════════════════════════════════════════════════════════════════════════════
add_paragraph()
add_centered("REPUBLIQUE DU BENIN", size=13, bold=True)
add_centered("**********", size=12)
add_centered("MINISTERE DE L'ENSEIGNEMENT SUPERIEUR ET", size=13, bold=True)
add_centered("DE LA RECHERCHE SCIENTIFIQUE", size=13, bold=True)
add_centered("**********", size=12)
add_paragraph()
add_centered("ECOLE PIGIER-BENIN", size=16, bold=True, space_before=10)
add_paragraph()
add_horizontal_rule()
add_paragraph()
add_centered("DOMAINE : SCIENCES ET TECHNOLOGIES", size=12)
add_paragraph()
add_centered("MENTION : INFORMATIQUE", size=12)
add_paragraph()
add_centered("SPÉCIALITÉ : INTELLIGENCE ARTIFICIELLE & BIG DATA", size=12)
add_paragraph()
add_paragraph()
add_centered("THÈME", size=14, bold=True, space_before=10)
add_paragraph()

# Boxed title
p_title = doc.add_paragraph()
p_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(6)
p_title.paragraph_format.space_after  = Pt(6)
pPr = p_title._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
for side in ['top','left','bottom','right']:
    bdr = OxmlElement(f'w:{side}')
    bdr.set(qn('w:val'), 'single')
    bdr.set(qn('w:sz'), '12')
    bdr.set(qn('w:space'), '6')
    bdr.set(qn('w:color'), 'E07000')
    pBdr.append(bdr)
pPr.append(pBdr)
run = p_title.add_run(
    "SYSTÈME MULTIMODAL DE TRADUCTION DE LA LANGUE DES SIGNES\n"
    "VERS LA PAROLE BASÉ SUR LA FUSION DES SIGNAUX CORPORELS,\n"
    "MANUELS ET DU CONTEXTE LINGUISTIQUE"
)
run.bold  = True
run.font.size = Pt(14)
run.font.name = "Times New Roman"

add_paragraph()
add_paragraph()
add_centered("Réalisé et présenté par :", size=11, space_before=10)
add_centered("[VOTRE NOM COMPLET]", size=13, bold=True)
add_paragraph()
add_centered("Sous la supervision :", size=11, space_before=6)
add_centered("[NOM ET TITRE DU DIRECTEUR DE MÉMOIRE]", size=12, bold=True)
add_paragraph()
add_paragraph()
add_centered("Session de Juin 2025", size=14, bold=True, space_before=20)

# ═══════════════════════════════════════════════════════════════════════════════
# ENGAGEMENT
# ═══════════════════════════════════════════════════════════════════════════════
add_section_title("Engagement")
p_eng = doc.add_paragraph()
p_eng.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_eng.paragraph_format.space_before = Pt(24)
p_eng.paragraph_format.space_after  = Pt(24)
pPr2 = p_eng._p.get_or_add_pPr()
pBdr2 = OxmlElement('w:pBdr')
for side in ['top','left','bottom','right']:
    bdr = OxmlElement(f'w:{side}')
    bdr.set(qn('w:val'), 'single')
    bdr.set(qn('w:sz'), '6')
    bdr.set(qn('w:space'), '10')
    bdr.set(qn('w:color'), 'E07000')
    pBdr2.append(bdr)
pPr2.append(pBdr2)
run_eng = p_eng.add_run(
    "L'école PIGIER-BÉNIN n'entend donner aucune approbation ni\n"
    "désapprobation aux opinions émises dans ce mémoire. Ces opinions doivent\n"
    "être considérées comme propres à leurs auteurs."
)
run_eng.font.size = Pt(12)
run_eng.font.name = "Times New Roman"

# ═══════════════════════════════════════════════════════════════════════════════
# DÉDICACES
# ═══════════════════════════════════════════════════════════════════════════════
add_section_title("Dédicaces")
add_paragraph()
add_paragraph()
add_justified(
    "À ma famille, pour son soutien constant et ses encouragements tout au long "
    "de mon parcours académique.",
    indent=False, space_before=12, space_after=12
)
add_justified(
    "À la communauté sourde et malentendante du Bénin, et en particulier aux "
    "élèves et enseignants du Centre d'Accueil et d'Intégration des Sourds de "
    "Louho (Porto-Novo), dont la résilience et la richesse linguistique sont la "
    "principale source d'inspiration de ce travail.",
    indent=False, space_before=6, space_after=12
)
add_justified(
    "À tous ceux qui croient que la technologie peut être un vecteur "
    "d'inclusion sociale.",
    indent=False, space_before=6
)
add_paragraph()
add_paragraph()
p_ded = doc.add_paragraph()
p_ded.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run_ded = p_ded.add_run("[VOTRE NOM]")
run_ded.font.size = Pt(12)
run_ded.font.name = "Times New Roman"
run_ded.bold = True
run_ded.font.color.rgb = RGBColor(0, 0, 139)

# ═══════════════════════════════════════════════════════════════════════════════
# REMERCIEMENTS
# ═══════════════════════════════════════════════════════════════════════════════
add_section_title("Remerciements")
add_justified(
    "Nous souhaitons exprimer notre profonde gratitude à toutes les personnes "
    "qui ont, de près ou de loin, contribué à la réalisation de ce travail. "
    "En signe de reconnaissance, nous tenons à leur adresser nos sincères "
    "remerciements. Nous remercions tout particulièrement :",
    indent=False, space_before=6, space_after=12
)
thanks = [
    "[NOM DU DIRECTEUR], pour avoir accepté de diriger ce mémoire avec "
    "bienveillance et pour ses précieux conseils tout au long de ce travail.",
    "L'école PIGIER-BÉNIN, ainsi que l'ensemble des membres de l'administration "
    "et du corps enseignant, pour la qualité de la formation dispensée et "
    "l'accompagnement tout au long de notre parcours académique.",
    "L'équipe du Centre d'Accueil et d'Intégration des Sourds (CAEIS) de Louho "
    "à Porto-Novo, pour les échanges enrichissants sur la langue des signes et "
    "les besoins technologiques de la communauté sourde béninoise.",
]
for t in thanks:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(t)
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

add_paragraph()
add_justified(
    "Nous tenons également à exprimer notre reconnaissance à nos camarades de "
    "promotion pour leur esprit de collaboration, ainsi qu'à notre famille pour "
    "leur appui constant durant cette période exigeante.",
    indent=False, space_before=6, space_after=6
)
add_justified(
    "Enfin, à toutes les personnes ayant apporté un regard critique et "
    "constructif sur ce travail.",
    indent=False, space_before=6
)

# ═══════════════════════════════════════════════════════════════════════════════
# SIGLES ET ABRÉVIATIONS
# ═══════════════════════════════════════════════════════════════════════════════
add_section_title("Sigles, Acronymes et Abréviations")
sigles = [
    ("AI", "Artificial Intelligence"),
    ("API", "Application Programming Interface"),
    ("ASL", "American Sign Language"),
    ("ASUNOES", "Association Universelle qui Œuvre pour l'Épanouissement des Sourds"),
    ("BiLSTM", "Bidirectional Long Short-Term Memory"),
    ("CAEIS", "Centre d'Accueil, d'Éducation et d'Intégration des Sourds"),
    ("CTC", "Connectionist Temporal Classification"),
    ("DL", "Deep Learning"),
    ("EOS", "End of Sequence"),
    ("FPS", "Frames Per Second"),
    ("GPU", "Graphics Processing Unit"),
    ("HOG", "Histogram of Oriented Gradients"),
    ("HMM", "Hidden Markov Model"),
    ("KNN", "K-Nearest Neighbors"),
    ("LSB", "Langue des Signes Béninoise"),
    ("LSF", "Langue des Signes Française"),
    ("LSTM", "Long Short-Term Memory"),
    ("ML", "Machine Learning"),
    ("MPS", "Metal Performance Shaders"),
    ("NLP", "Natural Language Processing"),
    ("OMS", "Organisation Mondiale de la Santé"),
    ("PAD", "Padding token"),
    ("RNN", "Recurrent Neural Network"),
    ("SOS", "Start of Sequence"),
    ("SVM", "Support Vector Machine"),
    ("WS", "WebSocket"),
]
for abbr, full in sigles:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    run1 = p.add_run(f"{abbr} : ")
    run1.bold = True
    run1.font.size = Pt(12)
    run1.font.name = "Times New Roman"
    run2 = p.add_run(full)
    run2.font.size = Pt(12)
    run2.font.name = "Times New Roman"

# ═══════════════════════════════════════════════════════════════════════════════
# RÉSUMÉ
# ═══════════════════════════════════════════════════════════════════════════════
add_section_title("Résumé")
add_justified(
    "Ce mémoire présente la conception, le développement et l'évaluation d'un "
    "système de traduction en temps réel de la Langue des Signes Française (LSF) "
    "vers le français écrit. Face aux barrières de communication que rencontrent "
    "quotidiennement les personnes sourdes et malentendantes, notamment au Bénin "
    "où les structures d'accompagnement technologique restent insuffisantes, nous "
    "proposons une architecture multimodale intégrant trois niveaux complémentaires "
    "de traitement : la capture des signaux corporels et manuels par MediaPipe "
    "Holistic (Google), la reconnaissance de séquences de signes par un réseau "
    "BiLSTM entraîné avec la perte CTC (Connectionist Temporal Classification), "
    "et la traduction vers le français naturel par un module de règles linguistiques.",
    space_before=6, space_after=8
)
add_justified(
    "Le système extrait 171 caractéristiques par image, représentant la pose "
    "corporelle supérieure (39 dimensions), la configuration de la main gauche "
    "(66 dimensions) et de la main droite (66 dimensions). Le modèle BiLSTM-CTC, "
    "fort de 2 471 455 paramètres entraînables, reconnaît 30 signes courants du "
    "LSF avec une précision de 99,97 % sur un corpus de 46 000 séquences "
    "synthétiques générées à partir de 2 100 enregistrements réels. La "
    "traduction couvre 6 209 combinaisons de signes vers des phrases françaises "
    "correctement structurées.",
    space_before=0, space_after=8
)
add_justified(
    "L'architecture client-serveur repose sur WebSocket pour une communication "
    "bidirectionnelle temps réel : MediaPipe s'exécute côté navigateur (JavaScript) "
    "tandis que l'inférence CTC est prise en charge par un serveur FastAPI (Python). "
    "Le système propose un affichage progressif des signes reconnus (mise à jour "
    "toutes les 20 images, soit ~0,67 s à 30 fps) et une finalisation automatique "
    "à la détection d'une pause ou d'une posture de repos.",
    space_before=0, space_after=8
)
add_justified(
    "Ce travail s'inscrit dans la Stratégie Nationale d'Intelligence Artificielle "
    "et de Mégadonnées du Bénin 2023-2027, qui identifie l'accessibilité pour les "
    "personnes en situation de handicap comme un axe d'intervention prioritaire.",
    space_before=0, space_after=12
)
p_kw = doc.add_paragraph()
run_kw1 = p_kw.add_run("Mots clés : ")
run_kw1.bold = True
run_kw1.font.size = Pt(12)
run_kw1.font.name = "Times New Roman"
run_kw2 = p_kw.add_run(
    "Langue des Signes Française, BiLSTM, CTC, MediaPipe, Traduction temps réel, "
    "Multimodalité, Accessibilité, Bénin, Apprentissage profond."
)
run_kw2.font.size = Pt(12)
run_kw2.font.name = "Times New Roman"

# ═══════════════════════════════════════════════════════════════════════════════
# ABSTRACT
# ═══════════════════════════════════════════════════════════════════════════════
add_section_title("Abstract")
add_justified(
    "This thesis presents the design, development, and evaluation of a real-time "
    "French Sign Language (LSF) to written French translation system. Addressing "
    "the persistent communication barriers faced by deaf and hard-of-hearing "
    "individuals — particularly in Benin, where dedicated technological support "
    "remains scarce — we propose a multimodal architecture integrating three "
    "complementary processing layers: body and hand signal capture via MediaPipe "
    "Holistic (Google), sign sequence recognition by a BiLSTM network trained with "
    "Connectionist Temporal Classification (CTC) loss, and natural French "
    "translation through a rule-based linguistic module.",
    space_before=6, space_after=8
)
add_justified(
    "The system extracts 171 features per frame, encoding upper-body pose "
    "(39 dimensions), left hand configuration (66 dimensions), and right hand "
    "configuration (66 dimensions). The BiLSTM-CTC model, with 2,471,455 "
    "trainable parameters, recognizes 30 common LSF signs with 99.97% accuracy "
    "on a corpus of 46,000 synthetic sequences generated from 2,100 real "
    "recordings. The translation module covers 6,209 sign combinations mapped "
    "to grammatically correct French sentences.",
    space_before=0, space_after=8
)
add_justified(
    "The client-server architecture uses WebSocket for real-time bidirectional "
    "communication: MediaPipe runs in the browser (JavaScript), while CTC "
    "inference is handled by a FastAPI server (Python). The system provides "
    "progressive sign display (updated every 20 frames, ~0.67 s at 30 fps) "
    "and automatic finalization on pause or rest posture detection.",
    space_before=0, space_after=12
)
p_kw2 = doc.add_paragraph()
run_kw3 = p_kw2.add_run("Keywords: ")
run_kw3.bold = True
run_kw3.font.size = Pt(12)
run_kw3.font.name = "Times New Roman"
run_kw4 = p_kw2.add_run(
    "French Sign Language, BiLSTM, CTC, MediaPipe, Real-time translation, "
    "Multimodality, Accessibility, Benin, Deep Learning."
)
run_kw4.font.size = Pt(12)
run_kw4.font.name = "Times New Roman"

# ═══════════════════════════════════════════════════════════════════════════════
# SOMMAIRE
# ═══════════════════════════════════════════════════════════════════════════════
add_section_title("Sommaire")
sommaire_items = [
    ("ENGAGEMENT", "i"),
    ("DÉDICACES", "ii"),
    ("REMERCIEMENTS", "iii"),
    ("SIGLES, ACRONYMES ET ABRÉVIATIONS", "iv"),
    ("RÉSUMÉ", "vi"),
    ("ABSTRACT", "vii"),
    ("SOMMAIRE", "viii"),
    ("INTRODUCTION", "1"),
    ("CHAPITRE 1 : REVUE DE LITTÉRATURE", "5"),
    ("  Section 1 : Définitions des thèmes et concepts", "5"),
    ("  Section 2 : La langue des signes et la communauté sourde au Bénin", "18"),
    ("  Section 3 : État de l'art des systèmes de reconnaissance", "23"),
    ("CHAPITRE 2 : ARCHITECTURE ET MÉTHODOLOGIE", "32"),
    ("  Section 1 : Architecture matérielle", "32"),
    ("  Section 2 : Architecture logicielle", "35"),
    ("  Section 3 : Architecture du système et pipeline de traitement", "42"),
    ("CHAPITRE 3 : RÉSULTATS ET DISCUSSIONS", "60"),
    ("  Section 1 : Construction du jeu de données", "60"),
    ("  Section 2 : Entraînement et évaluation du modèle BiLSTM-CTC", "66"),
    ("  Section 3 : Évaluation du système de traduction", "72"),
    ("  Section 4 : Interface et déploiement", "75"),
    ("  Section 5 : Discussions et perspectives", "78"),
    ("CONCLUSION", "83"),
    ("RÉFÉRENCES BIBLIOGRAPHIQUES ET WEBOGRAPHIQUES", "86"),
    ("TABLE DES MATIÈRES", "92"),
]
for item, page in sommaire_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    bold_item = not item.startswith("  ")
    run_s = p.add_run(f"{item}")
    run_s.bold = bold_item
    run_s.font.size = Pt(11)
    run_s.font.name = "Times New Roman"
    run_p = p.add_run(f" {'.'*max(1, 70-len(item))} {page}")
    run_p.font.size = Pt(11)
    run_p.font.name = "Times New Roman"

print("Front matter done. Writing Introduction...")

# ═══════════════════════════════════════════════════════════════════════════════
# INTRODUCTION
# ═══════════════════════════════════════════════════════════════════════════════
add_chapter_title("Introduction")
add_paragraph()

add_heading("1. Problématique", level=2, size=13)
add_justified(
    "Selon l'Organisation Mondiale de la Santé (OMS, 2023), plus de 430 millions "
    "de personnes dans le monde souffrent d'une déficience auditive invalidante, "
    "dont 80 % vivent dans des pays à revenus faibles ou intermédiaires. En Afrique "
    "subsaharienne, et plus particulièrement au Bénin, cette réalité est amplifiée "
    "par l'absence quasi totale d'outils technologiques adaptés et par des "
    "représentations sociales parfois défavorables envers les personnes sourdes "
    "(Prévot, 2011 ; Djissenou et al., 2021)."
)
add_justified(
    "La communication entre personnes sourdes ou malentendantes et la population "
    "entendante constitue l'une des barrières sociales les plus persistantes. "
    "Si des institutions comme le Centre d'Accueil et d'Intégration des Sourds "
    "(CAEIS) de Louho à Porto-Novo, fondé en 1993, ont entrepris des efforts "
    "remarquables pour l'éducation bilingue en Langue des Signes Française (LSF) "
    "et en français, le fossé communicationnel entre les deux communautés demeure "
    "très important dans la vie quotidienne (Prévot, 2011)."
)
add_justified(
    "Le recours à un interprète humain, solution traditionnelle, reste coûteux, "
    "peu disponible et non scalable. Les technologies de traduction automatique "
    "de la langue des signes pourraient combler ce manque, à condition d'être "
    "accessibles, robustes et adaptées au contexte local. Cette problématique "
    "soulève plusieurs défis techniques majeurs : comment capturer fidèlement les "
    "paramètres phonologiques d'un signe (configuration, emplacement, mouvement, "
    "orientation) sans équipement spécialisé ? Comment segmenter automatiquement "
    "des séquences continues de signes sans annotation temporelle explicite ? "
    "Comment assurer une latence suffisamment faible pour une utilisation en "
    "temps réel sur des équipements courants ?"
)

add_heading("2. Contexte et motivations", level=2, size=13)
add_justified(
    "Le Bénin s'est engagé dans une dynamique d'inclusion numérique à travers sa "
    "Stratégie Nationale d'Intelligence Artificielle et de Mégadonnées 2023-2027, "
    "définie par le Ministère du Numérique et de la Digitalisation. Parmi ses "
    "axes d'intervention prioritaires figurent explicitement l'accessibilité "
    "technologique pour les personnes en situation de handicap et l'inclusion "
    "sociale par le numérique."
)
add_justified(
    "Parallèlement, les avancées récentes en apprentissage profond — notamment "
    "les réseaux LSTM bidirectionnels (Hochreiter & Schmidhuber, 1997), la perte "
    "CTC (Graves et al., 2006) pour la reconnaissance de séquences, et les outils "
    "d'estimation de pose comme MediaPipe (Lugaresi et al., 2019) — offrent "
    "aujourd'hui les briques technologiques nécessaires pour concevoir un système "
    "de traduction de la langue des signes performant sans recourir à des capteurs "
    "coûteux ni à d'immenses corpus annotés."
)
add_justified(
    "Notre motivation est double : scientifique (explorer l'application de "
    "l'apprentissage profond à une langue peu dotée en ressources numériques) "
    "et sociale (contribuer à réduire les barrières de communication pour la "
    "communauté sourde béninoise)."
)

add_heading("3. Objectifs de l'étude", level=2, size=13)
add_justified(
    "L'objectif principal de ce travail est de concevoir et de mettre en œuvre "
    "un système multimodal de traduction en temps réel de la Langue des Signes "
    "Française (LSF) vers le français écrit, déployable dans un navigateur "
    "standard sans équipement spécialisé."
)
add_paragraph("Les objectifs spécifiques sont :", size=12)
specifiques = [
    "Constituer un corpus de templates de signes LSF (enregistrements réels) couvrant un vocabulaire de base de 30 signes courants ;",
    "Développer un pipeline de génération de données synthétiques pour pallier la rareté des données annotées, en combinant rééchantillonnage temporel, interpolation aux transitions et perturbation par bruit ;",
    "Concevoir et entraîner un modèle BiLSTM-CTC pour la reconnaissance de séquences continues de signes sans segmentation préalable ;",
    "Implémenter un système de traduction LSF→Français basé sur des règles linguistiques prenant en compte la structure grammaticale de la LSF ;",
    "Déployer le système complet en architecture client-serveur WebSocket pour une utilisation temps réel via navigateur web ;",
    "Évaluer les performances du système sur des critères quantitatifs (précision, latence) et qualitatifs (naturalité des traductions).",
]
for s in specifiques:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(s)
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

add_heading("4. Méthodologie générale", level=2, size=13)
add_justified(
    "Ce travail adopte une approche de développement agile et itérative. "
    "Le système a évolué à travers onze versions successives (V1 à V11), "
    "chaque version apportant des améliorations ciblées sur la qualité "
    "des données, l'architecture du modèle, le pipeline d'inférence ou "
    "l'interface utilisateur. La version finale (V11) constitue l'aboutissement "
    "de ce processus itératif."
)
add_justified(
    "La méthodologie s'articule autour de quatre phases principales : "
    "(1) collecte et constitution du corpus de templates, (2) génération synthétique "
    "et augmentation des données d'entraînement, (3) entraînement et validation "
    "du modèle BiLSTM-CTC, et (4) intégration dans le pipeline temps réel et "
    "évaluation système."
)

add_heading("5. Plan du mémoire", level=2, size=13)
add_justified(
    "Ce mémoire est organisé en trois chapitres. Le premier chapitre présente "
    "une revue de la littérature couvrant les concepts fondamentaux (langue des "
    "signes, apprentissage profond, CTC), le contexte de la surdité au Bénin "
    "et l'état de l'art des systèmes de reconnaissance de la langue des signes. "
    "Le deuxième chapitre décrit l'architecture matérielle, logicielle et "
    "méthodologique du système SignInterpreter V11. Le troisième chapitre "
    "présente les résultats obtenus, leur analyse et les perspectives "
    "d'amélioration. Une conclusion synthétise les apports du travail et "
    "ouvre sur de futures directions de recherche."
)

print("Introduction done. Writing Chapitre 1...")

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPITRE 1 : REVUE DE LITTÉRATURE
# ═══════════════════════════════════════════════════════════════════════════════
add_chapter_title("CHAPITRE 1 : REVUE DE LITTÉRATURE")
add_paragraph()

# SECTION 1
add_heading("Section 1 : Définitions des thèmes et concepts", level=1, size=14)
add_horizontal_rule()

add_heading("1.1 La langue des signes", level=2, size=13)
add_justified(
    "La langue des signes est une langue naturelle et complète qui utilise le "
    "canal visuo-gestuel pour communiquer. Contrairement aux idées reçues, "
    "la langue des signes n'est pas universelle : chaque communauté sourde "
    "nationale dispose de sa propre langue, avec sa phonologie, sa morphologie, "
    "sa syntaxe et sa sémantique propres (Cuxac, 1983 ; Armstrong, 2011). "
    "La Langue des Signes Française (LSF), utilisée en France et dans plusieurs "
    "pays francophones d'Afrique de l'Ouest dont le Bénin, est ainsi distincte "
    "de l'American Sign Language (ASL), du British Sign Language (BSL), "
    "ou de la Langue des Signes Béninoise (LSB) en cours d'émergence."
)
add_justified(
    "Un signe en LSF est analysé selon cinq paramètres phonologiques "
    "(Battison, 1978 ; Brentari, 2019) :"
)
params = [
    "La configuration de la main (chérème) : la forme que prend la main (poing fermé, main ouverte, index pointé, etc.) ;",
    "L'emplacement (tabulation) : la position du signe dans l'espace de signation par rapport au corps du signeur ;",
    "Le mouvement (kinème) : la trajectoire et la dynamique du déplacement de la main ;",
    "L'orientation de la paume : la direction vers laquelle la paume est tournée ;",
    "Les éléments non-manuels : les expressions faciales, le mouvement des lèvres, les mouvements de tête qui jouent un rôle grammatical (marquage des questions, de la négation, etc.).",
]
for param in params:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(param)
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

add_justified(
    "Sur le plan syntaxique, la LSF possède une structure grammaticale propre, "
    "différente du français oral. Elle fait usage de l'espace de signation pour "
    "établir des références, organise les constituants selon un ordre "
    "sujet-objet-verbe (SOV) plus flexible que le français, et utilise des "
    "classificateurs pour décrire des formes, des tailles et des mouvements "
    "(Cuxac, 2000)."
)

add_heading("1.2 La surdité et les personnes sourdes", level=2, size=13)
add_justified(
    "L'Organisation Mondiale de la Santé (OMS, 2001) définit la surdité comme "
    "le résultat de l'incapacité de l'environnement à assimiler et à gérer les "
    "altérités. Elle distingue deux types de déficience auditive : la surdité "
    "de transmission, résultant d'un problème de l'oreille externe ou moyenne "
    "(souvent curable médicalement), et la surdité neurosensorielle, liée à "
    "une atteinte de l'oreille interne ou du nerf auditif (généralement "
    "permanente) (Djissenou et al., 2021)."
)
add_justified(
    "La communauté sourde se distingue par une identité culturelle et linguistique "
    "forte, fondée sur le partage d'une langue visuelle et d'une expérience "
    "commune du monde. Dans cette perspective, la surdité n'est pas perçue "
    "comme un handicap mais comme une différence culturelle (Ladd, 2003)."
)

add_heading("1.3 L'intelligence artificielle et l'apprentissage profond", level=2, size=13)
add_justified(
    "L'intelligence artificielle (IA) désigne l'ensemble des théories et des "
    "techniques visant à permettre à des machines de simuler l'intelligence "
    "humaine. Le sous-domaine de l'apprentissage automatique (machine learning, "
    "ML) consiste à faire apprendre des patterns à un système à partir "
    "d'exemples, sans programmation explicite de règles (LeCun et al., 2015)."
)
add_justified(
    "L'apprentissage profond (deep learning, DL), sous-domaine du ML, utilise "
    "des réseaux de neurones artificiels à plusieurs couches (couches profondes) "
    "pour apprendre des représentations hiérarchiques des données. Ces architectures "
    "ont révolutionné la vision par ordinateur, le traitement du langage naturel "
    "et la reconnaissance vocale depuis 2012 (Krizhevsky et al., 2012)."
)

add_heading("1.4 Les réseaux LSTM et BiLSTM", level=2, size=13)
add_justified(
    "Les réseaux de neurones récurrents (RNN) sont conçus pour traiter des "
    "données séquentielles en maintenant un état caché propagé dans le temps. "
    "Cependant, les RNN classiques souffrent du problème du gradient vanishing, "
    "qui empêche l'apprentissage de dépendances à long terme (Hochreiter, 1991)."
)
add_justified(
    "Le réseau LSTM (Long Short-Term Memory), proposé par Hochreiter & Schmidhuber "
    "(1997), résout ce problème en introduisant une cellule à mémoire avec trois "
    "portes (gates) contrôlant le flux d'information : la porte d'oubli (forget gate) "
    "détermine quelle information effacer de la mémoire, la porte d'entrée (input gate) "
    "contrôle quelle nouvelle information stocker, et la porte de sortie (output gate) "
    "détermine quelle information produire en sortie."
)
add_justified(
    "Le BiLSTM (LSTM bidirectionnel) traite la séquence dans les deux sens temporels "
    "(passé→futur et futur→passé) et concatène les états cachés des deux directions. "
    "Cette architecture permet de capturer des contextes plus riches, particulièrement "
    "utiles pour la reconnaissance de gestes où la dynamique avant et après un signe "
    "contribue à son identification."
)

add_heading("1.5 La perte CTC (Connectionist Temporal Classification)", level=2, size=13)
add_justified(
    "La reconnaissance de la langue des signes en continu pose un problème "
    "fondamental d'alignement : les données d'entrée (séquences de frames vidéo) "
    "et les données de sortie (séquences de labels de signes) ont des longueurs "
    "différentes, et la segmentation temporelle de chaque signe est coûteuse "
    "à annoter manuellement."
)
add_justified(
    "La perte CTC (Connectionist Temporal Classification), introduite par Graves "
    "et al. (2006) pour la reconnaissance vocale, résout ce problème en introduisant "
    "un token spécial 'blank' (ε) et en marginalisant sur tous les alignements "
    "possibles entre la séquence d'entrée et la séquence de labels cible :"
)
add_centered("L_CTC = -ln P(y | x) = -ln Σ_{π∈B⁻¹(y)} P(π | x)", size=12, bold=False, space_before=8, space_after=8)
add_justified(
    "où y est la séquence de labels cible, x la séquence d'entrée, π un chemin "
    "d'alignement (séquence de labels incluant des blanks), et B la fonction de "
    "collapsage qui supprime les blanks et les répétitions consécutives. "
    "Le décodage greedy sélectionne le label le plus probable à chaque pas de "
    "temps, puis applique la fonction B sur la séquence résultante."
)
add_justified(
    "Appliqué à la langue des signes, CTC permet d'entraîner un modèle à reconnaître "
    "des séquences de signes sans connaître les instants précis de début et de fin "
    "de chaque signe, ce qui simplifie considérablement la collecte de données "
    "(Koller et al., 2016, 2017)."
)

add_heading("1.6 MediaPipe et l'estimation de pose", level=2, size=13)
add_justified(
    "MediaPipe (Lugaresi et al., 2019), développé par Google, est un framework "
    "d'analyse visuelle en temps réel proposant des solutions pré-entraînées "
    "pour l'estimation de pose, la détection des mains et du visage. "
    "La solution MediaPipe Holistic fusionne trois modèles spécialisés pour "
    "extraire simultanément 33 points de pose corporelle, 21 points de "
    "landmarks par main et 468 points pour le visage."
)
add_justified(
    "MediaPipe présente plusieurs avantages pour notre application : "
    "il fonctionne en temps réel dans un navigateur web (WebAssembly), "
    "ne nécessite aucun capteur spécialisé (simple webcam RGB), "
    "et est robuste à des conditions d'éclairage variées. Ces caractéristiques "
    "en font un choix adapté au contexte béninois où les équipements disponibles "
    "sont généralement d'entrée de gamme."
)

add_heading("1.7 La multimodalité dans la reconnaissance des gestes", level=2, size=13)
add_justified(
    "Un système multimodal combine plusieurs canaux ou modalités d'information "
    "pour améliorer sa robustesse et ses performances. Dans le contexte de la "
    "langue des signes, la multimodalité implique la fusion de signaux de "
    "différentes natures : les signaux corporels (position du tronc, des bras, "
    "des épaules), les signaux manuels (configuration détaillée des doigts, "
    "mouvement des mains), et potentiellement les signaux faciaux (expressions "
    "faciales jouant un rôle grammatical en LSF)."
)
add_justified(
    "La fusion de ces modalités peut s'opérer à différents niveaux : au niveau "
    "des features (feature-level fusion), au niveau des décisions (decision-level "
    "fusion), ou au niveau hybride (Atrey et al., 2010). Dans notre approche, "
    "la fusion est opérée au niveau des features par concaténation des vecteurs "
    "de caractéristiques corporelles et manuelles en un unique vecteur de 171 "
    "dimensions, traité conjointement par le modèle BiLSTM-CTC."
)

add_heading("1.8 Architectures seq2seq et mécanismes d'attention", level=2, size=13)
add_justified(
    "Les modèles séquence-à-séquence (seq2seq), introduits par Sutskever et al. "
    "(2014), permettent de traduire une séquence d'entrée de longueur variable "
    "vers une séquence de sortie de longueur différente. Ils sont composés d'un "
    "encodeur qui compresse la séquence d'entrée en un vecteur de contexte, "
    "et d'un décodeur qui génère la séquence de sortie à partir de ce vecteur."
)
add_justified(
    "Le mécanisme d'attention de Bahdanau et al. (2015) améliore ces modèles en "
    "permettant au décodeur de consulter sélectivement différentes parties de "
    "la séquence encodée à chaque pas de décodage, plutôt que de se reposer "
    "uniquement sur un vecteur de contexte fixe. Cette approche, explorée dans "
    "ce travail pour la traduction LSF→Français, est décrite au chapitre 3."
)

# SECTION 2
add_heading("Section 2 : La langue des signes et la communauté sourde au Bénin", level=1, size=14)
add_horizontal_rule()

add_heading("2.1 Contexte historique de la surdité au Bénin", level=2, size=13)
add_justified(
    "Le Bénin, pays d'Afrique de l'Ouest situé entre le Nigéria et le Togo, "
    "comptait environ 13 millions d'habitants en 2023 (ONU, 2023). Sa population, "
    "majoritairement jeune, est marquée par une grande diversité ethnique et "
    "linguistique (42 groupes ethniques, dont les Fon, les Yoruba, les Bariba "
    "et les Dendi)."
)
add_justified(
    "Historiquement, les personnes sourdes au Bénin étaient confrontées à une "
    "stigmatisation sociale profonde, souvent liée à des croyances traditionnelles "
    "associant la surdité à des causes surnaturelles (mauvais sorts, punitions "
    "divines) dans les cultes Vodoun et les pratiques géomantiques du Fa "
    "(Prévot, 2011). Cette stigmatisation induisait des comportements "
    "d'exclusion sociale, de déscolarisation et de marginalisation professionnelle."
)
add_justified(
    "L'histoire de l'éducation des sourds au Bénin est intimement liée à la "
    "création du CAEIS (Centre d'Accueil, d'Éducation et d'Intégration des "
    "Sourds) de Louho, quartier de Porto-Novo, fondé en 1993 par les membres "
    "de l'ASUNOES (Association Universelle qui Œuvre pour l'Épanouissement des "
    "Sourds), une ONG composée de jeunes étudiants désireux de défendre les "
    "droits des personnes sourdes (Prévot, 2011 ; Djissenou et al., 2021)."
)

add_heading("2.2 Le CAEIS de Louho : un modèle d'intégration", level=2, size=13)
add_justified(
    "Le CAEIS de Louho constitue la principale structure éducative dédiée aux "
    "enfants sourds au Bénin. Depuis 1994, le centre accueille à la fois des "
    "enfants sourds et des enfants entendants dans les mêmes classes, de la "
    "maternelle à la terminale, adoptant ainsi une approche d'éducation inclusive "
    "bilingue LSF/français (Djissenou et al., 2021)."
)
add_justified(
    "Le centre comptait, en 2018, 515 apprenants dont 235 déficients auditifs. "
    "Son approche pédagogique innovante permet aux enfants sourds de réussir "
    "sur les plans scolaire et parascolaire, bien que des difficultés persistent "
    "liées au manque de formation des enseignants et à l'insuffisance de "
    "matériels didactiques adaptés (Djissenou et al., 2021)."
)
add_justified(
    "La langue des signes utilisée dans ce contexte béninois est la LSF, "
    "introduite par les associations françaises partenaires comme ASUNOES-France. "
    "Il existe cependant une émergence de variantes locales qui tendent à "
    "constituer une Langue des Signes Béninoise (LSB) propre, mêlant LSF et "
    "innovations locales (Prévot, 2011)."
)

add_heading("2.3 Le besoin technologique dans le contexte béninois", level=2, size=13)
add_justified(
    "Malgré les progrès éducatifs accomplis, les personnes sourdes béninoises "
    "restent confrontées à de sérieuses barrières de communication dans leur "
    "vie quotidienne : accès aux services administratifs, échanges commerciaux, "
    "communication médicale. Les interprètes en langue des signes sont rares "
    "et leur service onéreux."
)
add_justified(
    "Le développement d'un système de traduction automatique de la LSF vers le "
    "français répond directement à ce besoin. Les contraintes du contexte "
    "béninois imposent cependant des choix technologiques spécifiques : "
    "fonctionnement sans équipement coûteux (simple webcam), "
    "légèreté computationnelle pour des appareils d'entrée de gamme, "
    "interface intuitive ne nécessitant pas de compétences informatiques "
    "avancées."
)

# SECTION 3
add_heading("Section 3 : État de l'art des systèmes de reconnaissance de la langue des signes", level=1, size=14)
add_horizontal_rule()

add_heading("3.1 Évolution historique des approches", level=2, size=13)
add_justified(
    "Les premières approches de reconnaissance automatique de la langue des "
    "signes, dans les années 1990, reposaient sur des capteurs spécialisés : "
    "gants de données (data gloves) équipés de capteurs de flexion et de "
    "position (Fels & Hinton, 1993), puis caméras de profondeur (Microsoft "
    "Kinect) permettant d'obtenir des données 3D du squelette. Ces systèmes, "
    "bien que précis, n'étaient pas adaptés à un usage quotidien en raison "
    "de leur coût et de leur contrainte matérielle."
)
add_justified(
    "L'avènement de méthodes de vision par ordinateur sans marqueur (markerless) "
    "a permis d'abandonner progressivement ces équipements au profit d'une simple "
    "caméra RGB. Les approches statistiques comme les modèles de Markov cachés "
    "(HMM) ont d'abord dominé, combinés avec des descripteurs artisanaux "
    "(HOG, optical flow) pour décrire le mouvement des mains (Koller et al., 2015)."
)

add_heading("3.2 Approches par réseaux de neurones profonds", level=2, size=13)
add_justified(
    "L'apprentissage profond a profondément transformé les performances en "
    "reconnaissance de la langue des signes. Les architectures CNN-RNN, "
    "combinant des convolutions pour l'extraction de caractéristiques spatiales "
    "et des réseaux récurrents pour la modélisation temporelle, ont produit "
    "des résultats remarquables sur les corpus de référence (Koller et al., "
    "2016 ; Huang et al., 2018)."
)
add_justified(
    "Les travaux de Koller et al. (2016, 2017) ont démontré l'efficacité de la "
    "combinaison CNN-HMM puis CNN-BiLSTM-CTC pour la reconnaissance continue "
    "de la langue des signes sur le corpus RWTH-PHOENIX-Weather 2014 (allemand), "
    "atteignant des taux de reconnaissance compétitifs en comparaison aux "
    "systèmes à base de HMM traditionnels."
)
add_justified(
    "Les transformers (Vaswani et al., 2017), introduits pour le traitement du "
    "langage naturel, ont été adaptés à la langue des signes par Camgoz et al. "
    "(2020) dans l'architecture Sign Language Transformers, permettant une "
    "traduction bout-en-bout de la langue des signes vers le texte. Ces "
    "modèles, néanmoins, nécessitent d'importants corpus annotés pour "
    "l'entraînement et une puissance de calcul importante pour l'inférence "
    "en temps réel."
)

add_heading("3.3 Approches basées sur l'estimation de pose", level=2, size=13)
add_justified(
    "L'émergence d'outils d'estimation de pose efficaces comme OpenPose "
    "(Cao et al., 2019) et MediaPipe (Lugaresi et al., 2019) a ouvert une "
    "nouvelle voie pour la reconnaissance de la langue des signes : plutôt "
    "que de traiter directement les pixels d'une image, on extrait d'abord "
    "un squelette normalisé (keypoints), puis on entraîne un modèle sur "
    "ces représentations compactes."
)
add_justified(
    "Cette approche présente plusieurs avantages : invariance à l'apparence "
    "(vêtements, couleur de peau), normalisation spatiale intégrée, "
    "réduction de la dimensionnalité des données, et compatibilité avec "
    "des architectures légères adaptées au temps réel. Des travaux comme "
    "ceux de Jiang et al. (2021) ont montré que les représentations squelettiques "
    "combinées à des LSTM ou des Graph Neural Networks (GNN) atteignent des "
    "performances comparables aux approches basées sur l'image brute, avec "
    "une fraction du coût computationnel."
)

add_heading("3.4 Systèmes pour les langues de signes africaines", level=2, size=13)
add_justified(
    "La recherche en reconnaissance automatique de la langue des signes a "
    "été dominée par les langues de signes bien dotées en ressources : "
    "ASL (américaine), GSL (allemande), CSL (chinoise). Les langues de signes "
    "africaines restent très peu étudiées, essentiellement en raison de l'absence "
    "de corpus annotés publiquement disponibles."
)
add_justified(
    "Quelques travaux pionniers existent pour des langues de signes africaines : "
    "Ojha et al. (2021) pour la langue des signes indienne, Bangham et al. "
    "pour la Kenya Sign Language. Pour la LSF dans le contexte béninois, "
    "à notre connaissance, aucun système de traduction automatique n'a encore "
    "été publié à ce jour. Notre travail constitue donc une contribution "
    "originale à ce domaine."
)

add_heading("3.5 Tableau comparatif et positionnement", level=2, size=13)
add_justified("Le tableau suivant synthétise les principales approches existantes et positionne notre travail :", space_before=4, space_after=4)
add_table(
    ["Approche", "Entrée", "Modèle", "Dataset", "Latence", "Sans capteur spécialisé"],
    [
        ["Fels & Hinton (1993)", "Gant", "MLP", "ASL", "Non temps réel", "Non"],
        ["Koller et al. (2016)", "Vidéo", "CNN+HMM", "PHOENIX-14", "Non temps réel", "Oui"],
        ["Koller et al. (2017)", "Vidéo", "CNN-BiLSTM-CTC", "PHOENIX-14", "Partiel", "Oui"],
        ["Camgoz et al. (2020)", "Vidéo", "Transformer", "PHOENIX-14T", "Partiel", "Oui"],
        ["Jiang et al. (2021)", "Squelette", "GNN+LSTM", "CSL", "Temps réel", "Oui"],
        ["Notre approche (V11)", "Squelette (MediaPipe)", "BiLSTM-CTC", "LSF synthétique", "Temps réel (<1s)", "Oui"],
    ]
)

print("Chapitre 1 done. Writing Chapitre 2...")

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPITRE 2 : ARCHITECTURE ET MÉTHODOLOGIE
# ═══════════════════════════════════════════════════════════════════════════════
add_chapter_title("CHAPITRE 2 : ARCHITECTURE MATÉRIELLE, LOGICIELLE ET MÉTHODOLOGIE DE MISE EN ŒUVRE")
add_paragraph()

# SECTION 1
add_heading("Section 1 : Architecture matérielle", level=1, size=14)
add_horizontal_rule()

add_heading("1.1 Environnement de développement", level=2, size=13)
add_justified(
    "Le développement du système SignInterpreter V11 a été réalisé sur un "
    "MacBook Pro équipé d'une puce Apple Silicon M1 (8 cœurs CPU, 8 cœurs GPU "
    "intégrés) et de 16 Go de mémoire unifiée. Cette architecture ARM64 "
    "dispose d'une technologie d'accélération GPU propre à Apple, les Metal "
    "Performance Shaders (MPS), supportée nativement par PyTorch depuis "
    "la version 1.12. Ce choix matériel s'avère pertinent pour le développement "
    "local d'un système à ressources limitées, reflétant les contraintes "
    "matérielles typiques du contexte africain."
)
add_table(
    ["Composant", "Spécification"],
    [
        ["Processeur", "Apple M1 (8 cœurs performance + efficacité)"],
        ["Mémoire RAM", "16 Go (unifiée CPU/GPU)"],
        ["Système d'exploitation", "macOS Sonoma (Darwin 24.1.0)"],
        ["Accélération ML", "Metal Performance Shaders (MPS)"],
        ["Stockage", "SSD NVMe"],
        ["Webcam", "Caméra intégrée FaceTime HD (720p, 30 fps)"],
    ]
)

add_heading("1.2 Environnement d'entraînement cloud", level=2, size=13)
add_justified(
    "Pour l'entraînement intensif du modèle BiLSTM-CTC sur 46 000 séquences, "
    "Google Colaboratory (Colab) a été utilisé, offrant un accès gratuit à "
    "des GPU NVIDIA T4 ou V100. Cette approche hybride (développement local "
    "MPS + entraînement cloud GPU) est particulièrement adaptée au contexte "
    "africain où les GPU locaux dédiés au deep learning sont rarement disponibles."
)

add_heading("1.3 Matériel de collecte de données", level=2, size=13)
add_justified(
    "La collecte des templates de signes a été réalisée avec une webcam standard "
    "(720p, 30 fps), sans équipement spécialisé. Cette contrainte volontaire "
    "garantit que les données reflètent les conditions réelles d'utilisation "
    "du système final. Les enregistrements ont été effectués dans des conditions "
    "d'éclairage naturel variées afin d'assurer la robustesse du système à "
    "des environnements différents."
)

# SECTION 2
add_heading("Section 2 : Architecture logicielle", level=1, size=14)
add_horizontal_rule()

add_heading("2.1 Technologies principales", level=2, size=13)
add_justified(
    "Le système est développé en Python 3.11, qui offre un écosystème riche "
    "en bibliothèques d'apprentissage profond et de vision par ordinateur. "
    "Le tableau suivant présente les technologies et bibliothèques utilisées :"
)
add_table(
    ["Composant", "Technologie", "Version", "Rôle"],
    [
        ["Apprentissage profond", "PyTorch", "≥2.0", "Modèles BiLSTM-CTC et Seq2Seq"],
        ["Extraction de pose", "MediaPipe", "≥0.10", "Landmarks corps et mains (côté client)"],
        ["Serveur web", "FastAPI", "≥0.100", "API REST + WebSocket"],
        ["Serveur ASGI", "Uvicorn", "≥0.20", "Serveur asynchrone Python"],
        ["Calcul numérique", "NumPy", "≥1.24", "Traitement des features"],
        ["Interface client", "JavaScript / HTML5", "ES2022", "Capture webcam, MediaPipe, UI"],
        ["Gestion données", "Python csv, json, pickle", "-", "Lecture/écriture des données"],
    ]
)

add_heading("2.2 Organisation du code source", level=2, size=13)
add_justified("Le projet est organisé selon une architecture modulaire claire :")
modules = [
    ("ml/", "Modules d'intelligence artificielle (moteurs d'inférence, modèles)"),
    ("ml/ctc_model.py", "Modèle BiLSTM-CTC et moteur d'inférence CTCEngine"),
    ("ml/seq2seq_model.py", "Modèle seq2seq LSTM avec attention de Bahdanau"),
    ("ml/features.py", "Module d'extraction et de normalisation des caractéristiques"),
    ("ml/sliding_window.py", "Classificateur à fenêtre glissante (fallback)"),
    ("ml/translator.py", "Module de traduction LSF→Français (rule_translate)"),
    ("ml/lstm_engine.py", "Moteur LSTMEngine (classificateur par KNN)"),
    ("server/", "Serveur FastAPI"),
    ("server/routers/inference_ws.py", "Endpoint WebSocket temps réel"),
    ("scripts/", "Scripts d'entraînement et de génération de données"),
    ("scripts/generate_ctc_data.py", "Génération de données CTC synthétiques"),
    ("scripts/train_ctc.py", "Entraînement du modèle BiLSTM-CTC"),
    ("scripts/train_seq2seq.py", "Entraînement du modèle seq2seq"),
    ("data/", "Données (templates, modèles entraînés, CSV)"),
    ("client/", "Interface navigateur (JavaScript + HTML)"),
]
add_table(["Chemin", "Description"], modules)

# SECTION 3
add_heading("Section 3 : Architecture logique du système et pipeline de traitement", level=1, size=14)
add_horizontal_rule()

add_heading("3.1 Vue d'ensemble du pipeline", level=2, size=13)
add_justified(
    "Le système SignInterpreter V11 est organisé autour d'un pipeline à sept "
    "étapes qui transforme un flux vidéo brut en une phrase française :"
)
pipeline_steps = [
    "Capture vidéo : la webcam du navigateur capture 30 images par seconde via l'API WebRTC ;",
    "Extraction de landmarks : MediaPipe Holistic extrait les points anatomiques clés de chaque image directement dans le navigateur (traitement côté client) ;",
    "Encodage de caractéristiques : les landmarks sont normalisés et encodés en un vecteur de 171 dimensions (module features.py) ;",
    "Transmission WebSocket : le vecteur de caractéristiques est envoyé au serveur FastAPI via une connexion WebSocket persistante ;",
    "Reconnaissance de signes : le serveur accumule les vecteurs dans un buffer et décode la séquence de signes avec le modèle BiLSTM-CTC ;",
    "Traduction : la séquence de signes reconnue est traduite en français par le module rule_translate ;",
    "Affichage : la phrase traduite et les signes reconnus sont envoyés au client via WebSocket pour affichage en temps réel.",
]
for step in pipeline_steps:
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(step)
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

add_heading("3.2 Module d'extraction de caractéristiques", level=2, size=13)
add_justified(
    "Le module d'extraction de caractéristiques (features.py) est au cœur du "
    "pipeline multimodal. Il transforme les landmarks bruts de MediaPipe en un "
    "vecteur compact de 171 dimensions par image, représentant à la fois les "
    "informations corporelles et manuelles."
)
add_justified(
    "Le vecteur de caractéristiques est structuré en trois blocs distincts, "
    "selon une philosophie de séparation intentionnelle des signaux :"
)

add_heading("Bloc 1 — Pose corporelle supérieure (39 dimensions)", level=3, size=12)
add_justified(
    "13 landmarks anatomiques du haut du corps × 3 coordonnées (x, y, z). "
    "Les landmarks sélectionnés sont : nez, œil gauche, œil droit, oreille "
    "gauche, oreille droite, épaule gauche, épaule droite, coude gauche, "
    "coude droit, poignet gauche, poignet droit, auriculaire gauche, "
    "auriculaire droit. Les coordonnées sont normalisées par le point milieu "
    "des épaules (origine) et la distance inter-épaules (échelle), "
    "assurant l'invariance à la position et à la taille du signeur devant "
    "la caméra."
)
add_justified(
    "Ce bloc encode la localisation des mains dans l'espace du corps (paramètre "
    "d'emplacement en phonologie de la LSF). Il est crucial car deux signes peuvent "
    "avoir une configuration de main identique mais différer uniquement par "
    "l'emplacement (ex. : 'moi' vs 'toi')."
)

add_heading("Bloc 2 & 3 — Configuration des mains (66 dimensions chacun)", level=3, size=12)
add_justified(
    "Pour chaque main (gauche et droite) : 21 landmarks de la main × 3 "
    "coordonnées + 3 valeurs d'orientation de la paume = 66 dimensions. "
    "Les coordonnées sont normalisées par la position du poignet (origine) "
    "et la taille propre de la main (distance poignet-majeur), "
    "puis pondérées par un facteur HAND_WEIGHT pour équilibrer leur "
    "contribution relative avec le bloc de pose."
)
add_justified(
    "Cette normalisation indépendante des mains par rapport à la pose globale "
    "est un choix architectural clé : elle permet au modèle de voir clairement "
    "la configuration des doigts (poing, main ouverte, pointage, etc.) sans "
    "que les caractéristiques fines des doigts ne soient écrasées par l'échelle "
    "plus grande des mouvements du bras."
)

add_heading("3.3 Architecture du modèle BiLSTM-CTC", level=2, size=13)
add_justified(
    "Le modèle de reconnaissance de signes est un réseau BiLSTM-CTC constitué "
    "de deux couches LSTM bidirectionnelles suivies d'une projection linéaire "
    "vers l'espace des classes. Son architecture complète est la suivante :"
)
add_table(
    ["Couche", "Configuration", "Sortie"],
    [
        ["Entrée", "Vecteur de 171 features par frame", "(B, T, 171)"],
        ["BiLSTM couche 1", "hidden=256, bidirectionnel, dropout=0.3", "(B, T, 512)"],
        ["BiLSTM couche 2", "hidden=256, bidirectionnel, dropout=0.3", "(B, T, 512)"],
        ["Projection linéaire", "Linear(512, 31)", "(B, T, 31)"],
        ["Log-Softmax + permutation", "Probabilités log par token", "(T, B, 31)"],
    ]
)
add_justified(
    "La sortie du modèle est une distribution de probabilité sur 31 tokens "
    "à chaque pas de temps : 30 classes de signes + 1 token 'blank' (indice 30). "
    "Pendant l'entraînement, la perte CTC est calculée en marginalisant sur tous "
    "les alignements possibles entre les frames vidéo et la séquence de labels. "
    "Pendant l'inférence, le décodage greedy sélectionne le label le plus probable "
    "à chaque frame, puis supprime les blanks et les répétitions consécutives."
)
add_justified(
    "Le modèle compte au total 2 471 455 paramètres entraînables, "
    "une taille délibérément modeste pour garantir une inférence rapide "
    "sur CPU (< 50 ms par séquence) et la compatibilité avec des appareils "
    "aux ressources limitées."
)

add_heading("3.4 Génération de données synthétiques", level=2, size=13)
add_justified(
    "La rareté des données annotées est l'un des principaux obstacles au "
    "développement de systèmes de reconnaissance de la langue des signes "
    "pour des langues peu dotées. Notre approche contourne ce problème par "
    "la génération de séquences synthétiques à partir d'un corpus limité "
    "d'enregistrements réels."
)

add_heading("Corpus de templates réels", level=3, size=12)
add_justified(
    "La base du pipeline de génération est un corpus de 2 100 enregistrements "
    "réels, couvrant 30 classes de signes avec environ 70 templates par classe. "
    "Chaque template est une séquence de vecteurs de 171 caractéristiques "
    "correspondant à une exécution unique d'un signe par un signeur. "
    "Ces templates ont été collectés dans des conditions variées (vitesse, "
    "amplitude, éclairage) pour maximiser la diversité inter-template."
)

add_heading("Pipeline d'augmentation", level=3, size=12)
add_justified("Pour chaque séquence synthétique générée, le pipeline procède comme suit :")
aug_steps = [
    "Sélection pondérée d'une combinaison de signes : tirée parmi les 6 209 paires LSF-Français valides, avec un poids inversement proportionnel à la fréquence du signe le plus rare de la combinaison (favorise les signes sous-représentés) ;",
    "Sélection round-robin des templates : pour chaque signe de la séquence, le template suivant dans l'ordre cyclique est sélectionné, garantissant une couverture uniforme de tous les templates disponibles ;",
    "Rééchantillonnage temporel (time-stretching) : chaque template est rééchantillonné linéairement par un facteur aléatoire entre 0,50 et 1,60, simulant des vitesses d'exécution variables (signe rapide ou lent) ;",
    "Interpolation aux transitions : 3 frames de transition sont interpolées linéairement entre le dernier frame d'un signe et le premier frame du suivant, simulant les mouvements de liaison naturels entre les signes ;",
    "Ajout de bruit gaussien : un bruit de distribution normale (σ = 0,008) est ajouté à toutes les coordonnées, simulant des variations de détection des landmarks.",
]
for step in aug_steps:
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(step)
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

add_justified(
    "Un aspect important de la génération est l'inclusion de 6 000 séquences "
    "mono-signe (200 par classe × 30 classes), conçues pour corriger le biais "
    "de position. Sans cet ajout, le modèle aurait tendance à ne reconnaître "
    "les verbes (souvent en position médiane) qu'au milieu des séquences, "
    "et non en première position."
)
add_justified(
    "Le pipeline de génération vérifie systématiquement la contrainte CTC : "
    "pour une séquence de T frames et L labels, il faut T ≥ 2L - 1 afin "
    "que la perte CTC soit définie. Aucune violation n'a été observée sur "
    "les 46 000 séquences générées."
)

add_heading("3.5 Entraînement du modèle BiLSTM-CTC", level=2, size=13)
add_justified("Les hyperparamètres d'entraînement ont été déterminés par expérimentation :")
add_table(
    ["Hyperparamètre", "Valeur"],
    [
        ["Taille du batch", "32"],
        ["Optimiseur", "Adam (lr = 1×10⁻³)"],
        ["Planificateur LR", "ReduceLROnPlateau (patience=10, factor=0.5)"],
        ["Early stopping", "patience = 20 epochs"],
        ["Dropout", "0.3"],
        ["Dimension cachée BiLSTM", "256"],
        ["Nombre de couches BiLSTM", "2"],
        ["Proportion train/validation", "90% / 10%"],
        ["Nombre max d'epochs", "150"],
    ]
)
add_justified(
    "La fonction de perte utilisée est la perte CTC native de PyTorch "
    "(torch.nn.CTCLoss), calculée sur des mini-batchs de 32 séquences "
    "de longueurs variables (avec padding). Le gradient est clipé à une norme "
    "maximale de 1.0 pour éviter le gradient explosion, fréquent dans les LSTM."
)

add_heading("3.6 Pipeline de reconnaissance en streaming", level=2, size=13)
add_justified(
    "L'inférence en temps réel repose sur un mécanisme de streaming CTC qui "
    "accumule progressivement les frames reçues via WebSocket et décode la "
    "séquence de signes à intervalles réguliers."
)
add_justified("Le pipeline de streaming fonctionne selon les règles suivantes :")
stream_rules = [
    "Accumulation : chaque frame reçue est ajoutée au buffer CTC (ctc_buf) ;",
    "Décodage périodique : lorsque le compteur de stride atteint 20 frames (CTC_STRIDE = 20, soit ~0,67 s à 30 fps), une passe CTC greedy est exécutée sur l'ensemble du buffer ;",
    "Déclenchement minimal : le premier décodage n'est effectué qu'après 30 frames au minimum (CTC_MIN_FRAMES = 30, soit ~1 s) pour éviter les prédictions sur des données insuffisantes ;",
    "Politique 'extend only' : les nouveaux signes détectés sont ajoutés à la liste current_signs uniquement si le décodage actuel produit plus de signes que le décodage précédent — les corrections rétroactives sont interdites pour éviter les instabilités ;",
    "Déduplification : les signes consécutifs identiques dans le décodage CTC sont fusionnés avant l'ajout (ex. : ['bonjour', 'bonjour', 'moi'] → ['bonjour', 'moi']) ;",
    "Finalisation : une phrase est finalisée (et le buffer réinitialisé) lors de la détection d'une pause (main absente ≥ 0,8 s), d'une posture de repos (poignets en position basse), ou d'un timeout d'inactivité (3 s sans nouveau signe).",
]
for rule in stream_rules:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(rule)
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

add_heading("3.7 Module de traduction LSF→Français (rule_translate)", level=2, size=13)
add_justified(
    "La traduction d'une séquence de signes LSF vers le français pose le "
    "problème de la divergence grammaticale entre les deux langues. La LSF "
    "utilise un ordre SOV flexible, exploite l'espace de signation pour les "
    "références pronominales, et n'intègre pas les articles ni les conjugaisons "
    "verbales de la même façon que le français."
)
add_justified(
    "Le module rule_translate implémente un ensemble de règles linguistiques "
    "qui transforment une séquence de signes en une phrase française "
    "grammaticalement correcte. Les principales règles appliquées sont :"
)
rules = [
    "Identification du sujet : 'moi' → 'Je', 'tu' → 'Tu', 'nous' → 'Nous' ; ajout du pronom sujet si absent ;",
    "Conjugaison des verbes : les verbes sont conjugués au présent en accord avec le sujet identifié ;",
    "Reconstruction SVO : reordonnancement des constituants selon la structure sujet-verbe-objet du français ;",
    "Gestion des questions : si 'comment' est en position initiale et qu'un autre signe suit, la phrase est mise en forme interrogative ; seul, il produit 'Comment ?' ;",
    "Gestion des adjectifs prédicatifs : insertion de 'est' entre le sujet et l'adjectif (ex. : 'moi bien' → 'Je vais bien') ;",
    "Salutations : 'bonjour', 'au revoir', 'merci' sont capitalisés et suivis d'un point d'exclamation ;",
    "Ponctuation : ajout du point final ou du point d'exclamation approprié.",
]
for rule in rules:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(rule)
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

add_justified(
    "Ces règles couvrent les 6 209 combinaisons de signes du corpus "
    "training_pairs.csv avec une cohérence grammaticale vérifiée. "
    "Une analyse comparative avec un modèle de langue statistique et un "
    "modèle seq2seq (détaillée au chapitre 3) a confirmé que cette approche "
    "basée sur des règles constitue la solution la plus fiable et la plus "
    "explicable pour ce vocabulaire de 30 signes."
)

add_heading("3.8 Architecture WebSocket temps réel", level=2, size=13)
add_justified(
    "La communication entre le client navigateur et le serveur d'inférence "
    "repose sur le protocole WebSocket (RFC 6455), qui offre une connexion "
    "bidirectionnelle persistante avec une latence minimale, sans le surcoût "
    "du protocole HTTP à chaque échange."
)
add_justified(
    "L'endpoint WebSocket (/ws/inference) est géré par le serveur FastAPI "
    "de manière asynchrone (asyncio). Pour chaque client connecté, le serveur "
    "maintient un état individuel complet : buffer CTC, signes en cours, "
    "phrases complétées, timestamps. "
    "Les messages échangés sont structurés en JSON avec les types suivants :"
)
add_table(
    ["Direction", "Type de message", "Contenu"],
    [
        ["Client → Serveur", "features", "Vecteur de 171 floats + hand_visible"],
        ["Client → Serveur", "finalize_sentence", "Demande de finalisation manuelle"],
        ["Client → Serveur", "clear_sentence", "Réinitialisation complète"],
        ["Serveur → Client", "sentence_update", "current_signs + phrases complétées"],
        ["Serveur → Client", "prediction", "Prédiction live + confiance (sliding window)"],
        ["Serveur → Client", "status", "Énergie de mouvement + état de la main"],
        ["Serveur → Client", "model_info", "Informations sur les modèles chargés"],
    ]
)

print("Chapitre 2 done. Writing Chapitre 3...")

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPITRE 3 : RÉSULTATS ET DISCUSSIONS
# ═══════════════════════════════════════════════════════════════════════════════
add_chapter_title("CHAPITRE 3 : RÉSULTATS ET DISCUSSIONS")
add_paragraph()

# SECTION 1
add_heading("Section 1 : Construction et caractéristiques du jeu de données", level=1, size=14)
add_horizontal_rule()

add_heading("1.1 Corpus de templates réels", level=2, size=13)
add_justified(
    "La constitution du corpus de templates réels a constitué la première étape "
    "et l'une des plus laborieuses du projet. Chaque template est un enregistrement "
    "vidéo d'un signe unique exécuté par un signeur, traité par MediaPipe pour "
    "produire une séquence de vecteurs de 171 caractéristiques."
)
add_justified(
    "Le tableau suivant présente les 30 signes du vocabulaire avec le nombre de "
    "templates disponibles pour chacun. L'objectif de 70 templates par classe a "
    "été atteint pour la quasi-totalité des signes :"
)
templates_data = [
    ["aide", "70"], ["aimer", "70"], ["aller", "70"], ["ami", "65"],
    ["amour", "70"], ["apprendre", "70"], ["attendre", "66"],
    ["au revoir", "70"], ["aujourd'hui", "73"], ["avec", "70"],
    ["bébé", "70"], ["bien", "70"], ["bientôt", "70"], ["boire", "74"],
    ["bon", "75"], ["bonjour", "70"], ["changer", "70"], ["chat", "70"],
    ["comment", "72"], ["comprendre", "70"], ["donner", "70"],
    ["ici", "70"], ["manger", "70"], ["merci", "70"], ["moi", "70"],
    ["non", "69"], ["nous", "70"], ["oui", "70"], ["tu", "69"],
    ["vouloir", "67"],
]
add_table(
    ["Signe", "Templates", "Signe", "Templates", "Signe", "Templates"],
    [
        [templates_data[i][0], templates_data[i][1],
         templates_data[i+10][0], templates_data[i+10][1],
         templates_data[i+20][0], templates_data[i+20][1]]
        for i in range(10)
    ]
)
add_justified(
    "Au total, le corpus compte 2 100 templates réels pour 30 classes, "
    "soit une moyenne de 70 templates par classe (min : 65 pour 'ami', "
    "max : 75 pour 'bon'). Cette distribution quasi-uniforme est le résultat "
    "d'un effort de collecte équilibré."
)

add_heading("1.2 Corpus synthétique d'entraînement", level=2, size=13)
add_justified(
    "À partir des 2 100 templates réels, le pipeline de génération synthétique "
    "a produit 46 000 séquences d'entraînement selon la répartition suivante :"
)
add_table(
    ["Type de séquences", "Nombre", "Description"],
    [
        ["Mono-signe", "6 000", "200 séquences par classe × 30 classes (correction biais position-0)"],
        ["Multi-signes", "40 000", "Tirées parmi les 6 209 paires LSF-Français valides"],
        ["Total", "46 000", "Corpus complet d'entraînement"],
    ]
)
add_justified(
    "La longueur des séquences générées varie selon le nombre de signes et "
    "le facteur de rééchantillonnage : les séquences mono-signe ont une longueur "
    "moyenne d'environ 25 à 40 frames, tandis que les séquences multi-signes "
    "(2 à 5 signes) atteignent en moyenne 80 à 200 frames. La contrainte CTC "
    "(T ≥ 2L - 1) a été vérifiée sur l'ensemble du corpus avec zéro violation."
)
add_justified(
    "Le corpus est stocké en format float16 (demi-précision) pour réduire "
    "l'empreinte mémoire tout en préservant la précision nécessaire. "
    "La taille totale du fichier pickle sérialisé est de l'ordre de 2 Go."
)

add_heading("1.3 Corpus de paires LSF-Français", level=2, size=13)
add_justified(
    "Le fichier training_pairs.csv contient 6 209 paires (séquence de signes, "
    "phrase française) couvrant les combinaisons linguistiquement valides des "
    "30 signes du vocabulaire. Ces paires ont été construites par combinaison "
    "exhaustive des signes selon des patrons grammaticaux de la LSF, avec "
    "traduction vers le français par le module rule_translate."
)
add_justified(
    "La couverture bigramme — proportion des bigrammes de signes possibles "
    "couverts par au moins une séquence d'entraînement — atteint 58,7 %, "
    "contre 36,1 % avant l'extension du corpus. Cette métrique reflète la "
    "capacité du système à généraliser à de nouvelles combinaisons de signes."
)

# SECTION 2
add_heading("Section 2 : Entraînement et évaluation du modèle BiLSTM-CTC", level=1, size=14)
add_horizontal_rule()

add_heading("2.1 Résultats d'entraînement", level=2, size=13)
add_justified(
    "Le modèle BiLSTM-CTC a été entraîné sur Google Colaboratory (GPU NVIDIA T4) "
    "pendant 80 epochs (early stopping activé avec patience=20). La convergence "
    "a été rapide grâce à la qualité du corpus synthétique et à l'architecture "
    "BiLSTM bien adaptée au problème."
)
add_table(
    ["Métrique", "Valeur"],
    [
        ["Précision sur les signes individuels (val)", "99,97 %"],
        ["Nombre d'epochs jusqu'à convergence", "80 (early stopping)"],
        ["Architecture", "BiLSTM (2 couches, hidden=256)"],
        ["Nombre de paramètres", "2 471 455"],
        ["Temps d'entraînement (GPU T4)", "~45 minutes"],
        ["Taille du modèle sauvegardé", "~10 Mo"],
    ]
)

add_heading("2.2 Analyse des performances", level=2, size=13)
add_justified(
    "La précision de 99,97 % sur les signes individuels, mesurée sur le jeu "
    "de validation (10 % du corpus), démontre que le modèle a correctement "
    "appris à discriminer les 30 classes de signes dans des conditions variées "
    "de vitesse et de bruit. Cette performance élevée s'explique par plusieurs "
    "facteurs :"
)
factors = [
    "La qualité de la représentation (171 features découplant pose corporelle et configuration des mains) qui capture les paramètres phonologiques essentiels des signes ;",
    "La diversité du corpus synthétique (facteur vitesse 0,50–1,60×) qui assure la robustesse aux variations d'exécution ;",
    "L'architecture BiLSTM qui capture les dépendances temporelles dans les deux directions, utile pour interpréter le début et la fin d'un signe ;",
    "Les 6 000 séquences mono-signe qui ont éliminé le biais position-0 observé dans les versions précédentes du système.",
]
for f in factors:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f)
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

add_heading("2.3 Comparaison avec l'architecture Transformer", level=2, size=13)
add_justified(
    "Au cours du développement, une architecture Transformer CTC a également "
    "été expérimentée, remplaçant le BiLSTM par un encodeur à auto-attention "
    "multi-têtes. Bien que le Transformer ait atteint une précision de 95,96 % "
    "sur le jeu de validation synthétique, son déploiement a révélé un problème "
    "critique de décalage de normalisation (normalization mismatch) : les "
    "statistiques de normalisation calculées sur les données d'entraînement "
    "synthétiques (moyenne ≈ -0,26) ne correspondaient pas aux distributions "
    "des features MediaPipe réelles (moyenne ≈ +0,10), conduisant à des "
    "prédictions erronées en conditions réelles."
)
add_justified(
    "Le BiLSTM, entraîné sans normalisation externe, ne souffre pas de ce "
    "problème puisqu'il traite les vecteurs de features bruts directement. "
    "Cette robustesse, combinée à une latence d'inférence inférieure, "
    "a motivé le retour à l'architecture BiLSTM-CTC."
)

# SECTION 3
add_heading("Section 3 : Évaluation du système de traduction", level=1, size=14)
add_horizontal_rule()

add_heading("3.1 Approches comparées", level=2, size=13)
add_justified(
    "Trois approches de traduction ont été développées et comparées :"
)
add_table(
    ["Approche", "Description", "Paramètres", "Val Loss"],
    [
        ["rule_translate", "Règles linguistiques LSF→Français", "0 (déterministe)", "N/A"],
        ["Modèle de langue n-gram", "Modèle de langue statistique", "~5K", "N/A"],
        ["Seq2Seq LSTM + Attention", "LSTM avec attention de Bahdanau", "1 143 360", "0,0354"],
    ]
)

add_heading("3.2 Modèle Seq2Seq LSTM avec attention de Bahdanau", level=2, size=13)
add_justified(
    "Le modèle seq2seq, entraîné sur les 6 209 paires LSF-Français, "
    "présente l'architecture suivante :"
)
add_table(
    ["Composant", "Configuration"],
    [
        ["Encodeur", "Embedding(40, 64) + BiLSTM(128, 2 couches, dropout=0.3)"],
        ["Mécanisme d'attention", "Attention de Bahdanau (couche linéaire + tanh)"],
        ["Décodeur", "Embedding(195, 64) + LSTM(128, 2 couches) + fc_out"],
        ["Vocabulaire source (signes)", "40 tokens (30 signes + tokens spéciaux)"],
        ["Vocabulaire cible (français)", "195 tokens (mots français + tokens spéciaux)"],
        ["Paramètres totaux", "1 143 360"],
        ["Val loss finale", "0,0354"],
    ]
)
add_justified(
    "Malgré une perte de validation basse (0,0354), l'évaluation qualitative "
    "a révélé que le modèle seq2seq ne surpasse pas rule_translate. "
    "Cette contre-performance apparente s'explique par la circularité des données : "
    "les phrases françaises du corpus d'entraînement ont été générées par "
    "rule_translate lui-même. Le modèle seq2seq apprend donc à approximer "
    "rule_translate avec des erreurs, sans jamais le dépasser."
)
add_justified(
    "Sur 10 paires de test, rule_translate produit la traduction attendue "
    "dans 10/10 cas, tandis que le seq2seq n'y parvient que dans 1/10 cas. "
    "Cette conclusion confirme que l'approche par règles est optimal pour "
    "un vocabulaire de 30 signes avec un corpus aligné sur ces règles."
)

add_heading("3.3 Exemples de traductions", level=2, size=13)
add_justified("Le tableau suivant présente des exemples de traductions LSF→Français produites par le système :")
add_table(
    ["Séquence de signes (LSF)", "Traduction produite"],
    [
        ["bonjour", "Bonjour !"],
        ["au revoir", "Au revoir !"],
        ["moi manger", "Je mange."],
        ["tu boire", "Tu bois."],
        ["moi aimer chat", "J'aime le chat."],
        ["comment tu", "Comment vas-tu ?"],
        ["nous aller", "Nous allons."],
        ["moi vouloir apprendre", "Je veux apprendre."],
        ["tu comprendre", "Tu comprends."],
        ["bonjour moi manger merci", "Bonjour ! Je mange. Merci !"],
    ]
)

# SECTION 4
add_heading("Section 4 : Interface et déploiement", level=1, size=14)
add_horizontal_rule()

add_heading("4.1 Interface utilisateur", level=2, size=13)
add_justified(
    "L'interface utilisateur est une application web mono-page (SPA) accessible "
    "depuis tout navigateur moderne (Chrome, Firefox, Edge) sans installation "
    "préalable. Elle intègre les composants suivants :"
)
ui_comps = [
    "Flux vidéo en direct : affichage de la webcam avec superposition des landmarks MediaPipe (squelette des mains et du corps) ;",
    "Zone de reconnaissance progressive : les signes reconnus en temps réel apparaissent progressivement au fil du streaming CTC ;",
    "Historique des phrases : les phrases finalisées sont accumulées avec leur traduction française et un score de confiance ;",
    "Indicateurs d'état : affichage de l'énergie de mouvement (barre de progression), de la visibilité de la main, et de la prédiction en cours (sliding window) ;",
    "Contrôles : bouton 'Finaliser' pour déclencher manuellement la fin d'une phrase, bouton 'Effacer' pour réinitialiser l'historique.",
]
for comp in ui_comps:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(comp)
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

add_heading("4.2 Performances en conditions réelles", level=2, size=13)
add_justified(
    "Les mesures de latence réalisées en conditions réelles (MacBook M1, "
    "navigateur Chrome, connexion locale) donnent les résultats suivants :"
)
add_table(
    ["Métrique", "Valeur mesurée"],
    [
        ["Latence MediaPipe (extraction)", "< 5 ms / frame"],
        ["Latence WebSocket (transmission)", "< 2 ms (connexion locale)"],
        ["Latence inférence BiLSTM-CTC", "< 30 ms / décodage"],
        ["Premier signe prédit", "~1 s après début du signe (30 frames min.)"],
        ["Mise à jour progressive", "~0,67 s (CTC_STRIDE = 20 frames à 30 fps)"],
        ["Latence perçue totale", "< 1 s (streaming progressif)"],
    ]
)
add_justified(
    "Ces valeurs de latence sont compatibles avec une utilisation fluide en "
    "conversation. L'approche de streaming progressif (affichage au fil du "
    "décodage) réduit considérablement la latence perçue par rapport à "
    "une attente de la phrase complète."
)

add_heading("4.3 Démarrage et déploiement", level=2, size=13)
add_justified(
    "Le système se déploie en deux commandes depuis le répertoire du projet : "
    "le script start.sh lance le serveur FastAPI (uvicorn) sur le port 8000, "
    "puis le client navigateur est accessible à l'adresse localhost:8000. "
    "L'architecture légère du projet (aucune dépendance cloud, inférence "
    "locale) assure son fonctionnement même en l'absence de connexion Internet, "
    "condition fréquente dans les zones rurales béninoises."
)

# SECTION 5
add_heading("Section 5 : Discussions et perspectives", level=1, size=14)
add_horizontal_rule()

add_heading("5.1 Forces du système", level=2, size=13)
strengths = [
    "Accessibilité matérielle : une simple webcam RGB suffit, sans capteur de profondeur ni gant de données ;",
    "Déploiement navigateur : aucune installation requise côté utilisateur, compatible avec tous les systèmes d'exploitation ;",
    "Architecture explicable : la chaîne extraction de features → BiLSTM-CTC → rule_translate est transparente et auditable, contrairement aux modèles de bout en bout ;",
    "Performance élevée : 99,97 % de précision sur les signes individuels sur le corpus de validation ;",
    "Entraînable avec des ressources limitées : un MacBook M1 suffit pour le développement ; Google Colab pour l'entraînement intensif ;",
    "Extensible : l'ajout d'un nouveau signe nécessite uniquement de collecter ~70 templates et de relancer le pipeline de génération synthétique.",
]
for s in strengths:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(s)
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

add_heading("5.2 Limitations du système", level=2, size=13)
limits = [
    "Vocabulaire limité à 30 signes : le système ne couvre qu'une fraction du vocabulaire LSF (estimé à plusieurs milliers de signes). La communication est limitée aux domaines couverts par les 30 signes ;",
    "Données d'un seul ou peu de signeurs : la diversité inter-signeurs des templates de collecte reste limitée. Un signeur non représenté dans les templates pourrait obtenir de moins bonnes performances ;",
    "Absence des éléments non-manuels : les expressions faciales, mouvements des lèvres et mouvements de tête, qui jouent un rôle grammatical en LSF (marquage des questions, de la négation), ne sont pas encore exploités ;",
    "LSF uniquement : le système n'est pas directement adapté à la Langue des Signes Béninoise (LSB) qui présente des différences lexicales par rapport à la LSF ;",
    "Sensibilité aux conditions d'éclairage : bien que MediaPipe soit robuste, une luminosité très faible ou un contre-jour fort peuvent dégrader la qualité de l'extraction des landmarks.",
]
for l in limits:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(l)
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

add_heading("5.3 Perspectives d'amélioration", level=2, size=13)
add_justified(
    "Plusieurs directions de recherche et de développement sont envisagées "
    "pour les versions futures du système :"
)
perspectives = [
    "Élargissement du vocabulaire : objectif de 200 signes à court terme, 500+ à moyen terme, en impliquant la communauté du CAEIS de Louho comme contributeurs de templates ;",
    "Adaptation à la LSB : recollecte des templates avec des signeurs utilisant la Langue des Signes Béninoise, avec des lexèmes locaux ; potentiellement un corpus de 1 000+ signes LSB ;",
    "Collecte collaborative et participative : développement d'une application mobile dédiée à la collecte de templates, permettant à la communauté sourde béninoise de contribuer directement à l'enrichissement du corpus ;",
    "Intégration des éléments non-manuels : ajout des landmarks du visage (468 points) pour capturer les expressions faciales et leur rôle grammatical en LSF ;",
    "Traduction vers la parole : intégration d'un moteur de synthèse vocale (Text-to-Speech) pour produire une sortie audio en plus du texte écrit, bénéficiant aux personnes illettrées ;",
    "Déploiement mobile : adaptation du système en Progressive Web App (PWA) pour un fonctionnement sur smartphone Android (dominant en Afrique de l'Ouest), en optimisant la taille du modèle (quantification INT8) ;",
    "Modèle de traduction neuronal : avec un corpus suffisamment large et diversifié, un modèle seq2seq entraîné sur des données réelles (non générées par rule_translate) pourrait surpasser l'approche par règles.",
]
for p_text in perspectives:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(p_text)
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

print("Chapitre 3 done. Writing Conclusion...")

# ═══════════════════════════════════════════════════════════════════════════════
# CONCLUSION
# ═══════════════════════════════════════════════════════════════════════════════
add_chapter_title("Conclusion")
add_paragraph()
add_justified(
    "Ce travail a présenté la conception, le développement et l'évaluation du "
    "système SignInterpreter V11, une solution multimodale de traduction en "
    "temps réel de la Langue des Signes Française (LSF) vers le français écrit. "
    "Motivé par les besoins réels de la communauté sourde béninoise — illustrés "
    "par l'exemple du CAEIS de Louho à Porto-Novo — et par les avancées récentes "
    "en apprentissage profond et en vision par ordinateur, ce système apporte "
    "une contribution technique et sociale originale."
)
add_justified(
    "L'approche multimodale adoptée fusionne trois niveaux complémentaires "
    "d'information : les signaux corporels (pose du haut du corps), les signaux "
    "manuels (configuration détaillée des deux mains), et le contexte linguistique "
    "(structure grammaticale de la LSF). Cette fusion est opérée au niveau des "
    "features par un vecteur de 171 dimensions extrait par MediaPipe Holistic "
    "et traité conjointement par un modèle BiLSTM-CTC."
)
add_justified(
    "Les résultats obtenus sont encourageants : une précision de 99,97 % sur "
    "les signes individuels du jeu de validation, une latence perçue inférieure "
    "à une seconde grâce au streaming progressif, et un déploiement entièrement "
    "dans le navigateur sans équipement spécialisé. L'ensemble du pipeline, "
    "de la collecte des templates à l'interface utilisateur finale, a été "
    "conçu pour fonctionner avec des ressources matérielles modestes, "
    "adaptées au contexte africain."
)
add_justified(
    "Ce travail a également mis en lumière des défis importants : la rareté "
    "des données annotées pour les langues de signes peu dotées (résolue "
    "par la génération synthétique), le risque de décalage de distribution "
    "entre données synthétiques et données réelles (ayant motivé l'abandon "
    "du Transformer au profit du BiLSTM), et les limitations inhérentes aux "
    "approches de traduction par règles pour les combinaisons de signes hors "
    "vocabulaire."
)
add_justified(
    "Les perspectives sont nombreuses et prometteuses. L'élargissement du "
    "vocabulaire, l'adaptation à la Langue des Signes Béninoise, l'intégration "
    "des éléments non-manuels et le déploiement mobile constituent des axes "
    "naturels de continuation. La dimension participative — impliquer la "
    "communauté sourde du CAEIS de Louho dans la collecte de données — "
    "est particulièrement importante pour ancrer le système dans les réalités "
    "locales et garantir son adoption."
)
add_justified(
    "Ce travail s'inscrit pleinement dans les objectifs de la Stratégie Nationale "
    "d'Intelligence Artificielle et de Mégadonnées du Bénin 2023-2027, qui "
    "promeut l'IA comme vecteur d'inclusion sociale et de développement durable. "
    "Nous espérons que ce système contribuera, même modestement, à réduire "
    "les barrières de communication entre personnes sourdes et entendantes "
    "au Bénin et dans la sous-région ouest-africaine."
)

print("Conclusion done. Writing References...")

# ═══════════════════════════════════════════════════════════════════════════════
# RÉFÉRENCES
# ═══════════════════════════════════════════════════════════════════════════════
add_section_title("Références Bibliographiques et Webographiques")

add_heading("Références bibliographiques", level=2, size=13)
biblio = [
    "Armstrong, D.F. (2011). Gesture and the Nature of Language. Cambridge University Press.",
    "Atrey, P.K., Hossain, M.A., El Saddik, A., & Kankanhalli, M.S. (2010). Multimodal fusion for multimedia analysis: a survey. Multimedia Systems, 16(6), 345-379.",
    "Bahdanau, D., Cho, K., & Bengio, Y. (2015). Neural Machine Translation by Jointly Learning to Align and Translate. ICLR 2015. arXiv:1409.0473.",
    "Battison, R. (1978). Lexical Borrowing in American Sign Language. Linstok Press.",
    "Brentari, D. (2019). Sign Languages. Cambridge Language Surveys. Cambridge University Press.",
    "Camgoz, N.C., Koller, O., Hadfield, S., & Bowden, R. (2020). Sign Language Transformers: Joint End-to-end Sign Language Recognition and Translation. CVPR 2020.",
    "Cao, Z., Hidalgo, G., Simon, T., Wei, S.E., & Sheikh, Y. (2019). OpenPose: Realtime Multi-Person 2D Pose Estimation using Part Affinity Fields. IEEE Transactions on Pattern Analysis and Machine Intelligence.",
    "Cuxac, C. (1983). Le langage des sourds. Payot, Paris.",
    "Cuxac, C. (2000). La Langue des Signes Française (LSF) : les voies de l'iconicité. Faits de Langues, Ophrys.",
    "Djissenou, A.S., Kélani, R.R. & Houessou, P. (2021). Education Inclusive En République Du Bénin : L'expérience Des Apprenants Du Centre De Malentendants De Louho À Porto-Novo. European Scientific Journal, ESJ, 17(15), 335. https://doi.org/10.19044/esj.2021.v17n15p335",
    "Fels, S.S., & Hinton, G.E. (1993). Glove-talk: A neural network interface between a data-glove and a speech synthesizer. IEEE Transactions on Neural Networks, 4(1), 2-8.",
    "Graves, A., Fernández, S., Gomez, F., & Schmidhuber, J. (2006). Connectionist Temporal Classification: Labelling Unsegmented Sequence Data with Recurrent Neural Networks. ICML 2006.",
    "Hochreiter, S. (1991). Untersuchungen zu dynamischen neuronalen Netzen. Diploma thesis, TU Munich.",
    "Hochreiter, S., & Schmidhuber, J. (1997). Long Short-Term Memory. Neural Computation, 9(8), 1735-1780.",
    "Huang, J., Zhou, W., Li, H., & Li, W. (2018). Attention-Based 3D-CNNs for Large-Vocabulary Sign Language Recognition. IEEE Transactions on Circuits and Systems for Video Technology, 29(8), 2822-2832.",
    "Jiang, S., Sun, B., Wang, L., Bai, Y., Li, K., & Fu, Y. (2021). Skeleton Aware Multi-modal Sign Language Recognition. CVPR Workshop 2021.",
    "Koller, O., Forster, J., & Ney, H. (2015). Continuous Sign Language Recognition: Towards Large Vocabulary Statistical Recognition Systems Handling Multiple Signers. Computer Vision and Image Understanding, 141, 108-125.",
    "Koller, O., Zargaran, S., Ney, H., & Bowden, R. (2016). Deep Sign: Hybrid CNN-HMM for Continuous Sign Language Recognition. BMVC 2016.",
    "Koller, O., Zargaran, O., & Ney, H. (2017). Re-sign: Re-aligned End-to-end Sequence Modelling with Deep Recurrent CNN-HMMs. CVPR 2017.",
    "Krizhevsky, A., Sutskever, I., & Hinton, G.E. (2012). ImageNet Classification with Deep Convolutional Neural Networks. NeurIPS 2012.",
    "Ladd, P. (2003). Understanding Deaf Culture: In Search of Deafhood. Multilingual Matters.",
    "LeCun, Y., Bengio, Y., & Hinton, G. (2015). Deep learning. Nature, 521, 436-444.",
    "Lugaresi, C., Tang, J., Nash, H., McClanahan, C., Uboweja, E., Hays, M., Zhang, F., Chang, C.L., Yong, M.G., Lee, J., Chang, W.T., Hua, W., Georg, M., & Grundmann, M. (2019). MediaPipe: A Framework for Perceiving and Processing Reality. Workshop on Perception for AR/VR at ICCV 2019.",
    "McCarthy, J. (1956). Proposal for the Dartmouth summer research project on artificial intelligence. Proposal.",
    "Organisation Mondiale de la Santé (OMS) (2001). Classification Internationale du Fonctionnement, du Handicap et de la Santé (CIF). OMS, Genève.",
    "Prévot, A. (2011). La surdité et les sourds au Bénin : la problématique du handicap et l'exemple du Centre d'Accueil et d'Intégration des Sourds de Louho (Porto-Novo) comme facteur de changement des représentations sociales. Mémoire de Master 1, Université Stendhal Grenoble 3. HAL: dumas-00620765.",
    "Sutskever, I., Vinyals, O., & Le, Q.V. (2014). Sequence to Sequence Learning with Neural Networks. NeurIPS 2014.",
    "Vaswani, A., Shazeer, N., Parmar, N., Uszkoreit, J., Jones, L., Gomez, A.N., Kaiser, L., & Polosukhin, I. (2017). Attention Is All You Need. NeurIPS 2017.",
]
for ref in biblio:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.first_line_indent = Cm(-0.8)
    run = p.add_run(ref)
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"

add_heading("Références webographiques", level=2, size=13)
webref = [
    "MediaPipe Holistic Solution. Google LLC. Consulté le 09 mai 2026. https://ai.google.dev/edge/mediapipe/solutions/vision/holistic_landmarker",
    "PyTorch Documentation — torch.nn.CTCLoss. Meta AI. Consulté le 09 mai 2026. https://pytorch.org/docs/stable/generated/torch.nn.CTCLoss.html",
    "FastAPI Documentation. Sebastián Ramírez. Consulté le 09 mai 2026. https://fastapi.tiangolo.com/",
    "Stratégie Nationale d'Intelligence Artificielle et de Mégadonnées du Bénin 2023-2027. Ministère du Numérique et de la Digitalisation, République du Bénin. 2023.",
    "ASUNOES — Association Universelle qui Œuvre pour l'Épanouissement des Sourds. Porto-Novo, Bénin.",
    "Organisation Mondiale de la Santé (OMS). Surdité et perte d'audition. Aide-mémoire, 2023. https://www.who.int/fr/news-room/fact-sheets/detail/deafness-and-hearing-loss",
]
for ref in webref:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.first_line_indent = Cm(-0.8)
    run = p.add_run(ref)
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"

# ═══════════════════════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════════════════════
doc.save(OUT)
print(f"\nMémoire généré : {OUT}")
print("Taille du fichier :", round(__import__('os').path.getsize(OUT)/1024, 1), "Ko")
