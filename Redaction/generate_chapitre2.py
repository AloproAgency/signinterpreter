#!/usr/bin/env python3
"""
Chapitre 2 — Matériels, méthodes et ingénierie du système SignInterpreter V11.
Ordre scientifique standard :
  Introduction → Vue d'ensemble → Matériels → Données → Caractéristiques →
  Augmentation → Modèles → Entraînement → Pipeline → Démarche itérative → Conclusion
Toutes les données proviennent du code source vérifié.
"""

import os, tempfile
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT   = "/Users/alopro/Desktop/AI/RECHERCHE/SignInterpreter/V11/Redaction/chapitre2_ingenierie.docx"
TMPDIR = tempfile.mkdtemp()

# ─── Palette cohérente avec le document ───────────────────────────────────────
C1 = "#1F3864"   # bleu nuit (principal)
C2 = "#2E75B6"   # bleu moyen
C3 = "#4BACC6"   # bleu clair
C4 = "#70AD47"   # vert
C5 = "#ED7D31"   # orange
C6 = "#A9D18E"   # vert pâle
GRID = "#E8EEF4"

plt.rcParams.update({
    "font.family":   "DejaVu Sans",
    "font.size":     9,
    "axes.spines.top":    False,
    "axes.spines.right":  False,
    "axes.grid":     True,
    "grid.color":    GRID,
    "grid.linewidth": 0.8,
})

def save_fig(name):
    path = os.path.join(TMPDIR, name)
    plt.savefig(path, dpi=160, bbox_inches='tight', facecolor='white')
    plt.close()
    return path

# ─── Chart generators ─────────────────────────────────────────────────────────

def chart_exemples_par_classe():
    classes = [
        "aide","aimer","aller","ami","amour","apprendre","attendre","au revoir",
        "aujourd'hui","avec","bébé","bien","bientôt","boire","bon","bonjour",
        "changer","chat","comment","comprendre","donner","ici","manger","merci",
        "moi","non","nous","oui","tu","vouloir",
    ]
    counts = [70,70,70,65,70,70,66,70,73,70,70,70,70,74,75,70,70,70,72,70,70,70,70,70,70,69,70,70,69,67]
    colors = [C2 if c == 70 else C5 for c in counts]

    fig, ax = plt.subplots(figsize=(7, 5.5))
    y = np.arange(len(classes))
    bars = ax.barh(y, counts, color=colors, height=0.65, edgecolor='white', linewidth=0.4)
    ax.set_yticks(y); ax.set_yticklabels(classes, fontsize=8)
    ax.set_xlabel("Nombre d'exemples enregistrés")
    ax.set_xlim(60, 78)
    ax.axvline(70, color=C1, linewidth=1.2, linestyle='--', alpha=0.6, label='Cible : 70')
    for bar, c in zip(bars, counts):
        ax.text(bar.get_width() + 0.2, bar.get_y() + bar.get_height()/2,
                str(c), va='center', fontsize=7.5, color=C1)
    patch_std = mpatches.Patch(color=C2, label='70 exemples (valeur cible)')
    patch_var = mpatches.Patch(color=C5, label='Légère variation (rejet qualité)')
    ax.legend(handles=[patch_std, patch_var], fontsize=8, loc='lower right')
    ax.set_title("Distribution des exemples enregistrés par classe\n(Total : 2 100)", fontsize=10, fontweight='bold', color=C1)
    plt.tight_layout()
    return save_fig("chart_classes.png")

def chart_paires_longueur():
    labels = ["1 signe","2 signes","3 signes","4 signes","5 signes","6 signes","7 signes","≥ 8 signes"]
    counts = [12, 140, 1099, 2042, 1530, 939, 342, 105]
    pcts   = [c/6209*100 for c in counts]
    colors = [C3, C3, C2, C1, C2, C3, C4, C4]

    fig, ax = plt.subplots(figsize=(7.5, 4))
    x = np.arange(len(labels))
    bars = ax.bar(x, counts, color=colors, width=0.65, edgecolor='white', linewidth=0.4)
    ax.set_xticks(x); ax.set_xticklabels(labels, fontsize=8.5)
    ax.set_ylabel("Nombre de paires")
    ax.set_xlabel("Longueur de la séquence de signes")
    for bar, c, p in zip(bars, counts, pcts):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 30,
                f"{c}\n({p:.1f}%)", ha='center', fontsize=7.5, color=C1)
    ax.set_title("Distribution des 6 209 paires d'entraînement\npar longueur de séquence", fontsize=10, fontweight='bold', color=C1)
    plt.tight_layout()
    return save_fig("chart_paires.png")

def chart_augmentation():
    techniques = [
        "Lissage d'images\n(fusion linéaire)",
        "Inversion de\nlatéralité (LH↔RH)",
        "Gigue temporelle\n(±15 % longueur)",
        "Mise à l'échelle\n([0.88 ; 1.12])",
        "Bruit gaussien\n(σ = 0.008)",
    ]
    probs = [40, 50, 60, 70, 80]
    colors = [C3, C4, C2, C2, C1]

    fig, ax = plt.subplots(figsize=(7, 3.8))
    y = np.arange(len(techniques))
    bars = ax.barh(y, probs, color=colors, height=0.55, edgecolor='white')
    ax.set_yticks(y); ax.set_yticklabels(techniques, fontsize=8.5)
    ax.set_xlabel("Probabilité d'application (%)")
    ax.set_xlim(0, 95)
    for bar, p in zip(bars, probs):
        ax.text(bar.get_width() + 1, bar.get_y() + bar.get_height()/2,
                f"{p} %", va='center', fontsize=9, fontweight='bold', color=C1)
    ax.set_title("Techniques d'augmentation — probabilités d'application\n(multiplicateur ×10 : 2 100 → 21 000 exemples)", fontsize=10, fontweight='bold', color=C1)
    plt.tight_layout()
    return save_fig("chart_augmentation.png")

def chart_volumes_donnees():
    fig, ax = plt.subplots(figsize=(7.5, 3.8))

    stages  = ["Enregistrements\nbruts", "LSTM\n(après augmentation)", "CTC\n(génération synthétique — 46 000)", "Seq2Seq\n(paires linguistiques)"]
    vals    = [2100, 21000, 46000, 6209]
    colors  = [C3, C2, C1, C4]

    x = np.arange(len(stages))
    bars = ax.bar(x, vals, color=colors, width=0.5, edgecolor='white')
    ax.set_xticks(x); ax.set_xticklabels(stages, fontsize=9)
    ax.set_ylabel("Nombre de séquences")
    ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda v, _: f"{int(v):,}".replace(",", " ")))

    for bar, v in zip(bars, vals):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 300,
                f"{v:,}".replace(",", " "), ha='center', fontsize=9, fontweight='bold', color=C1)

    arrows = [(0, 1, "×10 augmentation"), (0, 2, "×22 génération\nsynthétique"), (0, 3, "paires\nlinguistiques")]
    ax.set_title("Volumes de données à chaque étape de préparation", fontsize=10, fontweight='bold', color=C1)
    plt.tight_layout()
    return save_fig("chart_volumes.png")

# ─── Pre-generate all charts ──────────────────────────────────────────────────
IMG_CLASSES    = chart_exemples_par_classe()
IMG_PAIRES     = chart_paires_longueur()
IMG_AUGMENT    = chart_augmentation()
IMG_VOLUMES    = chart_volumes_donnees()
print("Charts générés.")

doc = Document()
sec = doc.sections[0]
sec.page_width  = Cm(21); sec.page_height = Cm(29.7)
sec.left_margin = sec.right_margin = Cm(2.5)
sec.top_margin  = sec.bottom_margin = Cm(2.5)

# ─── Style helpers ────────────────────────────────────────────────────────────

def add_img(path, width_cm=14.0):
    p = doc.add_paragraph()
    p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run()
    run.add_picture(path, width=Cm(width_cm))

def bg(cell, hex_c):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_c); tcPr.append(shd)

def h1(text, sb=18, sa=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    r = p.add_run(text); r.bold = True
    r.font.size = Pt(14); r.font.name = "Times New Roman"
    r.font.color.rgb = RGBColor.from_string("1F3864")

def h2(text, sb=12, sa=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    r = p.add_run(text); r.bold = True
    r.font.size = Pt(12); r.font.name = "Times New Roman"
    r.font.color.rgb = RGBColor.from_string("2E4057")

def h3(text, sb=8, sa=3):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    r = p.add_run(text); r.bold = True; r.italic = True
    r.font.size = Pt(12); r.font.name = "Times New Roman"

def j(text, indent=True, sb=0, sa=6):
    p = doc.add_paragraph()
    p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    r = p.add_run(text)
    r.font.size = Pt(12); r.font.name = "Times New Roman"

def j_b(parts, indent=True, sb=0, sa=6):
    """Mixed bold/normal paragraph. parts = [(text, bold)]"""
    p = doc.add_paragraph()
    p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(sb); p.paragraph_format.space_after = Pt(sa)
    if indent: p.paragraph_format.first_line_indent = Cm(1.25)
    for text, bold in parts:
        r = p.add_run(text); r.bold = bold
        r.font.size = Pt(12); r.font.name = "Times New Roman"

def dash(label, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1.5)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.space_before = Pt(0)
    r1 = p.add_run("– " + label + " "); r1.bold = True
    r1.font.size = Pt(12); r1.font.name = "Times New Roman"
    r2 = p.add_run(text)
    r2.font.size = Pt(12); r2.font.name = "Times New Roman"

def caption(text):
    p = doc.add_paragraph()
    p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(12)
    r = p.add_run(text); r.italic = True
    r.font.size = Pt(10); r.font.name = "Times New Roman"

def code_block(lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("\n".join(lines))
    r.font.name = "Courier New"; r.font.size = Pt(9)

def tbl(headers, rows, cap=None, widths=None, header_bg="2E4057"):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for i, hd in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = hd; c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        bg(c, header_bg)
        r = c.paragraphs[0].runs[0]
        r.bold = True; r.font.size = Pt(10); r.font.name = "Times New Roman"
        r.font.color.rgb = RGBColor(255, 255, 255)
    for ri, row in enumerate(rows):
        fill = "EBF0F7" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = str(val); bg(cell, fill)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10); run.font.name = "Times New Roman"
    if widths:
        for row in t.rows:
            for ci, w in enumerate(widths):
                row.cells[ci].width = Cm(w)
    if cap: caption(cap)
    return t

def flow_box(text, fill="DAE8FC", bold=False, arrow=True):
    tb = doc.add_table(rows=1, cols=1); tb.style = 'Table Grid'
    c = tb.rows[0].cells[0]; c.text = text
    bg(c, fill)
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in c.paragraphs[0].runs:
        run.font.size = Pt(10); run.font.name = "Times New Roman"
        run.bold = bold
    if arrow:
        ap = doc.add_paragraph()
        ap.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.CENTER
        ap.paragraph_format.space_before = Pt(0); ap.paragraph_format.space_after = Pt(0)
        ar = ap.add_run("↓"); ar.font.size = Pt(14); ar.font.name = "Times New Roman"
        ar.font.color.rgb = RGBColor(46, 64, 87)

def hr():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr(); pb = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '1'); bot.set(qn('w:color'), '2E4057')
    pb.append(bot); pPr.append(pb)

# ══════════════════════════════════════════════════════════════════════════════
# PAGE DE TITRE DU CHAPITRE
# ══════════════════════════════════════════════════════════════════════════════
def centered(text, size, bold=False, color="1F3864", sa=8):
    p = doc.add_paragraph()
    p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(sa)
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size)
    r.font.name = "Times New Roman"; r.font.color.rgb = RGBColor.from_string(color)

centered("CHAPITRE 2", 20, bold=True, sa=6)
centered("MATÉRIELS, MÉTHODES ET INGÉNIERIE DU SYSTÈME", 15, bold=True, sa=6)
centered("SignInterpreter V11 — Système de reconnaissance continue de la langue des signes", 12, bold=False, color="2E4057", sa=4)
hr()

# ══════════════════════════════════════════════════════════════════════════════
# 2.1 INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════
h1("2.1 Introduction")
j("Ce chapitre décrit les matériels, les données et les méthodes qui constituent le système SignInterpreter V11. Il est organisé selon l'ordre canonique d'une présentation méthodologique scientifique : après une vue d'ensemble de l'architecture globale, nous présentons l'environnement matériel et logiciel, la stratégie de collecte et de traitement des données, les architectures des modèles appris, le protocole d'entraînement, et enfin la pipeline d'inférence temps réel. Un bilan de la démarche de développement itératif clôt le chapitre en replaçant chaque choix dans son contexte d'évolution.")
j("Toutes les données numériques citées dans ce chapitre sont directement extraites du code source ou des fichiers de configuration des modèles entraînés (config.json, meta.json). Aucune valeur n'est approximée ou extrapolée.")

# ══════════════════════════════════════════════════════════════════════════════
# 2.2 ARCHITECTURE GLOBALE DU SYSTÈME
# ══════════════════════════════════════════════════════════════════════════════
h1("2.2 Architecture globale du système")
j("Avant d'entrer dans le détail des composants, il est nécessaire de présenter l'architecture d'ensemble afin de situer chaque module dans le flux global de traitement. Le système SignInterpreter V11 est une application client-serveur à deux niveaux, orientée temps réel, qui transforme un flux vidéo en phrases françaises.")

# Architecture layers diagram
t_arch = doc.add_table(rows=5, cols=3); t_arch.style = 'Table Grid'
arch_rows = [
    ("CLIENT",   "COUCHE ACQUISITION",   "Caméra 720p  →  MediaPipe Holistic (JS)\nExtraction : 171 features / image à ≈30 FPS", "D5E8D4"),
    ("RÉSEAU",   "COUCHE TRANSPORT",     "WebSocket JSON bidirectionnel\nClient → Serveur : frames brutes  |  Serveur → Client : texte traduit", "DAE8FC"),
    ("SERVEUR",  "COUCHE INFÉRENCE",     "FastAPI + Uvicorn (Python 3.11)\nCTC BiLSTM  +  LSTM fenêtre glissante  +  Traducteur", "E1D5E7"),
    ("SERVEUR",  "COUCHE TRADUCTION",    "Traducteur par règles (LSF → Français)\n+ Seq2Seq attention (module optionnel)", "FFF2CC"),
    ("STOCKAGE", "COUCHE PERSISTANCE",   "Templates .npy  |  Modèles .pt / .keras  |  SQLite", "F0F0F0"),
]
lbg = ["D5E8D4", "DAE8FC", "E1D5E7", "FFF2CC", "F0F0F0"]
for ri, (side, name, desc, fill) in enumerate(arch_rows):
    c0, c1, c2 = t_arch.rows[ri].cells
    c0.text = side;  bg(c0, fill); c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    c1.text = name;  bg(c1, fill); c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    c2.text = desc;  bg(c2, fill)
    for cell, w in [(c0, 2.5), (c1, 4.5), (c2, 9.5)]:
        cell.width = Cm(w)
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(10); run.font.name = "Times New Roman"
    c1.paragraphs[0].runs[0].bold = True
caption("Figure 2.1 — Architecture en couches du système SignInterpreter V11")

j("Le découplage client-serveur présente un double avantage. D'une part, l'extraction des points anatomiques par MediaPipe s'effectue entièrement côté navigateur, ce qui réduit la bande passante à 684 octets par image (171 flottants × 4 octets) au lieu des ~230 Ko d'une image JPEG brute. D'autre part, les composants d'inférence lourds (PyTorch, TensorFlow) restent sur le serveur, seul responsable de la logique ML.")

# ══════════════════════════════════════════════════════════════════════════════
# 2.3 ENVIRONNEMENT MATÉRIEL ET LOGICIEL
# ══════════════════════════════════════════════════════════════════════════════
h1("2.3 Environnement matériel et logiciel")

h2("2.3.1 Configuration matérielle")
j("L'intégralité du développement, de la collecte de données et des entraînements a été réalisée sur un seul poste de travail, sans accès à un serveur de calcul externe. Cette contrainte a directement influencé les choix architecturaux, notamment les tailles de modèles, les formats de données et la stratégie d'entraînement.", indent=False)

tbl(
    ["Composant", "Spécification"],
    [
        ["Processeur",  "Apple M1 (4 cœurs haute performance + 4 cœurs efficacité, ARM64)"],
        ["Mémoire",     "16 Go unifiée CPU/GPU (Metal Performance Shaders)"],
        ["GPU intégré", "Apple GPU 8 cœurs (MPS — Metal Performance Shaders)"],
        ["Stockage",    "SSD NVMe (lecture séquentielle ~3 Go/s)"],
        ["Caméra",      "FaceTime HD 720p intégrée (source vidéo principale, ~30 FPS)"],
        ["Système",     "macOS Darwin 24.x"],
    ],
    cap="Tableau 2.1 — Configuration matérielle du poste de développement et d'entraînement",
    widths=[5.0, 11.5]
)

j("L'architecture Apple Silicon présente une caractéristique importante pour PyTorch : l'opération CTCLoss n'est pas supportée sur le backend MPS. La stratégie adoptée consiste à calculer la perte sur CPU, puis à rétropropager les gradients vers les paramètres du modèle qui résident sur le GPU MPS. Les passes forward bénéficient ainsi de l'accélération matérielle, tandis que la perte est calculée sur CPU.", indent=False)

h2("2.3.2 Stack logiciel")
j("Le système repose sur deux écosystèmes d'apprentissage automatique distincts, coexistant délibérément : TensorFlow/Keras pour le classificateur LSTM, dont l'API de haut niveau simplifie la gestion des callbacks et de l'augmentation, et PyTorch pour les modèles CTC et Seq2Seq, dont le graphe de calcul dynamique est mieux adapté aux séquences de longueur variable et au contrôle fin requis par CTCLoss. Côté serveur, FastAPI a été retenu pour sa compatibilité native avec les WebSockets asynchrones et sa génération automatique de documentation d'API. L'extraction des points anatomiques par MediaPipe est exécutée en JavaScript directement dans le navigateur, évitant toute transmission d'images brutes sur le réseau.", indent=False)
tbl(
    ["Composant", "Version", "Rôle dans le système"],
    [
        ["Python",       "3.11",     "Langage principal (serveur + entraînement)"],
        ["TensorFlow",   "2.16.2",   "Entraînement et inférence du classificateur LSTM (Keras)"],
        ["PyTorch",      "2.x",      "Entraînement et inférence CTC + Seq2Seq (MPS/CPU)"],
        ["MediaPipe",    "0.10.14",  "Détection des points anatomiques (JavaScript, côté client)"],
        ["NumPy",        "≥ 1.24",   "Manipulations matricielles, génération de données"],
        ["scikit-learn", "≥ 1.3",    "Stratification train/val, métriques"],
        ["FastAPI",      "≥ 0.110",  "Serveur API REST + WebSocket"],
        ["Uvicorn",      "—",        "Serveur ASGI (interface FastAPI)"],
        ["React + TS",   "Vite",     "Interface utilisateur (client navigateur)"],
        ["SQLite",       "—",        "Persistance : contributions, vocabulaire"],
    ],
    cap="Tableau 2.2 — Stack logiciel complet du système SignInterpreter V11",
    widths=[3.0, 2.5, 11.0]
)

# ══════════════════════════════════════════════════════════════════════════════
# 2.4 CORPUS ET STRATÉGIE DE COLLECTE DES DONNÉES
# ══════════════════════════════════════════════════════════════════════════════
h1("2.4 Corpus et stratégie de collecte des données")
j("Le système repose sur deux types de données collectées indépendamment : (1) des gabarits de signes isolés (séquences de features pour chaque signe du vocabulaire) et (2) des paires linguistiques associant une séquence de signes à sa traduction française. Ces deux corpus ont des rôles distincts dans la chaîne d'entraînement.")

h2("2.4.1 Définition du vocabulaire cible")
j("Le vocabulaire du système comprend 30 signes couvrant les fonctions communicatives fondamentales : salutations, pronoms personnels, verbes transitifs courants, adjectifs évaluatifs et quelques noms. Ce vocabulaire forme un ensemble cohérent permettant la construction de phrases simples à complexes dans un registre quotidien.", indent=False)

# Vocabulary box
vp = doc.add_paragraph()
vp.paragraph_format.left_indent   = Cm(1.5)
vp.paragraph_format.space_after   = Pt(2)
vr = vp.add_run(
    "aide  •  aimer  •  aller  •  ami  •  amour  •  apprendre  •  attendre  •  au revoir\n"
    "aujourd'hui  •  avec  •  bébé  •  bien  •  bientôt  •  boire  •  bon  •  bonjour\n"
    "changer  •  chat  •  comment  •  comprendre  •  donner  •  ici  •  manger  •  merci\n"
    "moi  •  non  •  nous  •  oui  •  tu  •  vouloir"
)
vr.font.size = Pt(11); vr.font.name = "Times New Roman"; vr.italic = True
caption("Encadré 2.1 — Les 30 signes du vocabulaire SignInterpreter V11")

h2("2.4.2 Acquisition des gabarits de signes isolés")
j("Pour chaque signe, des séquences de référence ont été enregistrées par un signeur unique dans des conditions contrôlées (éclairage uniforme, fond neutre, caméra frontale à hauteur des épaules). Chaque enregistrement est automatiquement converti en une séquence de vecteurs de caractéristiques (T × 171) par le pipeline MediaPipe.")
j("Un script de collecte interactif guide le processus : affichage du signe à exécuter, compte à rebours, enregistrement de la séquence, vérification visuelle immédiate. Les séquences rejetées (mains non détectées sur plus de 30 % des images, ou mouvements parasites identifiés visuellement) sont écartées avant stockage.")

add_img(IMG_CLASSES, width_cm=14.5)
caption("Figure 2.2 — Distribution des exemples enregistrés par classe (total : 2 100)\n"
        "En bleu : 70 exemples (valeur cible). En orange : classes avec légère variation due aux rejets qualité (65–75).")

h2("2.4.3 Corpus de paires linguistiques")
j("En complément des gabarits de signes isolés, un corpus de 6 209 paires (séquence de signes → phrase française) a été constitué pour entraîner les modules de traduction. Ces paires couvrent l'ensemble des constructions permises par le vocabulaire sur différents sujets grammaticaux, temps verbaux et polarités (affirmatif / négatif).")

add_img(IMG_PAIRES, width_cm=14.5)
caption("Figure 2.3 — Distribution des 6 209 paires d'entraînement par longueur de séquence de signes\n"
        "Les séquences de 4 et 5 signes représentent à elles seules 57.5 % du corpus.")

# ══════════════════════════════════════════════════════════════════════════════
# 2.5 EXTRACTION ET REPRÉSENTATION DES CARACTÉRISTIQUES
# ══════════════════════════════════════════════════════════════════════════════
h1("2.5 Extraction et représentation des caractéristiques")
j("La qualité de la représentation des données est déterminante pour les performances de reconnaissance. Le module features.py définit le vecteur de 171 caractéristiques par image qui constitue l'entrée de tous les modèles du système. Ce module a été finalisé à la version 7 du projet et est resté stable jusqu'à V11.")

h2("2.5.1 Composition du vecteur de caractéristiques")
j("Chaque image traitée par MediaPipe Holistic produit un vecteur de 171 valeurs réelles organisé en trois blocs fonctionnels distinctement normalisés :", indent=False)

# Feature composition diagram
t_feat = doc.add_table(rows=2, cols=3); t_feat.style = 'Table Grid'
feat_cols = [
    ("POSTURE CORPORELLE  [0:39]\n\n13 points anatomiques × 3 coordonnées\n= 39 dimensions\n\nPoints : nez, yeux, oreilles, épaules,\ncoudes, poignets, doigts auriculaires\n\nRéférence : midpoint épaules\nÉchelle : distance inter-épaules", "D5E8D4"),
    ("MAIN GAUCHE  [39:105]\n\nForme : 21 pts × 3 = 63 dims\nPondération HAND_WEIGHT = 3.0\n\nOrientation palmaire : 3 dims\nPondération PALM_WEIGHT = 5.0\n\nTotal : 66 dimensions\n\nRéférence : poignet (pt 0)\nÉchelle : max distance poignet-doigt", "DAE8FC"),
    ("MAIN DROITE  [105:171]\n\nForme : 21 pts × 3 = 63 dims\nPondération HAND_WEIGHT = 3.0\n\nOrientation palmaire : 3 dims\nPondération PALM_WEIGHT = 5.0\n\nTotal : 66 dimensions\n\nRéférence : poignet (pt 0)\nÉchelle : max distance poignet-doigt", "FFF2CC"),
]
feat_totals = [("39 dims", "27AE60"), ("66 dims", "2980B9"), ("66 dims", "E67E22")]
for ci, (text, fill) in enumerate(feat_cols):
    c = t_feat.rows[0].cells[ci]; c.text = text; bg(c, fill)
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in c.paragraphs[0].runs:
        run.font.size = Pt(9); run.font.name = "Times New Roman"
    c2 = t_feat.rows[1].cells[ci]
    tot, col = feat_totals[ci]
    c2.text = tot; bg(c2, col)
    c2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = c2.paragraphs[0].runs[0]
    r2.bold = True; r2.font.size = Pt(11); r2.font.name = "Times New Roman"
    r2.font.color.rgb = RGBColor(255, 255, 255)
    for row in t_feat.rows:
        row.cells[ci].width = Cm(5.3)
caption("Figure 2.4 — Composition du vecteur de 171 caractéristiques par image\n"
        "Formule : 39 (posture) + 66 (main gauche) + 66 (main droite) = 171 dims\n"
        "Extension optionnelle : + 6 dims vitesse des poignets → 177 dims (non utilisée dans V11)")

h2("2.5.2 Normalisation et invariances")
j("La robustesse aux variations indépendantes du signe lui-même est assurée par trois niveaux de normalisation :")
dash("Invariance de position et de taille corporelle",
     "les coordonnées de posture sont traduites vers le point milieu des épaules puis divisées par la distance inter-épaules. Un signeur debout ou assis, grand ou petit, produit le même vecteur pour la même configuration corporelle.")
dash("Invariance de position et de taille de la main",
     "les coordonnées de chaque main sont centrées sur le poignet (landmark 0) et normalisées par la distance maximale poignet-point le plus éloigné. La forme de la main est ainsi indépendante de sa position dans l'espace et de la distance à la caméra.")
dash("Conservation de l'orientation palmaire",
     "l'orientation de la paume est délibérément conservée via le vecteur normal au plan de la paume, calculé par produit vectoriel : n⃗ = (poignet→index) × (poignet→auriculaire). Cette grandeur est phonologiquement distinctive (paume vers la caméra vs paume vers le signeur distinguent des signes différents).")

j("De plus, un mécanisme de canonicalisation de latéralité (to_right_dominant) assure que la main dominante est toujours représentée dans le bloc « main droite » du vecteur. Pour les signes bilatéraux (deux mains actives sur plus de 50 % des images), la latéralité est préservée afin de ne pas altérer l'information spatiale intrinsèque du signe.")

# ══════════════════════════════════════════════════════════════════════════════
# 2.6 AUGMENTATION ET GÉNÉRATION SYNTHÉTIQUE DES DONNÉES
# ══════════════════════════════════════════════════════════════════════════════
h1("2.6 Augmentation et génération synthétique des données")
j("Le corpus brut de 2 100 gabarits est insuffisant pour entraîner des modèles profonds robustes. Deux stratégies d'enrichissement sont appliquées selon le modèle cible : l'augmentation stochastique pour le LSTM et la génération synthétique pour le CTC.")

h2("2.6.1 Augmentation des données pour le classificateur LSTM")
j("Un multiplicateur d'augmentation ×10 transforme les 2 100 gabarits en 21 000 séquences d'entraînement. Cinq techniques sont appliquées de manière indépendante et stochastique à chaque séquence :", indent=False)

add_img(IMG_AUGMENT, width_cm=14.0)
caption("Figure 2.5 — Probabilités d'application des cinq techniques d'augmentation (corpus LSTM)\n"
        "Chaque technique est appliquée indépendamment. Un exemple peut recevoir 0 à 5 transformations simultanées.\n"
        "Les paramètres sont indiqués dans le code source ml/lstm_trainer.py.")

h2("2.6.2 Génération synthétique pour l'entraînement CTC")
j("L'entraînement d'un modèle CTC requiert des séquences multi-signes continues avec leurs étiquettes (liste de signes) sans aucun alignement temporel préalable. De telles séquences seraient prohibitivement longues à produire manuellement (annotation frame-par-frame). Un générateur synthétique produit 46 000 séquences à partir des exemples enregistrés :")
dash("6 000 séquences mono-signe",
     "(200 par classe × 30 classes). Ces séquences corrigent un biais de position : sans elles, le modèle tendrait à n'émettre des signes qu'après plusieurs images de contexte accumulées.")
dash("40 000 séquences multi-signes",
     "construites par concaténation de signes tirés des 6 209 paires linguistiques réelles, avec sur-représentation des signes rares (échantillonnage pondéré inversement proportionnel à la fréquence de la classe la moins présente dans chaque paire).")
j("Chaque séquence synthétique subit quatre transformations :", indent=False)
dash("Variation de tempo", "facteur aléatoire ∈ [0.50 ; 1.60] appliqué par rééchantillonnage linéaire, simulant des signeurs rapides et lents.")
dash("Bruit gaussien", "σ = 0.008 par dimension et par image.")
dash("Trames de transition", "3 images interpolées linéairement entre deux signes consécutifs pour simuler la fluidité naturelle des enchaînements.")
dash("Stockage float16", "pour contenir l'empreinte mémoire à 1,3 Go pour l'ensemble des 46 000 séquences (fichier ctc_sequences.pkl).")

add_img(IMG_VOLUMES, width_cm=14.0)
caption("Figure 2.6 — Volumes de données produits à chaque étape de préparation\n"
        "Les 2 100 exemples bruts sont amplifiés par ×10 (LSTM) ou ×22 (CTC synthétique).\n"
        "Les paires linguistiques (6 209) constituent un corpus indépendant pour l'entraînement Seq2Seq.")

# ══════════════════════════════════════════════════════════════════════════════
# 2.7 ARCHITECTURES DES MODÈLES APPRIS
# ══════════════════════════════════════════════════════════════════════════════
h1("2.7 Architectures des modèles appris")
j("Le système intègre trois modèles appris, chacun remplissant un rôle distinct et complémentaire dans la chaîne de traitement.")

h2("2.7.1 Classificateur LSTM pour la reconnaissance de signes isolés")
j("Le classificateur LSTM reconnaît un signe parmi 30 à partir d'une séquence rééchantillonnée à longueur fixe de 30 images × 171 dimensions. Il constitue le composant secondaire du système en mode continu (fenêtre glissante) et le composant principal en mode isolé.", indent=False)

t_lstm = doc.add_table(rows=5, cols=1); t_lstm.style = 'Table Grid'
lstm_l = [
    ("Entrée  ×  [30 images × 171 dims]  — rééchantillonné à longueur fixe", "C8E6C9"),
    ("BiLSTM  —  64 unités cachées → 128 (bidirectionnel)\n+ Dropout(0.50 → 0.40 selon l'époque)  +  L2 régularisation (λ=1e-4)", "B3D9F7"),
    ("Mécanisme d'attention temporelle\nDense(128) → tanh → score → softmax → somme pondérée des états cachés", "D7BDE2"),
    ("Dense(30 classes, softmax)  +  Dropout final", "FAD7A0"),
    ("Sortie  ×  probabilité sur 30 signes  →  argmax = signe prédit", "C8E6C9"),
]
for ri, (text, fill) in enumerate(lstm_l):
    c = t_lstm.rows[ri].cells[0]; c.text = text; bg(c, fill)
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in c.paragraphs[0].runs:
        run.font.size = Pt(10); run.font.name = "Times New Roman"
        if ri == 0 or ri == 4: run.bold = True
    if ri < 4:
        ap = doc.add_paragraph(); ap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ap.paragraph_format.space_before = Pt(0); ap.paragraph_format.space_after = Pt(0)
        ar = ap.add_run("↓"); ar.font.size = Pt(12); ar.font.name = "Times New Roman"
        ar.font.color.rgb = RGBColor(46, 64, 87)
caption("Figure 2.7 — Architecture du classificateur LSTM isolé (30 signes, entrée 30×171)")

h2("2.7.2 Modèle CTC BiLSTM pour la reconnaissance continue")
j("Le modèle CTC est le composant central de la reconnaissance continue. Il accepte une séquence de longueur variable T (toutes les images accumulées depuis le début d'une phrase) et prédit une séquence de signes sans alignement temporel préalable. La perte CTC (Connectionist Temporal Classification, Graves et al. 2006) marginalise sur tous les alignements possibles entre la séquence d'entrée et la séquence cible.", indent=False)

t_ctc = doc.add_table(rows=5, cols=1); t_ctc.style = 'Table Grid'
ctc_l = [
    ("Entrée  ×  [T images × 171 dims]  — longueur variable, séquence entière de la phrase", "C8E6C9"),
    ("BiLSTM couche 1  —  256 unités cachées → 512 (bidirectionnel)\nDropout(0.3) entre couches", "B3D9F7"),
    ("BiLSTM couche 2  —  256 unités cachées → 512 (bidirectionnel)\nDropout(0.3) en sortie", "9BC4EB"),
    ("Projection linéaire  512 → 31 logits\n30 classes + 1 token blank (index 30, convention standard CTC)", "D7BDE2"),
    ("Décodage glouton CTC  :  argmax(logits) → suppression doublons consécutifs → suppression blanks\nSortie  ×  séquence de signes  [signe₁, signe₂, …]", "C8E6C9"),
]
for ri, (text, fill) in enumerate(ctc_l):
    c = t_ctc.rows[ri].cells[0]; c.text = text; bg(c, fill)
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in c.paragraphs[0].runs:
        run.font.size = Pt(10); run.font.name = "Times New Roman"
        if ri == 0 or ri == 4: run.bold = True
    if ri < 4:
        ap = doc.add_paragraph(); ap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ap.paragraph_format.space_before = Pt(0); ap.paragraph_format.space_after = Pt(0)
        ar = ap.add_run("↓"); ar.font.size = Pt(12); ar.font.name = "Times New Roman"
        ar.font.color.rgb = RGBColor(46, 64, 87)
caption("Figure 2.8 — Architecture du modèle CTC BiLSTM\n"
        "(entrée : séquence variable T×171  |  sortie : séquence de signes sans segmentation préalable)")

h2("2.7.3 Modèle Seq2Seq avec attention de Bahdanau pour la traduction")
j("Le module de traduction apprise implémente une architecture encodeur-décodeur avec mécanisme d'attention de Bahdanau (1) pour transformer une séquence de signes en une phrase française. L'encodeur produit des représentations contextuelles de chaque signe ; le décodeur génère les mots français un par un en consultant ces représentations via l'attention.", indent=False)

t_s2s = doc.add_table(rows=3, cols=2); t_s2s.style = 'Table Grid'
s2s_top = [
    ("ENCODEUR\n\nEmbedding  :  40 tokens → 64 dims\n↓\nBiLSTM 2 couches (hidden = 128)\n↓\nProjection  256 → 128\n(état initial décodeur)\n\nSortie : (B, T_src, 256) context vectors", "B3D9F7"),
    ("DÉCODEUR\n\nEmbedding  :  195 tokens → 64 dims\n+ Attention Bahdanau\n  (Linear(256+128) → tanh → score → softmax)\n↓\nLSTM 2 couches (hidden = 128)\n  + context vector (256 dims) concaténé\n↓\nLinéaire  (128 + 256 + 64) → 195 logits", "D7BDE2"),
]
for ci, (text, fill) in enumerate(s2s_top):
    c = t_s2s.rows[0].cells[ci]; c.text = text; bg(c, fill)
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.width = Cm(8.0)
    for run in c.paragraphs[0].runs:
        run.font.size = Pt(9); run.font.name = "Times New Roman"
cl = t_s2s.rows[1].cells[0]; cl.text = "Entrée : séquence de signes  [s₁, s₂, …, sₙ]"
bg(cl, "C8E6C9"); cl.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
cr = t_s2s.rows[1].cells[1]; cr.text = "Sortie : phrase française  [mot₁, mot₂, …, <EOS>]  max 20 mots"
bg(cr, "C8E6C9"); cr.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
cl2 = t_s2s.rows[2].cells[0]
cl2.text = "Vocabulaire source : 40 tokens (30 signes + 4 tokens spéciaux + 6 bigrams)"
bg(cl2, "F0F0F0"); cl2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
cr2 = t_s2s.rows[2].cells[1]
cr2.text = "Vocabulaire cible : 195 tokens (mots français du corpus d'entraînement)"
bg(cr2, "F0F0F0"); cr2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
for row in t_s2s.rows:
    for ci in range(2):
        for run in row.cells[ci].paragraphs[0].runs:
            run.font.size = Pt(9); run.font.name = "Times New Roman"
caption("Figure 2.9 — Architecture du modèle Seq2Seq avec attention de Bahdanau\n"
        "(src_vocab = 40, tgt_vocab = 195, emb_dim = 64, hidden = 128, 2 couches LSTM)")

# ══════════════════════════════════════════════════════════════════════════════
# 2.8 PROTOCOLE D'ENTRAÎNEMENT
# ══════════════════════════════════════════════════════════════════════════════
h1("2.8 Protocole d'entraînement")

h2("2.8.1 Entraînement du classificateur LSTM")
j("Le classificateur LSTM est entraîné avec TensorFlow/Keras sur les 21 000 séquences augmentées, avec une division stratifiée train/validation de 85 %/15 % (17 850 / 3 150 séquences). La taille de batch est calculée dynamiquement : min(32, max(8, n_entraînement // 8)).", indent=False)
j("Trois callbacks assurent la régularisation du processus :", indent=False)
dash("ModelCheckpoint", "sauvegarde automatique du meilleur modèle selon la métrique val_accuracy.")
dash("EarlyStopping (patience = 25)", "arrêt anticipé si aucune amélioration sur 25 époques consécutives, sur 200 époques maximum.")
dash("ReduceLROnPlateau (facteur = 0.5, patience = 12, lr_min = 1e-5)", "réduction du taux d'apprentissage en cas de plateau de val_accuracy.")

h2("2.8.2 Entraînement du modèle CTC BiLSTM")
j("L'entraînement CTC présente deux difficultés spécifiques par rapport à un entraînement supervisé classique. Premièrement, la perte CTCLoss nécessite que la longueur T de la séquence d'entrée soit strictement supérieure à la longueur de la séquence cible (nombre de signes). Pour les séquences synthétiques avec tempo rapide (facteur 0.5), cette contrainte peut être violée. Le paramètre zero_infinity=True de CTCLoss gère ces cas en ignorant les pertes infinies.", indent=False)
j("Deuxièmement, le calcul de la perte CTC n'étant pas supporté sur le backend MPS d'Apple Silicon, la stratégie utilisée est de déplacer les logits sur CPU pour le calcul de la perte, puis de rétropropager les gradients vers les paramètres du modèle qui restent sur MPS.", indent=False)

tbl(
    ["Hyperparamètre", "Valeur", "Justification"],
    [
        ["Architecture",       "BiLSTM 2 couches (hidden = 256)",     "Modélisation des dépendances bidirectionnelles longue portée"],
        ["Dimension d'entrée", "171 dims",                            "Vecteur complet sans enrichissement vitesse"],
        ["Token blank CTC",    "Index 30 (= n_classes)",              "Convention standard PyTorch CTCLoss"],
        ["Batch size",         "32",                                  "Contrainte mémoire M1 16 Go (partagée CPU/GPU)"],
        ["Optimiseur",         "Adam (lr=1e-3, weight_decay=1e-5)",   "Convergence rapide + régularisation L2 légère"],
        ["Scheduler LR",       "CosineAnnealingLR (T_max=80, η_min=1e-5)", "Exploration large puis affinage fin de courbe"],
        ["Gradient clipping",  "norm = 5.0",                         "Prévention de l'explosion du gradient (LSTM profond)"],
        ["Époques max",        "80 (early stop patience = 15)",       "Convergence typique observée avant 80 époques"],
        ["Division train/val", "90 % / 10 %  (41 400 / 4 600)",      "Plus large que LSTM car séquences synthétiques"],
    ],
    cap="Tableau 2.6 — Hyperparamètres d'entraînement du modèle CTC BiLSTM",
    widths=[4.0, 5.0, 7.5]
)

h2("2.8.3 Entraînement du modèle Seq2Seq")
j("L'entraînement Seq2Seq utilise la technique du teacher forcing annealing pour éviter le biais d'exposition (exposure bias) : en mode pure teacher forcing, le modèle apprend à générer chaque mot en connaissant le mot précédent correct, mais en inférence il doit s'appuyer sur ses propres prédictions. La formule d'anneal est :", indent=False)

code_block([
    "tf = max(0.3,  1.0 − époque × 0.008)",
    "# tf = 1.0 en époque 1  →  tf = 0.3 à partir de l'époque 88",
    "# Avec probabilité tf : entrée = vrai mot précédent (teacher forcing)",
    "# Avec probabilité 1-tf : entrée = prédiction du modèle (scheduled sampling)",
])

j("Le décodage en inférence est glouton (greedy) avec une longueur maximale de 20 mots. La division train/validation est 90 %/10 % (5 588 / 621 paires).", indent=False)

h2("2.8.4 Résultats des entraînements")
tbl(
    ["Modèle", "Données entraînement", "Données validation", "Meilleure métrique validée"],
    [
        ["Classificateur LSTM",
         "17 850 séquences\n(85 % × 21 000 augmentées)",
         "3 150 séquences\n(15 % × 21 000)",
         "Précision validation : 100.0 %\nPrécision gabarits originaux : 100.0 %\n(données : data/lstm/meta.json)"],
        ["CTC BiLSTM",
         "41 400 séquences\n(90 % × 46 000 synthétiques)",
         "4 600 séquences\n(10 % × 46 000)",
         "Précision signe : 99.97 %\n(sign accuracy = 1 − taux d'erreur Levenshtein)\n(données : data/ctc_model/config.json)"],
        ["Seq2Seq attention",
         "5 588 paires\n(90 % × 6 209 paires linguistiques)",
         "621 paires\n(10 % × 6 209)",
         "Perte de validation : 0.0354\n(cross-entropie ignorant le padding)\n(données : data/seq2seq/config.json)"],
    ],
    cap="Tableau 2.7 — Résultats d'entraînement des trois modèles appris. Valeurs extraites directement des fichiers config.json.",
    widths=[3.5, 4.0, 3.5, 5.5]
)

j("La précision de 100 % du classificateur LSTM sur la validation s'explique par la combinaison de la normalisation rigoureuse (invariances de position et de taille) et de l'équilibre parfait des classes. La précision CTC de 99.97 % sur les séquences synthétiques confirme que le modèle apprend l'alignement temporel sans annotation. La perte Seq2Seq de 0.0354 correspond à des traductions grammaticalement correctes sur les paires du vocabulaire couvert, vérifiées par inspection manuelle.")

# ══════════════════════════════════════════════════════════════════════════════
# 2.9 PIPELINE D'INFÉRENCE TEMPS RÉEL
# ══════════════════════════════════════════════════════════════════════════════
h1("2.9 Pipeline d'inférence temps réel")
j("La pipeline d'inférence transforme un flux vidéo continu en une phrase française affichée dans l'interface utilisateur. Elle s'exécute en quasi-temps réel, organisée en sept étapes séquentielles.")

# Main pipeline flowchart — 7 steps
pipe_steps = [
    ("ÉTAPE 1  —  ACQUISITION VIDÉO\nCaméra FaceTime HD 720p  |  Cadence ~30 FPS", "C8E6C9", True),
    ("ÉTAPE 2  —  EXTRACTION DES CARACTÉRISTIQUES (CLIENT)\nMediaPipe Holistic (JavaScript)\n21 pts/main + 13 pts posture → vecteur 171 dims par image", "B3D9F7", True),
    ("ÉTAPE 3  —  TRANSMISSION WEBSOCKET\n{frames: [[171 floats], …], n_frames: N}\nBande passante : ~20 Ko / fenêtre de 30 images (vs ~7 Mo en JPEG brut)", "D7BDE2", True),
    ("ÉTAPE 4a  —  RECONNAISSANCE CTC (composant principal)\nAccumulation de toutes les frames depuis le début de la phrase\nRédécodage toutes les 20 images (CTC_STRIDE) | Min. 30 images avant 1er décodage\nPolitique extend-only : signes confirmés non supprimés rétrospectivement", "FAD7A0", True),
    ("ÉTAPE 4b  —  FENÊTRE GLISSANTE LSTM (composant secondaire — réactivité UI)\nFenêtre de 30 images | Pas de 3 images | Seuil de confiance -log(p) < 0.55\nRetour visuel immédiat ~100 ms | N'influence pas la décision finale", "FDEBD0", True),
    ("ÉTAPE 5  —  DÉTECTION DES FRONTIÈRES DE PHRASE\nRepos poignet (Y ≥ 0.8) pendant ≥ 6 frames consécutives AND énergie < 0.02\nOU mains absentes > 25 frames  |  OU timeout inactivité > 3 s", "E8DAEF", True),
    ("ÉTAPE 6  —  TRADUCTION LSF → FRANÇAIS\nTraducteur par règles (conjugaison, accord, négation) — toujours actif\n+ Seq2Seq attention (décodage glouton, max 20 mots) — module optionnel", "D5F5E3", True),
    ("ÉTAPE 7  —  AFFICHAGE CLIENT\nPhrase française affichée en temps réel dans l'interface React", "C8E6C9", False),
]
for text, fill, arrow in pipe_steps:
    flow_box(text, fill=fill, arrow=arrow)
caption("Figure 2.10 — Pipeline complète d'inférence temps réel SignInterpreter V11 (7 étapes)\n"
        "Étapes 4a et 4b s'exécutent en parallèle sur chaque réception de frames")

h2("2.9.1 Paramètres de contrôle de la pipeline")
tbl(
    ["Paramètre", "Valeur", "Rôle"],
    [
        ["SEQUENCE_LENGTH",     "30 images",  "Longueur fenêtre LSTM (≈ 1 s à 30 FPS)"],
        ["CTC_STRIDE",          "20 images",  "Fréquence de rédécodage CTC (≈ 667 ms)"],
        ["CTC_MIN_FRAMES",      "30 images",  "Seuil minimum avant premier décodage CTC"],
        ["THRESHOLD",           "0.55",       "Seuil de confiance LSTM : prob ≥ e^(−0.55) ≈ 0.58"],
        ["MOTION_START_THRESHOLD", "0.025",   "Delta moyen mains pour déclencher la détection"],
        ["MOTION_END_THRESHOLD",   "0.018",   "Delta moyen mains pour fermer un signe"],
        ["MOTION_END_FRAMES",   "6 images",   "Durée de repos en fin de signe (~200 ms)"],
        ["MIN_SIGN_FRAMES",     "6 images",   "Durée minimale d'un signe valide (~200 ms)"],
        ["MAX_SIGN_FRAMES",     "120 images", "Durée maximale d'un signe (~4 s)"],
        ["PREDICT_AFTER_N_FRAMES","5 images", "Premier retour visuel après N images"],
        ["PREDICT_EVERY_N_FRAMES","3 images", "Fréquence du retour visuel intermédiaire"],
        ["REST_FRAMES",         "6 images",   "Images repos consécutives pour fin de phrase"],
        ["HAND_TIMEOUT",        "25 images",  "Images sans mains avant fin de phrase (~833 ms)"],
        ["IDLE_TIMEOUT",        "3 s",        "Timeout d'inactivité complète"],
    ],
    cap="Tableau 2.8 — Paramètres de contrôle de la pipeline temps réel\n(fichiers ml/constants.py et server/routers/inference_ws.py)",
    widths=[5.0, 2.8, 8.7]
)

# ══════════════════════════════════════════════════════════════════════════════
# 2.10 DÉMARCHE DE DÉVELOPPEMENT ITÉRATIF
# ══════════════════════════════════════════════════════════════════════════════
h1("2.10 Démarche de développement itératif")
j("Le système V11 est l'aboutissement d'onze versions successives développées selon une démarche agile informelle. Chaque itération a été initiée par une limitation concrète observée lors de tests, et a apporté une réponse architecturale précise. Cette section décrit les grandes ruptures architecturales et les motivations qui les ont engendrées.")

tbl(
    ["Itération", "Versions", "Architecture", "Limitation observée → Action corrective"],
    [
        ["1",
         "V1–V3",
         "Classificateur LSTM\nRécognition signe par signe\nentrée (30×171), sortie 1 classe",
         "L'utilisateur devait marquer début/fin de chaque signe.\nLimite : inutilisable en flux naturel.\n→ Introduction d'un mécanisme de segmentation automatique."],
        ["2",
         "V4–V6",
         "Fenêtre glissante\nDétection continue par seuil\n(THRESHOLD sur -log(prob))",
         "Faux positifs aux transitions ; doublons pour les signes lents.\nAbsence de contrainte linguistique sur la séquence.\n→ Refonte des caractéristiques (V7), puis modèle séquentiel global (V11)."],
        ["3",
         "V7",
         "Refonte de l'extraction\n171 dims, orientation palmaire\npondérations différenciées",
         "Signes phonologiquement proches confondus (même mouvement,\norientations palmaires différentes).\n→ Ajout canal palmaire + pondérations HAND_WEIGHT / PALM_WEIGHT."],
        ["4",
         "V8",
         "Approche ensembliste\n(RF phonologique + RF brut)\ncombinés par score",
         "Architecture complexe, pas de contexte séquentiel global.\n→ Adoption du CTC, qui modélise explicitement la séquence entière."],
        ["5",
         "V11",
         "CTC BiLSTM\n+ fenêtre glissante hybride\n+ Seq2Seq traduction",
         "Version finale. Résout l'alignement automatiquement via CTC.\nHybride CTC/LSTM pour réactivité UI."],
    ],
    cap="Tableau 2.9 — Cinq itérations architecturales majeures : approche, limitation et action corrective",
    widths=[1.8, 1.8, 4.5, 8.4]
)

j("La progression des itérations illustre un principe d'ingénierie fondamental : chaque composant du système final (normalisation V7, CTC V11, traducteur par règles) a été introduit en réponse à une défaillance mesurable du système précédent, et non par anticipation théorique. Les règles de traduction, par exemple, ont été préférées à un modèle seq2seq pur précisément parce que celui-ci présentait des hallucinations sur les séquences hors distribution du vocabulaire couvert.")

# ══════════════════════════════════════════════════════════════════════════════
# 2.11 RÉCAPITULATIF DES PARAMÈTRES CLÉS
# ══════════════════════════════════════════════════════════════════════════════
h1("2.11 Récapitulatif des paramètres clés du système")
tbl(
    ["Dimension", "Valeur", "Source"],
    [
        ["Dimension du vecteur de caractéristiques", "171 dims (39+66+66)", "ml/features.py — FRAME_FEATURE_DIM"],
        ["Vocabulaire signes", "30 classes", "data/ctc_model/classes.json"],
        ["Gabarits collectés", "2 100 (≈70 / classe)", "data/templates/ — vérifiés"],
        ["Paires linguistiques", "6 209 paires", "data/training_pairs.csv"],
        ["Séquences CTC synthétiques", "46 000 (1,3 Go)", "data/ctc_sequences.pkl"],
        ["Séquences LSTM augmentées", "21 000 (×10)", "data/lstm/meta.json"],
        ["Architecture CTC", "2-couches BiLSTM, hidden=256", "data/ctc_model/config.json"],
        ["Architecture LSTM", "BiLSTM 64 + attention", "ml/lstm_engine.py"],
        ["Architecture Seq2Seq", "enc+dec LSTM 2 couches, hidden=128, emb=64", "data/seq2seq/config.json"],
        ["Précision LSTM (val)", "100.0 %", "data/lstm/meta.json"],
        ["Précision CTC (val, sign acc)", "99.97 %", "data/ctc_model/config.json"],
        ["Perte Seq2Seq (val)", "0.0354 (cross-entropie)", "data/seq2seq/config.json"],
    ],
    cap="Tableau 2.10 — Récapitulatif des paramètres clés du système V11 avec références aux fichiers sources",
    widths=[6.0, 4.0, 6.5]
)

# ══════════════════════════════════════════════════════════════════════════════
# 2.12 CONCLUSION DU CHAPITRE
# ══════════════════════════════════════════════════════════════════════════════
h1("2.12 Conclusion du chapitre")
j("Ce chapitre a présenté l'intégralité de la méthodologie d'ingénierie du système SignInterpreter V11, depuis la définition du vocabulaire cible jusqu'à la pipeline d'inférence temps réel, en passant par les stratégies de collecte de données, les architectures des modèles appris, et les protocoles d'entraînement.")
j("Les contributions méthodologiques majeures peuvent être résumées en quatre points. Premièrement, la décomposition du vecteur de caractéristiques en trois canaux distincts (posture, forme de main, orientation palmaire) avec pondérations différenciées assure une représentation phonologiquement pertinente. Deuxièmement, la génération synthétique de 46 000 séquences CTC avec variation de tempo et trames de transition évite l'annotation manuelle coûteuse tout en produisant des données d'entraînement réalistes. Troisièmement, l'architecture hybride CTC/fenêtre glissante combine la robustesse linguistique du décodage CTC avec la réactivité de l'interface. Quatrièmement, le traducteur par règles comme composant primaire de traduction garantit un comportement prédictible et débuggable sur le vocabulaire couvert.")
j("Les résultats d'entraînement (précision CTC de 99.97 %, précision LSTM de 100 %, perte Seq2Seq de 0.0354) constituent des indicateurs de convergence des modèles sur les données d'entraînement. Leur interprétation en termes de performances système réelles fait l'objet du chapitre 3.")

hr()

# RÉFÉRENCES
h2("Références du chapitre 2", sb=14, sa=6)
refs = [
    "Graves, A., Fernández, S., Gomez, F., & Schmidhuber, J. (2006). Connectionist Temporal Classification: Labelling unsegmented sequence data with recurrent neural networks. Proceedings of the 23rd International Conference on Machine Learning (ICML 2006), pp. 369–376.",
    "Bahdanau, D., Cho, K., & Bengio, Y. (2015). Neural Machine Translation by Jointly Learning to Align and Translate. International Conference on Learning Representations (ICLR 2015). arXiv:1409.0473.",
    "Luong, M.-T., Pham, H., & Manning, C. D. (2015). Effective Approaches to Attention-based Neural Machine Translation. Proceedings of EMNLP 2015. arXiv:1508.04025.",
    "MediaPipe Authors. (2023). MediaPipe Holistic Solution. Google LLC. mediapipe.dev/solutions/holistic.",
    "PyTorch Contributors. (2024). torch.nn.CTCLoss. PyTorch Documentation. pytorch.org/docs/stable/generated/torch.nn.CTCLoss.html.",
    "Abadi, M., et al. (2016). TensorFlow: A System for Large-Scale Machine Learning. 12th USENIX Symposium on Operating Systems Design and Implementation (OSDI 2016). (Utilisé via TensorFlow 2.16.2).",
    "Shorten, C., & Khoshgoftaar, T. M. (2019). A Survey on Image Data Augmentation for Deep Learning. Journal of Big Data, 6(1), 1–48. (Principes d'augmentation appliqués aux séquences temporelles.)",
    "Williams, R. J., & Zipser, D. (1989). A Learning Algorithm for Continually Running Fully Recurrent Neural Networks. Neural Computation, 1(2), 270–280. (Fondements BPTT pour l'entraînement LSTM.)",
]
for i, ref in enumerate(refs, 1):
    rp = doc.add_paragraph()
    rp.paragraph_format.left_indent       = Cm(1.0)
    rp.paragraph_format.first_line_indent = Cm(-1.0)
    rp.paragraph_format.space_after       = Pt(4)
    rr = rp.add_run(f"[{i}]  {ref}")
    rr.font.size = Pt(10); rr.font.name = "Times New Roman"

# ══════════════════════════════════════════════════════════════════════════════
doc.save(OUT)
size_kb = __import__('os').path.getsize(OUT) / 1024
print(f"Généré : {OUT}")
print(f"Taille : {size_kb:.1f} Ko")
