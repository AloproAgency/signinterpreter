#!/usr/bin/env python3
"""
Génère section_pipeline.docx — Tableau corrigé des paramètres de la pipeline
d'inférence temps réel, basé exclusivement sur inference_ws.py (mode CTC principal).
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(OUT_DIR, 'section_pipeline.docx')

# ── Helpers ────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, color='BFBFBF'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def cell_para(cell, text, bold=False, italic=False, size=10,
              align=WD_ALIGN_PARAGRAPH.LEFT, color=None, font='Calibri'):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return p


def add_paragraph(doc, text='', bold=False, italic=False, size=11,
                  align=WD_ALIGN_PARAGRAPH.LEFT, space_before=4, space_after=6,
                  font='Calibri'):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = font
        run.font.size = Pt(size)
    return p


def add_heading(doc, text, level=1, size=14):
    p = doc.add_paragraph()
    p.style = f'Heading {level}'
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.font.bold = True
    return p


def add_code_inline(paragraph, text):
    run = paragraph.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return run


# ── Document ───────────────────────────────────────────────────────────────────

doc = Document()

# Marges
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── Titre ─────────────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after  = Pt(16)
r = title.add_run("Pipeline d'inférence temps réel — Paramètres du mode CTC")
r.bold = True
r.font.name = 'Calibri'
r.font.size = Pt(16)
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_before = Pt(0)
subtitle.paragraph_format.space_after  = Pt(20)
r2 = subtitle.add_run("Valeurs extraites de inference_ws.py — Mode CTC principal (V11)")
r2.italic = True
r2.font.name = 'Calibri'
r2.font.size = Pt(10)
r2.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

# ── Section 1 : Introduction ───────────────────────────────────────────────────
add_heading(doc, "8.  Pipeline d'inférence temps réel", level=1, size=13)

add_paragraph(doc,
    "La pipeline d'inférence reçoit, à chaque trame vidéo, un vecteur de "
    "171 caractéristiques construit par MediaPipe côté navigateur. Le serveur "
    "Python (FastAPI + WebSocket) accumule ces vecteurs dans un tampon flottant "
    "et les soumet au modèle BiLSTM-CTC pour décodage progressif. Une couche de "
    "détection de repos identifie les limites de phrase ; lorsqu'une limite est "
    "franchie, la séquence de signes accumulée est transmise au module de "
    "traduction seq2seq pour produire la phrase française finale.",
    size=11, space_before=4, space_after=8)

add_paragraph(doc,
    "La pipeline est gouvernée par deux familles de paramètres : les paramètres "
    "CTC qui contrôlent le décodage séquentiel, et les paramètres de segmentation "
    "qui délimitent les phrases. Ces valeurs ont été calibrées empiriquement sur "
    "des sessions de test en conditions réelles (webcam 30 fps, locuteur unique).",
    size=11, space_before=0, space_after=12)

# ── Section 2 : Paramètres CTC ────────────────────────────────────────────────
add_heading(doc, "8.1  Paramètres de décodage CTC", level=2, size=12)

add_paragraph(doc,
    "Le décodage CTC s'effectue de manière progressive : le tampon de trames est "
    "ré-analysé toutes les CTC_STRIDE trames pour mettre à jour la liste des signes "
    "détectés sans attendre la fin de la phrase. Un seuil minimal CTC_MIN_FRAMES "
    "évite les décodages sur des séquences trop courtes, peu fiables.",
    size=11, space_before=4, space_after=10)

# Tableau CTC parameters
HEADER_BG   = '1F497D'
SUBHDR_BG   = 'D6E4F0'
ROW_ALT_BG  = 'F2F7FC'
ROW_BG      = 'FFFFFF'

ctc_cols = [('Paramètre', 3.2), ('Valeur', 2.0), ('Unité', 2.0), ('Rôle', 7.8)]
tbl_ctc = doc.add_table(rows=1, cols=len(ctc_cols))
tbl_ctc.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header
hdr = tbl_ctc.rows[0].cells
for i, (name, w) in enumerate(ctc_cols):
    hdr[i].width = Cm(w)
    set_cell_bg(hdr[i], HEADER_BG)
    set_cell_borders(hdr[i], '1F497D')
    cell_para(hdr[i], name, bold=True, size=10, color='FFFFFF',
              align=WD_ALIGN_PARAGRAPH.CENTER)

ctc_rows = [
    ('CTC_MIN_FRAMES', '30', 'trames', 'Taille minimale du tampon avant le premier décodage (~1 s à 30 fps)'),
    ('CTC_STRIDE', '20', 'trames', 'Intervalle entre deux décodages progressifs (~0,67 s à 30 fps)'),
]

for idx, (param, val, unit, role) in enumerate(ctc_rows):
    row = tbl_ctc.add_row().cells
    bg = ROW_ALT_BG if idx % 2 == 0 else ROW_BG
    widths = [c[1] for c in ctc_cols]
    for j, cell in enumerate(row):
        cell.width = Cm(widths[j])
        set_cell_bg(cell, bg)
        set_cell_borders(cell, 'BFBFBF')

    cell_para(row[0], param, bold=True, size=10, font='Courier New', color='1F497D')
    cell_para(row[1], val,   bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_para(row[2], unit,  italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_para(row[3], role,  size=10)

cap1 = add_paragraph(doc, '', space_before=4, space_after=14)
cap1.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_cap = cap1.add_run("Tableau 2.8a — Paramètres de décodage CTC (")
r_cap.italic = True; r_cap.font.size = Pt(9); r_cap.font.name = 'Calibri'
add_code_inline(cap1, 'inference_ws.py')
r_cap2 = cap1.add_run(", lignes 40–41)")
r_cap2.italic = True; r_cap2.font.size = Pt(9); r_cap2.font.name = 'Calibri'

# ── Section 3 : Paramètres de segmentation ────────────────────────────────────
add_heading(doc, "8.2  Paramètres de segmentation de phrase", level=2, size=12)

add_paragraph(doc,
    "La segmentation de phrase repose sur deux mécanismes complémentaires : "
    "la détection de repos (mains basses, immobiles) et la disparition des mains "
    "hors du champ. Un délai de grâce évite les coupures prématurées dues aux "
    "occultations momentanées. Un timeout global termine automatiquement la "
    "session après une période d'inactivité prolongée.",
    size=11, space_before=4, space_after=10)

seg_cols = [('Paramètre', 3.8), ('Valeur', 2.0), ('Unité', 2.2), ('Description', 7.0)]
tbl_seg = doc.add_table(rows=1, cols=len(seg_cols))
tbl_seg.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header
hdr2 = tbl_seg.rows[0].cells
for i, (name, w) in enumerate(seg_cols):
    hdr2[i].width = Cm(w)
    set_cell_bg(hdr2[i], HEADER_BG)
    set_cell_borders(hdr2[i], '1F497D')
    cell_para(hdr2[i], name, bold=True, size=10, color='FFFFFF',
              align=WD_ALIGN_PARAGRAPH.CENTER)

# Sous-en-tête : Détection de repos
def add_subheader(table, label, ncols, bg=SUBHDR_BG):
    row = table.add_row().cells
    # Merge all cells
    merged = row[0].merge(row[ncols - 1])
    set_cell_bg(merged, bg)
    set_cell_borders(merged, '1F497D')
    cell_para(merged, label, bold=True, size=10, color='1F497D',
              align=WD_ALIGN_PARAGRAPH.LEFT)

add_subheader(tbl_seg, '  Détection de la position de repos', len(seg_cols))

rest_rows = [
    ('REST_WRIST_Y', '0.8', 'norm.', 'Seuil de coordonnée Y normalisée du poignet (0 = haut, 1 = bas) ; au-delà, la main est considérée basse'),
    ('REST_MOTION_MAX', '0.02', 'norm./tr', "Énergie de mouvement maximale tolérée pour valider l'état de repos"),
    ('REST_FRAMES', '8', 'trames', "Nombre de trames consécutives de repos requises pour déclencher la finalisation (~267 ms à 30 fps)"),
    ('REST_MIN_GAP', '0.3', 's', 'Délai minimal depuis le dernier signe reconnu avant autorisation de finalisation'),
]

for idx, (param, val, unit, desc) in enumerate(rest_rows):
    row = tbl_seg.add_row().cells
    bg = ROW_ALT_BG if idx % 2 == 0 else ROW_BG
    widths = [c[1] for c in seg_cols]
    for j, cell in enumerate(row):
        cell.width = Cm(widths[j])
        set_cell_bg(cell, bg)
        set_cell_borders(cell, 'BFBFBF')
    cell_para(row[0], param, bold=True, size=10, font='Courier New', color='1F497D')
    cell_para(row[1], val,   bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_para(row[2], unit,  italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_para(row[3], desc,  size=10)

# Sous-en-tête : Disparition des mains
add_subheader(tbl_seg, '  Disparition des mains (hors champ)', len(seg_cols))

hand_rows = [
    ('PAUSE_THRESHOLD', '0.8', 's', "Durée d'absence des mains déclenchant la finalisation de la phrase en cours"),
    ('HAND_GRACE', '25', 'trames', "Délai de grâce avant de considérer les mains comme perdues (~833 ms à 30 fps)"),
]

for idx, (param, val, unit, desc) in enumerate(hand_rows):
    row = tbl_seg.add_row().cells
    bg = ROW_ALT_BG if idx % 2 == 0 else ROW_BG
    widths = [c[1] for c in seg_cols]
    for j, cell in enumerate(row):
        cell.width = Cm(widths[j])
        set_cell_bg(cell, bg)
        set_cell_borders(cell, 'BFBFBF')
    cell_para(row[0], param, bold=True, size=10, font='Courier New', color='1F497D')
    cell_para(row[1], val,   bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_para(row[2], unit,  italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_para(row[3], desc,  size=10)

# Sous-en-tête : Timeout global
add_subheader(tbl_seg, '  Timeout global (inactivité)', len(seg_cols))

idle_rows = [
    ('IDLE_TIMEOUT', '3.0', 's', "Durée d'inactivité totale au-delà de laquelle la session WebSocket est considérée terminée"),
]

for idx, (param, val, unit, desc) in enumerate(idle_rows):
    row = tbl_seg.add_row().cells
    bg = ROW_ALT_BG
    widths = [c[1] for c in seg_cols]
    for j, cell in enumerate(row):
        cell.width = Cm(widths[j])
        set_cell_bg(cell, bg)
        set_cell_borders(cell, 'BFBFBF')
    cell_para(row[0], param, bold=True, size=10, font='Courier New', color='1F497D')
    cell_para(row[1], val,   bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_para(row[2], unit,  italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_para(row[3], desc,  size=10)

cap2 = add_paragraph(doc, '', space_before=4, space_after=16)
cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_c2 = cap2.add_run("Tableau 2.8b — Paramètres de segmentation de phrase (")
r_c2.italic = True; r_c2.font.size = Pt(9); r_c2.font.name = 'Calibri'
add_code_inline(cap2, 'inference_ws.py')
r_c2b = cap2.add_run(", lignes 33–91)")
r_c2b.italic = True; r_c2b.font.size = Pt(9); r_c2b.font.name = 'Calibri'

# ── Section 4 : Note sur SlidingWindowClassifier ──────────────────────────────
add_heading(doc, "Note — Rôle du SlidingWindowClassifier en mode CTC", level=2, size=12)

add_paragraph(doc,
    "Lorsque le modèle CTC est chargé, le SlidingWindowClassifier continue de "
    "s'exécuter à chaque trame, mais son rôle est réduit à deux fonctions "
    "auxiliaires : (1) calculer l'énergie de mouvement utilisée par la détection "
    "de repos, et (2) fournir une prédiction instantanée et le streak_progress "
    "pour le retour visuel de l'interface. Sa sortie de reconnaissance de signe "
    "n'est utilisée que si le moteur CTC n'est pas disponible (mode dégradé).",
    size=11, space_before=4, space_after=8)

# Box note
note_box = doc.add_paragraph()
note_box.paragraph_format.space_before = Pt(4)
note_box.paragraph_format.space_after  = Pt(16)
note_box.paragraph_format.left_indent  = Cm(0.8)
note_box.paragraph_format.right_indent = Cm(0.8)
r_note = note_box.add_run(
    "Les paramètres propres au SlidingWindowClassifier (WINDOW_SIZE, "
    "CONFIDENCE_THRESHOLD, STREAK_REQUIRED, etc.) définis dans constants.py "
    "ne figurent pas dans ce tableau : ils ne participent pas au chemin de "
    "reconnaissance principal en mode CTC. Ils restent configurables pour "
    "le mode dégradé."
)
r_note.italic = True
r_note.font.size = Pt(10)
r_note.font.name = 'Calibri'
r_note.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

# ── Section 5 : Flux de décision ──────────────────────────────────────────────
add_heading(doc, "8.3  Flux de décision simplifié", level=2, size=12)

add_paragraph(doc,
    "À chaque trame reçue, la pipeline effectue les opérations suivantes dans "
    "l'ordre :",
    size=11, space_before=4, space_after=6)

steps = [
    ("1.", "Extraction des caractéristiques",
     "Le vecteur de 171 flottants est validé (dimension, absence de NaN) et "
     "ajouté au tampon CTC ctc_buf."),
    ("2.", "Mise à jour du SlidingWindowClassifier",
     "Le classificateur à fenêtre glissante est alimenté pour calculer l'énergie "
     "de mouvement (motion_energy) et mettre à jour le retour visuel UI."),
    ("3.", "Détection de présence des mains",
     "Si aucune main n'est détectée, hand_miss_count est incrémenté. "
     "Passé HAND_GRACE = 25 trames (~833 ms), la phrase est finalisée."),
    ("4.", "Détection de repos",
     "Si les deux poignets dépassent REST_WRIST_Y = 0.8 et que "
     "motion_energy < REST_MOTION_MAX = 0.02 pendant REST_FRAMES = 8 trames "
     "consécutives, et si le dernier signe remonte à plus de REST_MIN_GAP = 0.3 s, "
     "la phrase est finalisée."),
    ("5.", "Décodage CTC progressif",
     "Toutes les CTC_STRIDE = 20 trames, si le tampon contient au moins "
     "CTC_MIN_FRAMES = 30 trames, le modèle BiLSTM-CTC décode la séquence "
     "et met à jour la liste de signes courants."),
    ("6.", "Finalisation et traduction",
     "Lors d'une limite de phrase (repos ou disparition), une passe CTC finale "
     "est effectuée, la séquence de signes est transmise au module seq2seq, "
     "et la phrase traduite est envoyée au client via WebSocket."),
]

for num, titre, desc in steps:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Cm(0.5)
    r_num = p.add_run(f"{num} ")
    r_num.bold = True; r_num.font.size = Pt(11); r_num.font.name = 'Calibri'
    r_num.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    r_tit = p.add_run(f"{titre} — ")
    r_tit.bold = True; r_tit.font.size = Pt(11); r_tit.font.name = 'Calibri'
    r_desc = p.add_run(desc)
    r_desc.font.size = Pt(11); r_desc.font.name = 'Calibri'

add_paragraph(doc, '', space_before=10, space_after=4)

# ── Pied de page ──────────────────────────────────────────────────────────────
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_p.paragraph_format.space_before = Pt(16)
r_f = footer_p.add_run(
    "Toutes les valeurs sont extraites de "
    "V11/server/routers/inference_ws.py · Aucune valeur de fallback incluse"
)
r_f.italic = True
r_f.font.size = Pt(9)
r_f.font.name = 'Calibri'
r_f.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(OUT_PATH)
print(f"Saved → {OUT_PATH}")
print(f"Size  → {os.path.getsize(OUT_PATH) / 1024:.1f} Ko")
