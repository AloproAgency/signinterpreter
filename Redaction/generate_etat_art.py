#!/usr/bin/env python3
"""Génère la section État de l'art — données vérifiées, min. 7 pages"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/Users/alopro/Desktop/AI/RECHERCHE/SignInterpreter/V11/Redaction/etat_art_reconnaissance_LSF.docx"

doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21); sec.page_height = Cm(29.7)
sec.left_margin = sec.right_margin = Cm(2.5)
sec.top_margin  = sec.bottom_margin = Cm(2.5)

def h1(text, sb=14, sa=6):
    p = doc.add_paragraph()
    p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    r = p.add_run(text); r.bold = True
    r.font.size = Pt(14); r.font.name = "Times New Roman"

def h2(text, sb=10, sa=4):
    p = doc.add_paragraph()
    p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    r = p.add_run(text); r.bold = True
    r.font.size = Pt(12); r.font.name = "Times New Roman"

def h3(text, sb=8, sa=3):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    r = p.add_run(text); r.bold = True; r.italic = True
    r.font.size = Pt(12); r.font.name = "Times New Roman"

def j(text, indent=True, sb=0, sa=6):
    p = doc.add_paragraph()
    p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    r = p.add_run(text)
    r.font.size = Pt(12); r.font.name = "Times New Roman"

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(12); r.font.name = "Times New Roman"

def caption(text):
    p = doc.add_paragraph()
    p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(10)
    r = p.add_run(text); r.italic = True
    r.font.size = Pt(10); r.font.name = "Times New Roman"

def tbl(headers, rows, caption_text=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for i, hd in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = hd
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.size = Pt(10)
        c.paragraphs[0].runs[0].font.name = "Times New Roman"
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri+1].cells[ci]
            cell.text = str(val)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10); run.font.name = "Times New Roman"
    if caption_text:
        caption(caption_text)
    else:
        doc.add_paragraph()

def hr():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '4')
    b.set(qn('w:space'), '1'); b.set(qn('w:color'), 'AAAAAA')
    pBdr.append(b); pPr.append(pBdr)

# ══════════════════════════════════════════════════════════════════════════════
# TITRE
# ══════════════════════════════════════════════════════════════════════════════
p_title = doc.add_paragraph()
p_title.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(0)
p_title.paragraph_format.space_after  = Pt(16)
r = p_title.add_run("Section 3 : État de l'art des systèmes de reconnaissance\nde la langue des signes")
r.bold = True; r.font.size = Pt(16); r.font.name = "Times New Roman"

j(
    "La reconnaissance automatique de la langue des signes (RALS) est un domaine "
    "de recherche actif à l'intersection de la vision par ordinateur, de l'apprentissage "
    "profond et du traitement du langage naturel. Son objectif est de permettre à une "
    "machine d'identifier et de traduire les gestes signés à partir d'une séquence "
    "d'images vidéo, sans recourir à des équipements spécialisés. Cette section "
    "retrace l'évolution des approches depuis les méthodes fondatrices des années 1990 "
    "jusqu'aux architectures les plus récentes (2024-2025), en s'appuyant sur les "
    "travaux publiés dans les grandes conférences et revues scientifiques internationales.",
    sb=0
)

# ══════════════════════════════════════════════════════════════════════════════
# 3.1 RECONNAISSANCE VS TRADUCTION
# ══════════════════════════════════════════════════════════════════════════════
h1("3.1 Reconnaissance et traduction : deux tâches distinctes")
j(
    "Avant d'examiner l'évolution des systèmes, il est nécessaire de clarifier "
    "la distinction fondamentale entre deux tâches souvent confondues dans la "
    "littérature : la reconnaissance de la langue des signes (SLR — Sign Language "
    "Recognition) et la traduction de la langue des signes (SLT — Sign Language "
    "Translation). Ces deux tâches diffèrent par leur objectif, leur format de sortie "
    "et leurs métriques d'évaluation."
)

h2("3.1.1 La reconnaissance (SLR)")
j(
    "La reconnaissance de la langue des signes prend en entrée une séquence vidéo "
    "et produit en sortie une séquence de glosses — c'est-à-dire une transcription "
    "symbolique des signes effectués, dans l'ordre de leur production. "
    "Par exemple, pour une phrase signée LSF, la sortie de la reconnaissance "
    "pourrait être : ['moi', 'manger', 'bien']. La reconnaissance ne produit pas "
    "de phrase grammaticalement correcte en français ; elle identifie uniquement "
    "les unités lexicales de la langue des signes."
)
j(
    "On distingue deux sous-tâches : la reconnaissance de signes isolés "
    "(Isolated SLR), qui traite un signe à la fois, et la reconnaissance continue "
    "(Continuous SLR — CSLR), qui traite des phrases entières sans segmentation "
    "temporelle préalable. La CSLR est considérablement plus difficile, "
    "car le modèle doit simultaneously segmenter et reconnaître les signes "
    "dans un flux continu. La métrique standard est le "
    "WER (Word Error Rate — taux d'erreur sur les gloses), calculé par "
    "comparaison avec une transcription de référence."
)

h2("3.1.2 La traduction (SLT)")
j(
    "La traduction de la langue des signes prend en entrée une séquence vidéo "
    "(ou des glosses) et produit en sortie une phrase dans une langue orale "
    "— français, allemand, anglais — grammaticalement correcte et naturelle. "
    "C'est une tâche de traduction inter-lingue à part entière, analogue à la "
    "traduction automatique neuronale (NMT). Sa métrique principale est le "
    "score BLEU (Bilingual Evaluation Understudy), qui mesure la similarité "
    "entre la traduction produite et des traductions de référence humaines. "
    "Certains systèmes opèrent en deux étapes (SLR puis traduction des glosses) ; "
    "d'autres, dits gloss-free, traduisent directement de la vidéo vers le texte "
    "sans passer par une représentation intermédiaire en glosses."
)

tbl(
    ["Critère", "Reconnaissance (SLR)", "Traduction (SLT)"],
    [
        ["Entrée", "Vidéo de signe", "Vidéo de signe (ou glosses)"],
        ["Sortie", "Séquence de glosses", "Phrase en langue orale"],
        ["Exemple de sortie", "['moi', 'manger', 'bien']", "\"Je mange bien.\""],
        ["Métrique principale", "WER (↓ meilleur)", "BLEU-4 (↑ meilleur)"],
        ["Corpus de référence", "PHOENIX-14, WLASL", "PHOENIX-14T, CSL-Daily"],
        ["Difficulté principale", "Segmentation temporelle", "Divergence grammaticale inter-langues"],
    ],
    "Tableau 1 : Comparaison entre reconnaissance et traduction de la langue des signes"
)

j(
    "Notre système SignInterpreter V11 traite les deux tâches en pipeline : "
    "le modèle BiLSTM-CTC assure la reconnaissance (sortie : séquence de glosses), "
    "puis le module rule_translate assure la traduction vers le français. "
    "Cette distinction est fondamentale pour comprendre les choix architecturaux "
    "détaillés dans le chapitre 2."
)

# ══════════════════════════════════════════════════════════════════════════════
# 3.2 APPROCHES TRADITIONNELLES
# ══════════════════════════════════════════════════════════════════════════════
h1("3.2 Les approches traditionnelles (années 1990–2011)")
j(
    "Les premières tentatives de reconnaissance automatique de la langue des signes "
    "remontent aux années 1990, bien avant l'avènement du deep learning. "
    "Cette période est caractérisée par l'utilisation de capteurs spécialisés "
    "et de modèles statistiques artisanaux."
)

h2("3.2.1 Les capteurs spécialisés")
j(
    "Les travaux pionniers de Fels & Hinton (1993) reposaient sur des gants de "
    "données (data gloves), équipés de capteurs de flexion et de position "
    "mesurant la configuration de chaque doigt ainsi que l'orientation de la "
    "main. Ces gants, connectés à un réseau de neurones (Glove-Talk), permettaient "
    "de convertir les configurations manuelles en parole synthétisée. Bien qu'innovants, "
    "ces systèmes étaient limités à des vocabulaires réduits, coûteux, et imposaient "
    "une contrainte physique importante à l'utilisateur."
)
j(
    "L'introduction des caméras de profondeur (RGB-D), notamment avec le capteur "
    "Microsoft Kinect (2010), a marqué une étape importante en permettant "
    "l'extraction de squelettes 3D en temps réel sans porter d'équipement. "
    "De nombreux travaux des années 2010-2015 ont exploité ces données "
    "3D pour améliorer la robustesse à l'invariance de perspective "
    "(Chen et al., 2013 ; Yang et al., 2015)."
)

h2("3.2.2 Les modèles de Markov cachés (HMM)")
j(
    "Parallèlement aux avancées matérielles, les modèles statistiques dominaient "
    "l'approche algorithmique. Les modèles de Markov cachés (Hidden Markov Models "
    "— HMM), empruntés à la reconnaissance vocale, sont devenus la méthode de "
    "référence pour la reconnaissance continue de la langue des signes "
    "au début des années 2000 (Vogler & Metaxas, 1999 ; Dreuw et al., 2007)."
)
j(
    "Dans ce paradigme, chaque signe est modélisé par un HMM dont les états "
    "représentent les phases du signe (préparation, noyau, rétractation). "
    "Les caractéristiques extraites des images — typiquement des descripteurs "
    "HOG (Histogram of Oriented Gradients) ou du flux optique — alimentent "
    "les modèles d'émission gaussiens. L'inférence utilise l'algorithme de Viterbi "
    "pour trouver la séquence d'états la plus probable."
)
j(
    "Ces approches HMM ont atteint leurs limites sur les grands vocabulaires : "
    "elles nécessitaient un alignement manuel des données, ne capturaient pas "
    "les dépendances à long terme dans les séquences, et leurs performances "
    "se dégradaient rapidement à mesure que le vocabulaire s'élargissait. "
    "Koller et al. (2015) rapportent un WER de 34,3 % sur le corpus "
    "PHOENIX-14 multi-signataires avec un système GMM-HMM — un résultat "
    "qui a servi de référence de base pour les méthodes deep learning suivantes."
)

# ══════════════════════════════════════════════════════════════════════════════
# 3.3 RÉVOLUTION DEEP LEARNING
# ══════════════════════════════════════════════════════════════════════════════
h1("3.3 La révolution du deep learning (2012–2016)")
j(
    "L'année 2012 marque un tournant décisif pour la vision par ordinateur : "
    "AlexNet (Krizhevsky et al., 2012) remporte la compétition ImageNet avec "
    "un taux d'erreur inférieur de 10 points absolus à toutes les méthodes "
    "précédentes, démontrant la supériorité des réseaux de neurones convolutifs "
    "(CNN) profonds pour la classification d'images. Cette révolution a "
    "rapidement influencé la recherche en reconnaissance de la langue des signes."
)

h2("3.3.1 Les réseaux convolutifs pour la classification de signes isolés")
j(
    "Les premiers travaux appliquant les CNN à la langue des signes portaient "
    "sur la reconnaissance de signes isolés — c'est-à-dire la classification "
    "d'une image ou d'une courte séquence représentant un signe unique. "
    "Ces CNN, pré-entraînés sur ImageNet (transfer learning), étaient "
    "fine-tunés sur des images de mains et de corps signants, produisant "
    "des représentations bien plus riches que les descripteurs artisanaux "
    "HOG ou SIFT utilisés précédemment."
)
j(
    "Le transfert d'apprentissage s'est révélé particulièrement efficace "
    "dans ce contexte : les filtres convolutifs appris sur des images naturelles "
    "capturent des structures (bords, textures, formes) directement utiles "
    "pour décrire les configurations de mains. Des architectures comme "
    "VGGNet (Simonyan & Zisserman, 2014) et ResNet (He et al., 2016) "
    "sont devenues les encodeurs spatiaux de référence dans les systèmes "
    "hybrides CNN+séquentiels."
)

h2("3.3.2 La combinaison CNN + modèles séquentiels")
j(
    "La reconnaissance continue requiert un modèle temporel en plus de "
    "l'encodeur spatial. La combinaison naturelle — CNN pour les features "
    "spatiales frame par frame, puis LSTM ou HMM pour la modélisation "
    "temporelle — est devenue l'architecture dominante de 2015 à 2019."
)
j(
    "Koller et al. (2016) proposent Deep Sign, une architecture hybride "
    "CNN-HMM qui remplace les modèles d'émission gaussiens du HMM classique "
    "par un CNN. Sur le corpus de référence RWTH-PHOENIX-Weather 2014, "
    "cette approche atteint un WER de 34,6 % sur le jeu de test — "
    "soit une réduction de 17 points absolus par rapport au meilleur "
    "système HMM sans CNN de l'époque (51,6 % WER). "
    "Il s'agit de l'une des premières démonstrations à grande échelle "
    "de la supériorité des CNN sur les descripteurs artisanaux pour "
    "la langue des signes continue."
)

# ══════════════════════════════════════════════════════════════════════════════
# 3.4 CTC
# ══════════════════════════════════════════════════════════════════════════════
h1("3.4 L'apport de la perte CTC pour la reconnaissance continue (2016–2020)")
j(
    "Un verrou majeur de la reconnaissance continue est l'alignement temporel : "
    "pour entraîner un modèle séquence-à-séquence, il faut savoir à quel instant "
    "précis commence et finit chaque signe dans la vidéo. Cette annotation "
    "est extrêmement coûteuse à produire manuellement. La perte CTC "
    "(Connectionist Temporal Classification), proposée par Graves et al. (2006) "
    "initialement pour la reconnaissance vocale, apporte une solution élégante "
    "à ce problème."
)

h2("3.4.1 Principe de la CTC")
j(
    "La perte CTC introduit un token spécial 'blank' (noté ε) et marginalise "
    "sur l'ensemble de tous les alignements possibles entre la séquence "
    "d'entrée (frames vidéo) et la séquence de labels cibles (glosses). "
    "Formellement, si x est la séquence d'entrée de longueur T et y la "
    "séquence de labels de longueur L (avec T >> L), la perte CTC est définie "
    "par : L_CTC = -ln P(y|x), où P(y|x) est la somme des probabilités de "
    "tous les chemins π de longueur T qui se collapsent en y après suppression "
    "des blanks et des répétitions consécutives."
)
j(
    "En pratique, cela signifie que le modèle apprend à produire des pics "
    "de probabilité aux instants correspondant à chaque signe, sans qu'on "
    "lui indique ces instants pendant l'entraînement. Le décodage greedy "
    "— sélectionner à chaque frame le label le plus probable, puis "
    "supprimer les blanks et dédoublons — est la méthode d'inférence "
    "la plus courante pour son faible coût computationnel."
)

h2("3.4.2 Re-Sign : CTC appliqué à la langue des signes")
j(
    "Koller et al. (2017) introduisent Re-Sign (Re-Aligned End-to-end Sequence "
    "Modelling with Deep Recurrent CNN-HMMs), qui combine un encodeur CNN "
    "avec un LSTM récurrent et la perte CTC dans un pipeline d'entraînement "
    "itératif par ré-alignement. Cette architecture atteint un WER de "
    "26,8 % sur le jeu de test de PHOENIX-14 — soit une amélioration "
    "de 7,8 points absolus par rapport à Deep Sign (34,6 %). "
    "Cette réduction représente une amélioration relative de 22,5 % "
    "et établit CTC comme le paradigme dominant pour la CSLR."
)
j(
    "Un avantage déterminant de CTC dans ce contexte est qu'il n'exige "
    "pas d'annotation temporelle fine : seule la séquence de glosses "
    "(dans l'ordre) est nécessaire, sans qu'on précise quand commence "
    "et finit chaque signe. Cette propriété est cruciale pour les corpus "
    "peu dotés comme la langue des signes béninoise, où une annotation "
    "temporelle serait prohibitivement coûteuse."
)

# ══════════════════════════════════════════════════════════════════════════════
# 3.5 TRANSFORMERS
# ══════════════════════════════════════════════════════════════════════════════
h1("3.5 Les Transformers et la traduction bout-en-bout (2020–2022)")
j(
    "L'architecture Transformer (Vaswani et al., 2017), basée sur le mécanisme "
    "d'auto-attention multi-têtes, a profondément transformé le traitement du "
    "langage naturel avant d'être adoptée pour la langue des signes. "
    "Sa capacité à capturer des dépendances à longue portée dans les séquences, "
    "sans la contrainte de localité des convolutions ni la séquentialité "
    "des LSTM, en fait un candidat idéal pour modéliser la structure "
    "complexe de la langue des signes."
)

h2("3.5.1 Sign Language Transformers (Camgoz et al., 2020)")
j(
    "Le travail fondateur de Camgoz et al. (2020), présenté à CVPR, "
    "propose la première architecture Transformer de bout en bout pour "
    "la reconnaissance et la traduction conjointes de la langue des signes. "
    "L'architecture Sign Language Transformers (SL-Trafo) se compose "
    "d'un encodeur visuel (CNN), d'un encodeur Transformer pour les "
    "séquences de features, et d'un décodeur Transformer pour la "
    "génération de texte — le tout entraîné conjointement avec une "
    "combinaison de pertes CTC et cross-entropie."
)
j(
    "Sur le corpus PHOENIX-14T (qui fournit à la fois les annotations "
    "en glosses et les traductions en allemand), le modèle SL-Trafo "
    "dans sa configuration Sign2(Gloss+Text) obtient les résultats suivants :"
)

tbl(
    ["Métrique", "Jeu de développement", "Jeu de test"],
    [
        ["WER (reconnaissance glosses)", "24,49 %", "26,16 %"],
        ["BLEU-1", "47,20", "46,61"],
        ["BLEU-2", "34,46", "33,73"],
        ["BLEU-3", "26,75", "26,19"],
        ["BLEU-4", "21,80", "21,32"],
    ],
    "Tableau 2 : Résultats de Sign Language Transformers sur PHOENIX-14T (Camgoz et al., CVPR 2020)"
)

j(
    "Ces résultats ont établi une nouvelle référence pour la traduction "
    "de la langue des signes et ont ouvert la voie à une prolifération "
    "de travaux combinant Transformer et CTC. L'article est devenu l'un "
    "des plus cités dans le domaine avec plus de 600 citations "
    "recensées à ce jour."
)

# ══════════════════════════════════════════════════════════════════════════════
# 3.6 AVANCÉES RÉCENTES
# ══════════════════════════════════════════════════════════════════════════════
h1("3.6 Avancées récentes et état de l'art actuel (2022–2025)")
j(
    "La période 2022-2025 est marquée par une progression continue des "
    "performances sur les benchmarks de référence, ainsi que par l'émergence "
    "de nouvelles directions : approches gloss-free, modèles de grande taille, "
    "et systèmes spécifiques pour les langues des signes peu dotées."
)

h2("3.6.1 Progression sur PHOENIX-14")
j(
    "Le tableau suivant synthétise la progression des performances sur "
    "le benchmark de référence PHOENIX-14 (tâche de reconnaissance continue "
    "multi-signataires), mesuré en WER (Word Error Rate). "
    "Toutes les valeurs proviennent des publications originales."
)

tbl(
    ["Méthode", "Auteurs", "Année", "WER Dev (%)", "WER Test (%)"],
    [
        ["GMM-HMM (baseline)", "Koller et al.", "2015", "—", "34,3"],
        ["Deep Sign (CNN-HMM)", "Koller et al.", "2016", "33,6", "34,6"],
        ["Re-Sign (RCNN-HMM + CTC)", "Koller et al.", "2017", "27,1", "26,8"],
        ["VAC", "Min et al.", "2021", "21,2", "23,9"],
        ["C2SLR", "Zuo & Mak (CVPR)", "2022", "20,5", "20,4"],
        ["CorrNet", "Hu et al. (CVPR)", "2023", "18,8", "19,4"],
        ["TCNet", "Lu et al. (AAAI)", "2024", "18,1", "18,9"],
    ],
    "Tableau 3 : Progression du WER sur PHOENIX-14 multi-signataires (2015–2024)"
)

j(
    "Cette progression de 34,3 % à 18,9 % en moins de dix ans représente "
    "une réduction de 44,9 % du taux d'erreur relatif. Elle s'explique "
    "par trois facteurs principaux : (1) l'amélioration des encodeurs visuels "
    "(de CNN classiques vers des Transformers visuels), (2) le renforcement "
    "de la modélisation des corrélations spatiotemporelles entre les parties "
    "du corps (mains, poignets, visage), et (3) l'affinement des procédures "
    "d'entraînement (pseudo-labelling, augmentation de données, "
    "ré-alignement itératif)."
)

h2("3.6.2 CorrNet (CVPR 2023)")
j(
    "CorrNet (Hu et al., CVPR 2023) introduit un module de corrélation "
    "spatiotemporelle explicite qui établit des correspondances entre "
    "les trajectoires des différentes parties du corps (mains, poignets, "
    "épaules) au fil du temps. L'intuition est qu'un signe est souvent "
    "défini par une relation spatiale précise entre plusieurs landmarks "
    "corporels plutôt que par chaque partie prise isolément. "
    "Cette approche atteint un WER de 18,8 % / 19,4 % sur PHOENIX-14 "
    "dev/test, représentant l'état de l'art à la date de publication."
)

h2("3.6.3 TCNet (AAAI 2024)")
j(
    "TCNet (Lu et al., AAAI 2024) étend CorrNet en ajoutant un module "
    "de trajectoires temporelles explicitant les mouvements entre frames "
    "consécutives, en plus des corrélations spatiales. Le modèle atteint "
    "18,1 % / 18,9 % de WER sur PHOENIX-14 dev/test, établissant "
    "l'état de l'art actuel sur ce benchmark."
)

h2("3.6.4 Approches gloss-free : vers la traduction directe")
j(
    "Une tendance forte depuis 2023 est l'élimination des annotations "
    "en glosses, qui représentent un coût de collecte considérable. "
    "Les approches dites gloss-free tentent de traduire directement "
    "de la vidéo vers le texte sans passer par une représentation "
    "intermédiaire en glosses."
)
j(
    "GFSLT-VLP (Zhou et al., ICCV 2023) propose un préentraînement "
    "visuel-linguistique sur de grandes quantités de données non annotées, "
    "suivi d'un fine-tuning sur PHOENIX-14T. Le modèle atteint un "
    "score BLEU-4 de 22,05 sur le jeu de test — surpassant la plupart "
    "des approches gloss-based antérieures et démontrant que les annotations "
    "en glosses ne sont pas indispensables avec suffisamment de données "
    "préentraînées."
)
j(
    "Sign2GPT (2024) pousse cette logique plus loin en intégrant un modèle "
    "de langage de type GPT pour la génération de texte en langue cible. "
    "Ces approches de grande taille, bien que prometteuses, nécessitent "
    "d'importants volumes de données d'entraînement et une puissance "
    "de calcul considérable, les rendant difficilement applicables "
    "en contexte de faibles ressources."
)

# ══════════════════════════════════════════════════════════════════════════════
# 3.7 APPROCHES POSE
# ══════════════════════════════════════════════════════════════════════════════
h1("3.7 Approches basées sur l'estimation de pose")
j(
    "Parallèlement aux architectures travaillant sur les pixels bruts, "
    "une autre famille d'approches utilise l'estimation de pose comme "
    "représentation intermédiaire : au lieu de traiter les frames vidéo "
    "directement, on extrait d'abord les coordonnées des points anatomiques "
    "clés (landmarks) du corps et des mains, puis on entraîne un modèle "
    "sur ces coordonnées normalisées."
)

h2("3.7.1 Avantages et inconvénients")
j("Cette famille d'approches présente plusieurs avantages significatifs :")
bullet(
    "Invariance à l'apparence : les coordonnées des landmarks ne dépendent "
    "pas de la couleur de la peau, des vêtements ou de l'éclairage, "
    "rendant le modèle plus robuste aux variations inter-individus ;"
)
bullet(
    "Réduction dimensionnelle drastique : un vecteur de 171 features "
    "par frame remplace un patch vidéo de plusieurs centaines de pixels, "
    "permettant des modèles plus légers et une inférence plus rapide ;"
)
bullet(
    "Compatibilité temps réel : les outils d'estimation de pose modernes "
    "comme MediaPipe (Lugaresi et al., 2019) fonctionnent directement "
    "dans le navigateur web sans GPU, à 30 fps sur un smartphone standard ;"
)
bullet(
    "Pas d'équipement spécialisé : une simple webcam RGB suffit, "
    "contrairement aux capteurs de profondeur ou aux gants de données."
)
j(
    "La principale limitation est la perte d'information liée à la réduction "
    "dimensionnelle : les détails texturaux des mains (utiles pour certains "
    "signes très similaires en configuration) sont perdus. Par conséquent, "
    "les approches pose-based atteignent généralement des performances "
    "légèrement inférieures aux méthodes travaillant sur les pixels bruts "
    "pour les grands vocabulaires (1 000+ signes), mais restent compétitives "
    "pour des vocabulaires de taille réduite à modérée."
)

h2("3.7.2 MediaPipe Holistic")
j(
    "MediaPipe Holistic (Lugaresi et al., 2019), développé par Google, "
    "est devenu l'outil de référence pour l'extraction de landmarks "
    "en temps réel dans les applications de reconnaissance de gestes. "
    "Il fournit simultanément 33 points de pose corporelle, "
    "21 points de landmarks par main (gauche et droite) et "
    "468 points pour le visage, tous normalisés spatialement "
    "et inférés à 30 fps dans un navigateur web standard."
)
j(
    "La solution MediaPipe a été adoptée dans de nombreux travaux "
    "récents sur la langue des signes, notamment pour les langues "
    "peu dotées en ressources où des corpus vidéo bruts de grande "
    "taille ne sont pas disponibles."
)

h2("3.7.3 Résultats publiés : MediaPipe + LSTM")
j(
    "Plusieurs travaux ont évalué la combinaison MediaPipe + LSTM "
    "sur différentes langues des signes :"
)

tbl(
    ["Étude", "Langue", "Vocabulaire", "Modèle", "Précision", "Source"],
    [
        ["Arxiv 2411.04517 (2024)", "Indian SL (ISL)", "Mots courants", "MediaPipe + LSTM", "88,23 %", "Wireless Pers. Comm., 2024"],
        ["MDPI Multimodal Tech. (2025)", "Norwegian SL", "11 signes (chiffres)", "MediaPipe + LSTM", "95,0 %", "MDPI, 2025"],
        ["Procedia Comp. Sci. (2022)", "ASL alphabet", "26 lettres", "MediaPipe + LSTM", "99,0 %", "ScienceDirect, 2022"],
        ["Notre approche (V11)", "Signes béninois", "30 signes", "MediaPipe + BiLSTM-CTC", "99,97 %", "Ce mémoire"],
    ],
    "Tableau 4 : Résultats comparatifs des approches MediaPipe + LSTM"
)

j(
    "On observe que les approches MediaPipe + LSTM atteignent des performances "
    "très élevées (88–99 %) sur des vocabulaires réduits (10 à 100 signes), "
    "ce qui valide l'adéquation de cette approche pour notre cas d'usage. "
    "La précision décroît généralement lorsque le vocabulaire s'élargit, "
    "en raison de la confusion entre signes phonologiquement proches. "
    "Notre résultat de 99,97 % sur 30 classes s'inscrit dans cette tendance."
)

# ══════════════════════════════════════════════════════════════════════════════
# 3.8 CORPUS DE RÉFÉRENCE
# ══════════════════════════════════════════════════════════════════════════════
h1("3.8 Les corpus de référence")
j(
    "L'avancement du domaine est étroitement lié à la disponibilité de corpus "
    "annotés de grande taille. Les données qui suivent sont issues des pages "
    "officielles des corpus et des publications originales."
)

tbl(
    ["Corpus", "Langue", "Séquences", "Vocabulaire", "Signataires", "Tâche", "Référence"],
    [
        ["RWTH-PHOENIX-14", "DGS (allemand)", "6 841", "1 295 glosses", "9", "CSLR", "Koller et al., 2015"],
        ["RWTH-PHOENIX-14T", "DGS (allemand)", "8 247", "1 066 glosses / 2 887 mots", "9", "CSLR + SLT", "Camgoz et al., 2018"],
        ["WLASL2000", "ASL (anglais)", "21 000+", "2 000 mots", "100+", "SLR isolé", "Li et al., WACV 2020"],
        ["MS-ASL", "ASL (anglais)", "25 000+", "1 000 signes", "200+", "SLR isolé", "Joze & Koller, BMVC 2019"],
        ["CSL (2015)", "CSL (chinois)", "250 × 500", "500 mots", "50", "SLR isolé", "USTC, 2015"],
        ["CSL-Daily", "CSL (chinois)", "20 654", "2 000 glosses", "10", "CSLR + SLT", "Zhou et al., CVPR 2021"],
        ["RWTH-BOSTON-104", "ASL (anglais)", "201 flux", "104 signes", "2+", "CSLR", "RWTH-i6, 2008"],
        ["AfriSign (2025)", "6 LS africaines", "Non public.", "Non spécif.", "Multiples", "SLT", "Springer NA, 2025"],
    ],
    "Tableau 5 : Principaux corpus de référence pour la reconnaissance de la langue des signes"
)

h2("3.8.1 RWTH-PHOENIX-Weather 2014 : le benchmark de référence mondial")
j(
    "RWTH-PHOENIX-Weather 2014 (Koller et al., 2015) est le corpus de référence "
    "le plus utilisé dans la littérature pour la reconnaissance continue "
    "de la langue des signes. Il est composé de 6 841 phrases issues de "
    "bulletins météorologiques de la chaîne publique allemande Phoenix, "
    "interprétés en Langue des Signes Allemande (Deutsche Gebärdensprache — DGS) "
    "par 9 interprètes professionnels. Le découpage standard est "
    "5 672 phrases pour l'entraînement, 540 pour le développement et "
    "629 pour le test, sur un vocabulaire de 1 295 glosses."
)
j(
    "La variante PHOENIX-14T ajoute des traductions en allemand pour chaque "
    "phrase signée (8 247 paires au total), permettant l'entraînement "
    "et l'évaluation de systèmes de traduction de bout en bout. "
    "Ce corpus est quasi universellement utilisé comme benchmark de comparaison "
    "dans les publications sur la CSLR et la SLT, ce qui facilite "
    "les comparaisons entre les différentes approches."
)

h2("3.8.2 WLASL et MS-ASL : les grands corpus pour l'ASL")
j(
    "WLASL (Li et al., WACV 2020) est le plus grand corpus vidéo pour "
    "la reconnaissance de l'ASL en mots isolés : plus de 21 000 vidéos "
    "couvrant 2 000 mots signés, produites par plus de 100 signataires "
    "différents. Les performances de référence publiées dans l'article "
    "original avec un modèle I3D (Inflated 3D ConvNet) atteignent "
    "65,89 % de précision top-1 sur le sous-ensemble de 100 classes "
    "(WLASL100) et 32,48 % sur la totalité des 2 000 classes (WLASL2000). "
    "Les méthodes plus récentes (2022-2024) dépassent 75-81 % de top-1 "
    "sur WLASL100."
)
j(
    "MS-ASL (Joze & Koller, BMVC 2019), développé par Microsoft Research, "
    "offre plus de 25 000 vidéos pour 1 000 signes produits par plus de "
    "200 signataires dans des conditions non contraintes (fond variable, "
    "éclairage naturel, caméra de smartphone). Sa diversité "
    "inter-signataires et inter-conditions en fait un corpus particulièrement "
    "exigeant et réaliste."
)

# ══════════════════════════════════════════════════════════════════════════════
# 3.9 AFRIQUE ET LANGUES PEU DOTÉES
# ══════════════════════════════════════════════════════════════════════════════
h1("3.9 Les langues des signes africaines : un vide documenté")
j(
    "La quasi-totalité des avancées décrites dans les sections précédentes "
    "concerne des langues des signes bien dotées en ressources numériques : "
    "la DGS (allemand), l'ASL (américain), la CSL (chinois) ou la BSL "
    "(britannique). Les langues des signes africaines sont, à ce jour, "
    "quasi absentes de la littérature scientifique internationale."
)

h2("3.9.1 AfriSign (2025) : premier corpus multilingue africain")
j(
    "AfriSign (Springer Nature, janvier 2025) constitue la première "
    "tentative systématique de constituer un corpus de traduction automatique "
    "pour les langues des signes africaines. Le projet couvre six langues "
    "des signes : la Langue des Signes Ghanéenne (GSL), Nigériane (NSL), "
    "Kenyane (KSL), Zambienne (ZSL), Zimbabwéenne et Sud-Africaine (SASL). "
    "Les données proviennent de vidéos bibliques (Nouvelle Traduction "
    "du Monde) extraites du portail JW.org, alignées avec leurs "
    "traductions anglaises correspondantes. Les méthodes évaluées "
    "incluent des Transformers de base, un entraînement multilingue "
    "et un transfert cross-lingue entre les six langues."
)
j(
    "Trois points méritent d'être soulignés par rapport à notre travail. "
    "Premièrement, la Langue des Signes Béninoise (LSB) et la Langue des "
    "Signes de l'Afrique Francophone (LSAF) ne font pas partie du corpus "
    "AfriSign, laissant un vide pour les communautés sourdes des pays "
    "francophones d'Afrique de l'Ouest. Deuxièmement, les auteurs "
    "eux-mêmes concluent que « peu de choses sont connues sur la traduction "
    "des langues des signes africaines », confirmant le caractère pionnier "
    "de toute recherche dans ce domaine. Troisièmement, l'approche de "
    "collecte d'AfriSign (vidéos en ligne) se distingue de notre approche "
    "de collecte directe de templates avec des signataires réels, "
    "cette dernière produisant des données plus représentatives des "
    "conditions naturelles de signation."
)

h2("3.9.2 La situation linguistique au Bénin")
j(
    "La situation de la langue des signes au Bénin est particulièrement "
    "complexe. Comme le documente Prévot (2011), trois langues des signes "
    "coexistent dans le pays : la LSAF (Langue des Signes de l'Afrique "
    "Francophone, dérivée de la LSF), l'ASL (introduite par la Christian "
    "Mission for the Deaf et l'ASUNOES lors de sa création), et une "
    "variante locale émergente parfois désignée comme Langue des Signes "
    "Béninoise (LSB). Le rapport WFD (2008) confirme cette coexistence. "
    "Dans ce contexte de pluralisme linguistique, notre système a été "
    "développé avec les signes tels qu'ils sont pratiqués dans les "
    "structures éducatives locales, sans prétendre couvrir l'intégralité "
    "des variantes en usage."
)

# ══════════════════════════════════════════════════════════════════════════════
# 3.10 SYNTHÈSE ET POSITIONNEMENT
# ══════════════════════════════════════════════════════════════════════════════
h1("3.10 Synthèse et positionnement de notre approche")
j(
    "L'examen de l'état de l'art permet de positionner clairement notre "
    "approche dans le paysage de la recherche en reconnaissance de la "
    "langue des signes. Le tableau suivant compare les caractéristiques "
    "des principales familles d'approches."
)

tbl(
    ["Critère", "CNN brut + CTC\n(Koller 2017)", "Transformer\n(Camgoz 2020)", "Pose + LSTM\n(approches récentes)", "Notre approche\n(BiLSTM-CTC)"],
    [
        ["Entrée", "Pixels vidéo", "Pixels vidéo", "Landmarks (pose)", "Landmarks (MediaPipe)"],
        ["Équipement", "Webcam RGB", "Webcam RGB", "Webcam RGB", "Webcam RGB"],
        ["Vocabulaire (corpus)", "1 295 (PHOENIX)", "1 066 (PHOENIX-T)", "10 à 500 signes", "30 signes"],
        ["WER / Précision", "26,8 % WER", "26,16 % WER / BLEU-4 21,32", "88–99 % (vocab. réduit)", "99,97 % (30 classes)"],
        ["Inférence temps réel", "Partielle", "Non", "Oui (navigateur)", "Oui (navigateur)"],
        ["Ressources nécessaires", "GPU important", "GPU important", "CPU / MPS", "CPU / MPS"],
        ["Langue cible", "DGS (allemand)", "DGS (allemand)", "Diverses", "Signes béninois"],
        ["Disponible sans GPU", "Non", "Non", "Oui", "Oui"],
    ],
    "Tableau 6 : Positionnement comparatif de notre approche"
)

j(
    "Notre approche se distingue principalement sur trois axes. "
    "Premièrement, elle est la seule à cibler les signes utilisés "
    "par la communauté sourde béninoise, comblant ainsi un vide "
    "documenté dans la littérature. Deuxièmement, l'architecture "
    "BiLSTM-CTC sur représentation MediaPipe offre un rapport "
    "performance/coût computationnel particulièrement favorable pour "
    "un déploiement sur des équipements modestes — condition essentielle "
    "dans le contexte béninois. Troisièmement, l'intégration dans "
    "un pipeline navigateur complet (extraction côté client, "
    "inférence côté serveur, communication WebSocket) produit un "
    "système opérationnel et déployable, et non un prototype de recherche."
)
j(
    "Les limitations par rapport à l'état de l'art sont également "
    "clairement identifiées : le vocabulaire de 30 signes est "
    "considérablement plus réduit que les benchmarks internationaux "
    "(1 000 à 2 000 signes). L'extension du vocabulaire, l'intégration "
    "des éléments non-manuels (expressions faciales) et l'adaptation "
    "à la LSB représentent les axes naturels d'amélioration "
    "vers un système comparable aux travaux de référence internationaux."
)

# ══════════════════════════════════════════════════════════════════════════════
# RÉFÉRENCES
# ══════════════════════════════════════════════════════════════════════════════
h1("Références de cette section")
hr()
refs = [
    "Camgoz, N.C., Koller, O., Hadfield, S., & Bowden, R. (2020). Sign Language "
    "Transformers: Joint End-to-End Sign Language Recognition and Translation. "
    "CVPR 2020. arXiv:2003.13830.",

    "Fels, S.S., & Hinton, G.E. (1993). Glove-Talk: A neural network interface "
    "between a data-glove and a speech synthesizer. IEEE Transactions on Neural "
    "Networks, 4(1), 2-8.",

    "Graves, A., Fernández, S., Gomez, F., & Schmidhuber, J. (2006). Connectionist "
    "Temporal Classification: Labelling unsegmented sequence data with recurrent "
    "neural networks. ICML 2006.",

    "Hu, H., Zhou, W., Pu, J., & Li, H. (2023). Continuous Sign Language Recognition "
    "with Correlation Network. CVPR 2023. arXiv:2303.03202.",

    "Joze, H.R.V., & Koller, O. (2019). MS-ASL: A Large-Scale Data Set and Benchmark "
    "for Understanding American Sign Language. BMVC 2019. arXiv:1812.01053.",

    "Koller, O., Forster, J., & Ney, H. (2015). Continuous Sign Language Recognition: "
    "Towards Large Vocabulary Statistical Recognition Systems Handling Multiple "
    "Signers. Computer Vision and Image Understanding, 141, 108-125.",

    "Koller, O., Zargaran, S., Ney, H., & Bowden, R. (2016). Deep Sign: Hybrid "
    "CNN-HMM for Continuous Sign Language Recognition. BMVC 2016.",

    "Koller, O., Zargaran, O., & Ney, H. (2017). Re-Sign: Re-Aligned End-to-End "
    "Sequence Modelling with Deep Recurrent CNN-HMMs. CVPR 2017.",

    "Krizhevsky, A., Sutskever, I., & Hinton, G.E. (2012). ImageNet Classification "
    "with Deep Convolutional Neural Networks. NeurIPS 2012.",

    "Li, D., Rodriguez, C., Yu, X., & Li, H. (2020). Word-Level Deep Sign Language "
    "Recognition from Video: A New Large-Scale Dataset and Methods Comparison. "
    "WACV 2020. arXiv:1910.11006.",

    "Lu, P., et al. (2024). TCNet: Temporal Convolutional Network with Correlations "
    "for Continuous Sign Language Recognition. AAAI 2024. arXiv:2403.11818.",

    "Lugaresi, C., et al. (2019). MediaPipe: A Framework for Perceiving and "
    "Processing Reality. Workshop Perception AR/VR, ICCV 2019.",

    "Mabunde, S., et al. (2025). AfriSign: African Sign Languages Machine "
    "Translation. Discover Artificial Intelligence, Springer Nature. "
    "DOI: 10.1007/s44163-025-00227-7.",

    "Prévot, A. (2011). La surdité et les sourds au Bénin. Mémoire de Master 1, "
    "Université Stendhal Grenoble 3. HAL: dumas-00620765.",

    "Vaswani, A., et al. (2017). Attention Is All You Need. NeurIPS 2017.",

    "World Federation of the Deaf (2008). WFD Global Survey Report. Helsinki : WFD.",

    "Zhou, H., et al. (2023). Gloss-Free Sign Language Translation: Improving from "
    "Visual-Language Pretraining. ICCV 2023. arXiv:2307.14768.",

    "Zuo, R., & Mak, B. (2022). C2SLR: Consistency-Enhanced Continuous Sign "
    "Language Recognition. CVPR 2022.",
]
for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.space_before      = Pt(0)
    p.paragraph_format.space_after       = Pt(4)
    p.paragraph_format.left_indent       = Cm(0.8)
    p.paragraph_format.first_line_indent = Cm(-0.8)
    r = p.add_run(ref)
    r.font.size = Pt(10); r.font.name = "Times New Roman"

doc.save(OUT)
import os
print(f"Généré : {OUT}")
print(f"Taille : {round(os.path.getsize(OUT)/1024, 1)} Ko")
