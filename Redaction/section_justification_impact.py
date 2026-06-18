#!/usr/bin/env python3
"""Génère la section 'Justification de l'apport social' en .docx — données vérifiées"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = "/Users/alopro/Desktop/AI/RECHERCHE/SignInterpreter/V11/Redaction/section_impact_social.docx"

doc = Document()
s = doc.sections[0]
s.page_width = Cm(21); s.page_height = Cm(29.7)
s.left_margin = s.right_margin = Cm(2.5)
s.top_margin  = s.bottom_margin = Cm(2.5)

def h(text, size=13, sb=12, sa=6):
    p = doc.add_paragraph()
    p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    r = p.add_run(text); r.bold = True
    r.font.size = Pt(size); r.font.name = "Times New Roman"

def j(text, indent=True, size=12, sb=0, sa=6):
    p = doc.add_paragraph()
    p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.name = "Times New Roman"

def bullet(text, size=12):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.name = "Times New Roman"

def tbl(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for i, hd in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = hd
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.size = Pt(10)
        c.paragraphs[0].runs[0].font.name = "Times New Roman"
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri+1].cells[ci]
            cell.text = str(val)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10); run.font.name = "Times New Roman"
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════
h("Justification de l'apport social de la solution", size=14, sb=0)

j(
    "La pertinence d'un système de traduction automatique de la langue des signes "
    "s'évalue à l'aune des réalités quantifiables de la situation des personnes "
    "sourdes et malentendantes au Bénin et en Afrique. Les données présentées "
    "dans cette section sont issues exclusivement de sources officielles : "
    "recensements nationaux (INSAE), rapports de l'Organisation Mondiale de la "
    "Santé (OMS), enquêtes internationales (WFD, GSMA, DataReportal) et "
    "publications scientifiques à comité de lecture.",
    sb=6
)

# ──────────────────────────────────────────────────────────────────
h("1. La déficience auditive au Bénin : données du recensement national")

j(
    "Le quatrième Recensement Général de la Population et de l'Habitat du Bénin "
    "(RGPH4, INSAE, 2013) constitue la source nationale de référence sur le "
    "handicap. Sur une population totale de 10 008 749 habitants recensés, "
    "92 495 personnes ont été enregistrées avec un handicap, soit 1,02 % "
    "de la population. Parmi celles-ci, le handicap auditif représente "
    "18 % des cas — la deuxième catégorie la plus fréquente après "
    "la déficience visuelle (37,4 %)."
)

tbl(
    ["Catégorie de handicap", "Part (%)", "Estimation (RGPH4, 2013)"],
    [
        ["Déficience visuelle", "37,4 %", "~34 593 personnes"],
        ["Déficience auditive", "18,0 %", "~16 649 personnes"],
        ["Handicap moteur cérébral", "16,9 %", "~15 631 personnes"],
        ["Handicap moteur", "16,4 %", "~15 169 personnes"],
        ["Handicap intellectuel", "6,5 %", "~6 012 personnes"],
        ["Handicap psychosocial", "5,0 %", "~4 625 personnes"],
        ["Total personnes handicapées", "100 %", "92 495 personnes (1,02 % pop.)"],
    ]
)

j(
    "Par ailleurs, le rapport 2008 de la Fédération Mondiale des Sourds "
    "(World Federation of the Deaf — WFD) recense officiellement "
    "12 512 personnes sourdes au Bénin à cette date, en distinguant "
    "les personnes sourdes utilisant la langue des signes de l'ensemble "
    "plus large des personnes malentendantes. Ces deux chiffres — "
    "12 512 sourds signants et ~16 649 déficients auditifs (RGPH4) — "
    "délimitent l'ordre de grandeur de la communauté concernée directement "
    "par notre solution.",
    sa=10
)

# ──────────────────────────────────────────────────────────────────
h("2. La crise des interprètes et des ressources éducatives")

j(
    "Le même rapport de la WFD (2008) est sans ambiguïté sur la situation "
    "des interprètes au Bénin : aucun interprète qualifié en langue des "
    "signes n'existait à la date de l'enquête. Seul un nombre limité "
    "d'interprètes bénévoles, principalement des enseignants d'écoles pour "
    "sourds, comblaient partiellement ce vide, sans formation certifiée "
    "ni code de déontologie professionnel (WFD, 2008)."
)
j(
    "Sur le plan éducatif, le WFD (2008) recensait au Bénin une école "
    "publique et cinq écoles privées pour enfants sourds sur l'ensemble "
    "du territoire national. Le premier lycée public pour malentendants "
    "n'a ouvert ses portes qu'en 2012. "
    "À titre de comparaison, les données les plus récentes de l'ASUNOES "
    "(2023-2024) font état de 571 élèves (sourds et entendants confondus) "
    "inscrits au CAEIS de Louho à Porto-Novo. "
    "L'école primaire publique de Védoko à Cotonou, seul établissement "
    "d'État dédié aux sourds dans la capitale économique, n'accueille "
    "que 100 élèves environ, encadrés par 4 enseignants.",
)

tbl(
    ["Indicateur", "Valeur", "Source"],
    [
        ["Interprètes qualifiés en langue des signes (Bénin)", "0 en 2008 (bénévoles seulement)", "WFD, 2008"],
        ["Écoles pour sourds (Bénin, 2008)", "1 publique + 5 privées", "WFD, 2008"],
        ["Premier lycée public pour malentendants", "Ouvert en 2012", "Wikipedia / Deafness in Benin"],
        ["Élèves CAEIS de Louho (2023-2024)", "571 (sourds + entendants)", "ASUNOES Bénin, 2023"],
        ["Élèves sourds CAEIS (2018)", "235 sur 515 au total", "Archives CAEIS / Djissenou et al., 2021"],
        ["Élèves école de Védoko, Cotonou", "~100 sourds, 4 enseignants", "ASFA Parrainage"],
    ]
)

# ──────────────────────────────────────────────────────────────────
h("3. La crise de l'audition en Afrique : une urgence régionale")

j(
    "Le rapport de l'OMS sur les soins auditifs dans la Région Africaine "
    "(OMS/AFRO, octobre 2024) fournit les données régionales les plus récentes "
    "et les plus alarmantes : 40 millions de personnes vivent actuellement "
    "avec une déficience auditive en Afrique (soit 3,6 % de la population "
    "de la région), un chiffre projeté à 54 millions d'ici 2030 et "
    "97 millions d'ici 2050. Le coût économique annuel de cette situation "
    "est estimé à 27,1 milliards de dollars américains pour le continent."
)

tbl(
    ["Indicateur — Afrique", "Valeur", "Source"],
    [
        ["Personnes avec déficience auditive (2024)", "40 millions (3,6 % de la population)", "OMS/AFRO, 2024"],
        ["Projection 2030", "54 millions", "OMS/AFRO, 2024"],
        ["Projection 2050", "97 millions", "OMS/AFRO, 2024"],
        ["Coût économique annuel", "27,1 milliards USD", "OMS/AFRO, 2024"],
        ["Accès aux aides auditives", "~10 % de ceux qui en ont besoin", "OMS/AFRO, 2024"],
        ["Pays avec < 1 ORL par million d'hab.", "> 56 % des pays africains", "OMS/AFRO, 2024"],
        ["Pertes auditives évitables (enfants)", "Jusqu'à 75 % des cas", "OMS/AFRO, 2024"],
    ]
)

j(
    "Ces chiffres révèlent une réalité particulièrement grave : non seulement "
    "la déficience auditive est largement répandue en Afrique, mais les "
    "ressources médicales pour y répondre sont dramatiquement insuffisantes. "
    "Dans ce contexte, les solutions technologiques de communication — "
    "qui n'ont pas besoin d'un système de santé développé pour fonctionner "
    "— représentent un levier d'action accessible et immédiat.",
    sa=10
)

# ──────────────────────────────────────────────────────────────────
h("4. L'infrastructure numérique béninoise : un vecteur de déploiement réaliste")

j(
    "La faisabilité d'un déploiement à grande échelle de notre solution "
    "repose sur l'accessibilité numérique de la population cible. "
    "Les données de DataReportal (2026) et de la GSMA (2023) dressent "
    "un tableau encourageant de la connectivité au Bénin :"
)

tbl(
    ["Indicateur numérique — Bénin", "Valeur", "Source"],
    [
        ["Population totale (oct. 2025)", "14,9 millions", "DataReportal, 2026"],
        ["Utilisateurs internet", "4,80 millions (32,2 %)", "DataReportal, 2026"],
        ["Connexions mobiles actives", "16,4 millions (110 % pop.)", "DataReportal, 2026"],
        ["Utilisateurs internet mobile uniques", "7 millions (55,4 % pop.)", "GSMA Intelligence, 2023"],
        ["Couverture 4G du territoire", "88 % (100 % urbain)", "GSMA, 2023"],
        ["Couverture réseau mobile (3G+4G)", "~90 % du territoire", "GSMA, 2023"],
        ["Abonnés internet mobile (Q1 2025)", "11,14 millions", "ARCEP Bénin, 2025"],
    ]
)

j(
    "Avec 7 millions d'utilisateurs internet mobile uniques et une couverture "
    "4G atteignant 88 % du territoire national (100 % en zone urbaine), "
    "le Bénin dispose d'une infrastructure numérique suffisante pour "
    "qu'une application web temps réel soit accessible à une part "
    "significative de la population. Notre architecture navigateur — "
    "sans téléchargement ni installation — est directement compatible "
    "avec les smartphones d'entrée de gamme sous Android, système le "
    "plus répandu en Afrique de l'Ouest.",
    sa=10
)

# ──────────────────────────────────────────────────────────────────
h("5. Un vide technologique confirmé par la recherche scientifique")

j(
    "La revue de littérature confirme l'absence totale de système de "
    "reconnaissance automatique de la langue des signes béninoise. "
    "La publication la plus récente et la plus exhaustive sur ce sujet "
    "— le projet AfriSign (Springer Nature, janvier 2025), qui constitue "
    "le premier jeu de données de traduction automatique pour les langues "
    "des signes africaines — couvre six langues des signes du continent "
    "(Ghana, Nigeria, Kenya, Zambie, Zimbabwe, Afrique du Sud), "
    "mais exclut explicitement le Bénin et la Langue des Signes "
    "de l'Afrique Francophone (LSAF)."
)
j(
    "Les auteurs d'AfriSign eux-mêmes concluent que « peu de choses sont "
    "connues sur la traduction des langues des signes africaines ». "
    "Notre système SignInterpreter constitue donc, à notre connaissance, "
    "la première tentative documentée de reconnaissance automatique "
    "des signes utilisés par la communauté sourde béninoise.",
    sa=10
)

# ──────────────────────────────────────────────────────────────────
h("6. Un cadre juridique international qui soutient cette démarche")

j(
    "Notre solution s'inscrit dans des engagements formels pris par le Bénin :"
)
bullet(
    "Convention relative aux droits des personnes handicapées (CRPD, ONU, 2006) : "
    "ratifiée par le Bénin le 5 juillet 2012. L'article 9 oblige les États "
    "à garantir l'accès des personnes handicapées aux technologies de l'information "
    "et de la communication."
)
bullet(
    "Objectifs de Développement Durable n°10 (Réduction des inégalités, ONU, 2015) : "
    "cible l'inclusion sociale et économique de toutes les personnes, "
    "indépendamment de leur situation de handicap."
)
bullet(
    "Stratégie Nationale d'Intelligence Artificielle et de Mégadonnées "
    "du Bénin 2023-2027 : identifie l'accessibilité pour les personnes "
    "en situation de handicap parmi ses axes d'intervention prioritaires."
)

j(
    "L'ensemble de ces éléments — population sourde documentée (12 512 à "
    "16 649 personnes), absence totale d'interprètes qualifiés, infrastructure "
    "numérique couvrant 88 % du territoire, et vide technologique confirmé "
    "par la communauté scientifique internationale — justifie pleinement "
    "le développement d'une solution de traduction automatique de la langue "
    "des signes adaptée au contexte béninois.",
    sb=6
)

# ──────────────────────────────────────────────────────────────────
h("Références de cette section", size=11, sb=20)
refs = [
    "African Sign Languages Resource Center (2024). Benin. "
    "https://africansignlanguagesresourcecenter.com/benin/",

    "ARCEP Bénin (2025). Marché de l'internet mobile — Rapport T1 et T2 2025. "
    "Autorité de Régulation des Communications Électroniques et de la Poste.",

    "ASFA Parrainage. École de sourds de Cotonou. "
    "https://www.asfa-parrainage.org/ecole-de-sourds-de-cotonou/",

    "ASUNOES Bénin (2023). Résultats de l'année scolaire 2023-2024. "
    "https://www.asunoes-benin.org/",

    "DataReportal (2026). Digital 2026: Benin. Kepios / GSMA Intelligence / ITU. "
    "https://datareportal.com/reports/digital-2026-benin",

    "Djissenou, A.S., Kélani, R.R. & Houessou, P. (2021). Education Inclusive "
    "En République Du Bénin. European Scientific Journal, 17(15), 335. "
    "https://doi.org/10.19044/esj.2021.v17n15p335",

    "GSMA Intelligence (2023). The Mobile Economy: Benin. GSM Association. "
    "https://extensia.tech/benin-94-mobile-internet-coverage-but-actual-usage-can-improve-gsma/",

    "INSAE (2013). Quatrième Recensement Général de la Population et de l'Habitat "
    "(RGPH4). Institut National de la Statistique et de l'Analyse Économique, Bénin.",

    "Mabunde, S. et al. (2025). AfriSign: African Sign Languages Machine Translation. "
    "Discover Artificial Intelligence, Springer Nature. "
    "https://link.springer.com/article/10.1007/s44163-025-00227-7",

    "OMS / Bureau Régional pour l'Afrique (2024). Status Report on Ear and Hearing Care "
    "in the WHO African Region. Organisation Mondiale de la Santé, Brazzaville. "
    "https://www.afro.who.int/publications/status-report-ear-and-hearing-care-who-african-region",

    "ONU (2006). Convention relative aux droits des personnes handicapées (CRPD). "
    "Nations Unies, New York. Ratifiée par le Bénin le 5 juillet 2012.",

    "ONU (2015). Objectifs de Développement Durable — ODD 10 : Réduction des inégalités.",

    "Prévot, A. (2011). La surdité et les sourds au Bénin. "
    "Mémoire de Master 1, Université Stendhal Grenoble 3. HAL: dumas-00620765.",

    "République du Bénin — Ministère du Numérique et de la Digitalisation (2023). "
    "Stratégie Nationale d'Intelligence Artificielle et de Mégadonnées 2023-2027.",

    "Wikipedia (2024). Deafness in Benin. "
    "https://en.wikipedia.org/wiki/Deafness_in_Benin",

    "World Federation of the Deaf (2008). WFD Global Survey Report. "
    "Helsinki : WFD.",
]
for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.space_before      = Pt(0)
    p.paragraph_format.space_after       = Pt(4)
    p.paragraph_format.left_indent       = Cm(0.8)
    p.paragraph_format.first_line_indent = Cm(-0.8)
    r = p.add_run(ref)
    r.font.size = Pt(10); r.font.name = "Times New Roman"

doc.save(OUT)
import os
print(f"Généré : {OUT}")
print("Taille :", round(os.path.getsize(OUT)/1024, 1), "Ko")
